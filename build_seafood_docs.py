from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import seafood_project_data as d


ROOT = Path("/Volumes/WD_BLACK/汾酒尼泊尔")
OUT = ROOT / "尼泊尔海鲜AI线上销售系统"
ASSETS = ROOT / "_qa_seafood_assets"
FONT = "Hiragino Sans GB"
FONT_PATH = "/Users/fan/.local/share/fonts/codex-report-fonts/NotoSansCJKsc-Regular.otf"
BLUE = "1F4D78"
BLUE_2 = "2E74B5"
INK = "1F2937"
MUTED = "667085"
PALE = "EDF4FA"
PALE_GOLD = "FFF4CE"
PALE_RED = "FDECEC"
PALE_GREEN = "EAF6EF"
TABLE_WIDTH = 9360


def clean_text(value) -> str:
    if value is None:
        return ""
    return str(value).replace("客户", "合作对象").replace("字段", "信息项").replace("用户", "使用方")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_run_font(run, size=11, bold=None, color=INK, italic=None):
    run.font.name = FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str, placeholder: str = "1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])


def add_hyperlink(paragraph, label: str, url: str):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), FONT)
    r_pr.extend([fonts, color, underline])
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = clean_text(label)
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_table_geometry(table, proportions: Sequence[float]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(proportions)
    widths = [max(400, round(TABLE_WIDTH * p / total)) for p in proportions]
    widths[-1] += TABLE_WIDTH - sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def new_doc(title: str, subtitle: str, compact=False) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.4)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT if compact else WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6 if compact else 8)
    normal.paragraph_format.line_spacing = 1.25 if compact else 1.333
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE_2, 18, 10),
        ("Heading 2", 13, BLUE, 12 if not compact else 14, 6 if not compact else 7),
        ("Heading 3", 12, BLUE, 8 if not compact else 10, 4 if not compact else 5),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("尼泊尔海鲜项目｜合作执行资料")
    set_run_font(run, size=8, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{d.ACCESS_DATE}  |  ")
    set_run_font(run, size=8, color=MUTED)
    add_field(footer, "PAGE")
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    cover = doc.add_paragraph()
    cover.paragraph_format.space_before = Pt(56 if not compact else 32)
    cover.paragraph_format.space_after = Pt(12)
    run = cover.add_run(clean_text(title))
    set_run_font(run, size=25 if not compact else 22, bold=True, color=BLUE)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(20)
    run = sub.add_run(clean_text(subtitle))
    set_run_font(run, size=12, color=MUTED)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta.paragraph_format.space_after = Pt(18)
    for label, value in (
        ("研究日期", d.ACCESS_DATE),
        ("当前状态", "部分完成：公开研究已完成，真实商品与本地执行条件待核"),
        ("判断方式", "已确认 / 部分成立 / 待验证 / 推测"),
        ("版本", "v1.0 合作讨论稿"),
    ):
        run = meta.add_run(f"{label}：")
        set_run_font(run, size=9, bold=True)
        run = meta.add_run(f"{value}\n")
        set_run_font(run, size=9, color=MUTED)
    add_callout(doc, "使用提醒", "本文件用于合作讨论和首阶段验证，不替代尼泊尔当地法律、报关、食品许可、税务或冷链专业意见。", PALE_GOLD)
    doc.add_page_break()
    return doc


def add_para(doc: Document, text: str, label: str | None = None, color=INK):
    p = doc.add_paragraph()
    if label:
        run = p.add_run(clean_text(label))
        set_run_font(run, bold=True, color=BLUE)
    run = p.add_run(clean_text(text))
    set_run_font(run, color=color)
    return p


def add_callout(doc: Document, label: str, text: str, fill=PALE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(9)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    run = p.add_run(f"{clean_text(label)}  ")
    set_run_font(run, size=10.5, bold=True, color=BLUE)
    run = p.add_run(clean_text(text))
    set_run_font(run, size=10.5)
    return p


def add_step_list(doc: Document, items: Iterable[str], start=1):
    rows = [[str(index), clean_text(item)] for index, item in enumerate(items, start=start)]
    return add_table(doc, ["顺序", "说明"], rows, [0.6, 5.9], 9.2)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Iterable],
    proportions: Sequence[float] | None = None,
    font_size=8.4,
    header_fill=BLUE,
):
    matrix = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(clean_text(header))
        set_run_font(run, size=font_size, bold=True, color="FFFFFF")
    for row_index, row in enumerate(matrix):
        cells = table.add_row().cells
        for index in range(len(headers)):
            value = row[index] if index < len(row) else ""
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if isinstance(value, tuple) and len(value) == 2 and str(value[1]).startswith("http"):
                add_hyperlink(p, clean_text(value[0]), str(value[1]))
            else:
                run = p.add_run(clean_text(value))
                set_run_font(run, size=font_size)
            if row_index % 2 == 1:
                set_cell_shading(cells[index], "F8FAFC")
    if proportions is None:
        proportions = [1] * len(headers)
    set_table_geometry(table, proportions)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_picture(doc: Document, path: Path, caption: str, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(clean_text(caption))
    set_run_font(run, size=8.5, color=MUTED, italic=True)


def add_toc(doc: Document, sections: Sequence[str]):
    doc.add_heading("阅读目录", level=1)
    rows = [[str(i), clean_text(name)] for i, name in enumerate(sections, 1)]
    add_table(doc, ["章节", "内容"], rows, [0.7, 5.8], 9.4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("可更新目录：")
    set_run_font(run, size=9, bold=True, color=MUTED)
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "在 Word/WPS 中更新目录")


def add_basis(doc: Document):
    doc.add_heading("判断口径", level=1)
    add_table(
        doc,
        ["标记", "含义", "如何使用"],
        [
            ["已确认", "有官方、平台官方、企业官网或公开统计支持", "可以作为当前讨论依据；动态页面在行动前再看一次"],
            ["部分成立", "核心方向有证据，但地区、账号、商品或执行条件尚未全部核实", "只用于试点，不用于做强承诺"],
            ["待验证", "缺少本地主体、报关、商品、库存、价格、配送或账号实测", "由责任方补证后才能推进"],
            ["推测", "从多项证据形成的工作假设", "用小规模实验检验，不当作市场事实"],
        ],
        [1.05, 2.5, 2.95],
        8.8,
    )


def source_map():
    return {source["id"]: source for source in d.SOURCES}


SOURCE_BY_ID = source_map()


def selected_sources(ids: str | Sequence[str] | None = None):
    if ids is None:
        return d.SOURCES
    if isinstance(ids, str):
        ids = [part.strip() for part in ids.split(",") if part.strip()]
    return [SOURCE_BY_ID[item] for item in ids if item in SOURCE_BY_ID]


def add_sources(doc: Document, ids: str | Sequence[str] | None = None, title="来源与引用"):
    sources = selected_sources(ids)
    doc.add_heading(title, level=1)
    add_callout(
        doc,
        "阅读方法",
        "来源按原始页面记录。官方与平台官方来源优先；企业网页用于证明在售、价格和渠道存在，不代表全部市场。动态价格、库存与平台能力在行动前需要重新核对。",
    )
    rows = []
    for source in sources:
        rows.append(
            [
                source["id"],
                (source["title"], source["url"]),
                source["institution"],
                source["year"],
                source["type"],
                source["use"],
                source["confidence"],
            ]
        )
    add_table(doc, ["编号", "来源", "机构", "时间", "类别", "用途", "可信度"], rows, [0.55, 1.65, 1.3, 0.75, 0.85, 1.45, 0.55], 7.2)


def add_open_items(doc: Document):
    doc.add_heading("需要合作各方补充的关键信息", level=1)
    add_table(
        doc,
        ["优先级", "需要补充", "由谁确认", "未确认时的处理"],
        [
            ["P0", "尼泊尔进口主体、有效 EXIM Code、DFTQC 许可路径", "尼泊尔进口方、报关行", "不发货、不承诺到货"],
            ["P0", "每个商品的准确海关编码、当期税费与所需证明", "报关行、当地合规顾问", "不做落地成本和正式报价"],
            ["P0", "真实商品、规格、产地、可供价格、库存和有效期", "中国供应方", "只展示候选方向，不对外销售"],
            ["P0", "仓储温控、运输方式、配送区域、时效与异常处理", "进口方、仓储与配送方", "只做资料准备，不接单"],
            ["P1", "包装语言、过敏提示、净重、批号与日期排版", "进口方、食品合规顾问", "不上架"],
            ["P1", "收款主体、Fonepay/eSewa/Khalti/COD 接入条件", "尼泊尔经营主体、支付服务方", "保留人工收款演练"],
            ["P1", "TikTok 实际账号的 LIVE、广告、地区定向后台", "账号负责人、平台代表", "只做自然内容草稿"],
            ["P1", "首批样品、试单预算、人员分工和回款规则", "项目负责人、合作各方", "不扩大外联"],
        ],
        [0.55, 2.65, 1.65, 1.65],
        8.1,
    )


def add_next_steps(doc: Document, items: Sequence[str]):
    doc.add_heading("下一步", level=1)
    add_step_list(doc, items)


def save(doc: Document, filename: str, subdir: str | None = None):
    existing = {clean_text(p.text).strip() for p in doc.paragraphs}
    if not any("需要合作各方补充" in text or "执行前确认" in text for text in existing):
        doc.add_heading("执行前确认", level=1)
        add_table(
            doc,
            ["事项", "当前状态", "行动"],
            [
                ["真实商品、价格、库存与有效期", "待验证", "由中国供应方提供并持续更新"],
                ["进口许可、海关编码、税费与标签", "待验证", "由尼泊尔进口方和报关行逐品复核"],
                ["冷链、配送、支付与售后", "待验证", "在加德满都谷地完成真实演练"],
                ["真实联系、投放、订单与复购", "未执行", "人工批准后小范围开展"],
            ],
            [2.6, 1.1, 2.8],
            8.5,
        )
    if not any(text in {"下一步", "验证计划", "30天安排", "对接顺序"} for text in existing):
        add_next_steps(
            doc,
            [
                "先补齐与本文件直接相关的 P0 信息。",
                "由尼泊尔当地责任方复核许可、税费、标签和执行条件。",
                "只在加德满都谷地做小范围真实验证。",
                "把结果写回 CRM，再决定继续、调整或暂停。",
            ],
        )
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        paragraph = doc.paragraphs[-1]._element
        paragraph.getparent().remove(paragraph)
    target_dir = OUT / subdir if subdir else OUT
    target_dir.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = filename.rsplit(".", 1)[0]
    doc.core_properties.subject = "中国供应链向尼泊尔销售海鲜产品的合作执行资料"
    doc.core_properties.author = "项目研究与规划工作组"
    doc.core_properties.comments = "公开研究与试点方案；未确认事项保持待验证。"
    doc.save(target_dir / filename)


def create_charts():
    ASSETS.mkdir(parents=True, exist_ok=True)
    title_font = ImageFont.truetype(FONT_PATH, 34)
    label_font = ImageFont.truetype(FONT_PATH, 24)
    small_font = ImageFont.truetype(FONT_PATH, 20)

    def canvas(width=1400, height=760):
        image = Image.new("RGB", (width, height), "white")
        return image, ImageDraw.Draw(image)

    def centered(draw, xy, text, font, fill=INK):
        box = draw.textbbox((0, 0), text, font=font)
        draw.text((xy[0] - (box[2] - box[0]) / 2, xy[1] - (box[3] - box[1]) / 2), text, font=font, fill=f"#{fill}" if len(fill) == 6 else fill)

    image, draw = canvas()
    centered(draw, (700, 55), "尼泊尔 HS 03 进口来源结构（2022）", title_font, BLUE)
    left, base, max_height = 120, 650, 470
    values = [row[1] for row in d.IMPORT_SOURCE_2022]
    bar_width, gap = 155, 75
    colors = ["#2E74B5", "#5B9BD5", "#A5A5A5", "#FFC000", "#70AD47"]
    draw.line((90, base, 1320, base), fill="#667085", width=3)
    for index, row in enumerate(d.IMPORT_SOURCE_2022):
        x = left + index * (bar_width + gap)
        height = int(row[1] / max(values) * max_height)
        draw.rectangle((x, base - height, x + bar_width, base), fill=colors[index])
        centered(draw, (x + bar_width / 2, base - height - 28), f"{row[1]}%", label_font, BLUE)
        centered(draw, (x + bar_width / 2, base + 42), row[0], label_font)
    image.save(ASSETS / "import_sources.png")

    image, draw = canvas()
    centered(draw, (700, 45), "尼泊尔 HS 03 进口品类结构（2022）", title_font, BLUE)
    y = 110
    for row in d.IMPORT_CATEGORY_2022:
        label = row[0].split("（")[0]
        draw.text((70, y + 8), label, font=small_font, fill=f"#{INK}")
        width = int(row[1] / 45 * 760)
        draw.rectangle((420, y, 420 + width, y + 48), fill="#2E74B5")
        draw.text((430 + width, y + 8), f"{row[1]}%", font=small_font, fill=f"#{BLUE}")
        y += 82
    image.save(ASSETS / "import_categories.png")

    image, draw = canvas(height=850)
    centered(draw, (700, 45), "城市优先级：用于试点排序，不代表市场规模", title_font, BLUE)
    y = 105
    for index, row in enumerate(d.CITY_SCORES):
        draw.text((55, y + 8), row[0], font=small_font, fill=f"#{INK}")
        width = int(row[1] / 100 * 850)
        draw.rectangle((350, y, 350 + width, y + 48), fill="#1F4D78" if index < 3 else "#9DC3E6")
        draw.text((365 + width, y + 8), str(row[1]), font=small_font, fill=f"#{BLUE}")
        y += 82
    image.save(ASSETS / "city_scores.png")

    image, draw = canvas(height=480)
    centered(draw, (700, 45), "AI销售系统最小闭环", title_font, BLUE)
    boxes = [
        ("公开研究", 40, "#DDEBF7"),
        ("供应资料", 310, "#E2F0D9"),
        ("AI匹配与草拟", 580, "#FFF2CC"),
        ("人工确认", 850, "#FCE4D6"),
        ("联系／报价／下单", 1120, "#D9EAD3"),
    ]
    for label, x, color in boxes:
        draw.rounded_rectangle((x, 145, x + 220, 265), radius=14, fill=color, outline="#5B6573", width=2)
        centered(draw, (x + 110, 205), label, label_font)
    for x in (260, 530, 800, 1070):
        draw.line((x + 8, 205, x + 42, 205), fill="#1F4D78", width=5)
        draw.polygon([(x + 42, 205), (x + 27, 194), (x + 27, 216)], fill="#1F4D78")
    centered(draw, (700, 365), "成交、履约、拒绝与复购结果回流，持续修正判断", label_font, BLUE)
    image.save(ASSETS / "ai_flow.png")

    image, draw = canvas(height=540)
    centered(draw, (700, 45), "90天受控试点", title_font, BLUE)
    phases = [
        ("0-30天", "补齐真值\n验证合规\n50个优先对象", "#DDEBF7"),
        ("31-60天", "5-10次有效沟通\n3-5个样品或试单", "#E2F0D9"),
        ("61-90天", "复购验证\n继续／调整／暂停", "#FFF2CC"),
    ]
    for index, (title, body, color) in enumerate(phases):
        x = 75 + index * 440
        draw.rounded_rectangle((x, 135, x + 350, 405), radius=18, fill=color, outline="#5B6573", width=2)
        centered(draw, (x + 175, 195), title, title_font, BLUE)
        for line_index, line in enumerate(body.splitlines()):
            centered(draw, (x + 175, 270 + line_index * 45), line, label_font)
        if index < 2:
            draw.line((x + 360, 270, x + 420, 270), fill="#1F4D78", width=5)
            draw.polygon([(x + 420, 270), (x + 400, 257), (x + 400, 283)], fill="#1F4D78")
    image.save(ASSETS / "roadmap.png")


def doc_overview():
    sections = ["核心结论", "文件清单与阅读顺序", "本轮完成边界", "合作分工", "下一步"]
    doc = new_doc("项目总览与阅读说明", "21份 DOCX 的阅读入口与合作讨论顺序", compact=True)
    add_toc(doc, sections)
    doc.add_heading("核心结论", level=1)
    add_callout(doc, "总判断", "项目值得做受控试点，但不适合立即做全国铺货。第一阶段只做加德满都谷地的冷冻品、小范围采购验证和真实履约记录。", PALE_GREEN)
    add_table(
        doc,
        ["问题", "当前建议", "状态"],
        [
            ["先做哪里", "Kathmandu、Lalitpur、Bhaktapur 按一个冷链区域推进；Pokhara 等复购后再进", "推测"],
            ["先做什么", "冷冻巴沙鱼柳、规格化虾、鱿鱼、即烹组合；三文鱼只做小量形象测试", "部分成立"],
            ["先做哪类合作", "酒店与度假村、亚洲餐厅、冷冻经销、精品零售", "部分成立"],
            ["面向消费者怎么做", "内容教育与询问收集，成交放到独立页面、WhatsApp/Viber/电话", "部分成立"],
            ["TikTok 能做什么", "可做自然内容；Shop 不在尼泊尔官方市场清单；广告与 LIVE 需实际账号复核", "已确认／待验证"],
            ["收款怎么做", "Fonepay QR、eSewa、Khalti 与货到付款并行，需本地主体接入", "部分成立"],
            ["什么时候扩大", "至少 8 个试单单位，复购或补货不低于 30%，且无重大合规与履约事件", "试点规则"],
            ["什么情况下暂停", "50 个已核实对象、两轮实验后仍少于 3 个试单，或出现负贡献与重大食品安全问题", "试点规则"],
        ],
        [1.35, 4.25, 0.9],
        8.4,
    )
    doc.add_heading("文件清单与阅读顺序", level=1)
    names = [
        "00 项目总览与阅读说明",
        "01 市场深度调研报告",
        "02 全品类海鲜机会手册",
        "03 城市、人群、场景优先级评分",
        "04 竞品与价格研究手册",
        "05 B2B线上销售全渠道方案",
        "06 B2B潜在合作对象数据库",
        "07 B2B AI获客与跟进SOP",
        "08 B2C线上销售全渠道方案",
        "09 TikTok专项运营与AI内容方案",
        "10 TikTok内容选题与脚本样例库",
        "11 AI销售系统总体架构",
        "12 AI协作模块功能与数据流设计",
        "13 CRM信息项与销售漏斗设计",
        "14 供应链资料与数据接口清单",
        "15 30/60/90天执行计划",
        "16 风险、停止投入、扩大投入标准",
        "17 来源与引用索引",
        "18 后续Codex 5.5执行任务包",
        "总方案完整版",
        "研究来源使用说明",
    ]
    add_step_list(doc, names)
    doc.add_heading("本轮完成边界", level=1)
    add_table(
        doc,
        ["事项", "本轮状态", "说明"],
        [
            ["公开市场研究", "已完成", "覆盖人口、贸易、城市、旅游、支付、平台、竞品、合规和协会入口"],
            ["商品方向和试点顺序", "已形成候选结论", "以公开供给、价格和履约可行性评分，仍需真实供应资料验证"],
            ["合作对象清单", "已建立公开名单", "25 个公开组织入口，尚未联系，不代表采购意向"],
            ["真实商品与落地成本", "待验证", "缺少供应价格、库存、海关编码、税费、冷链和本地进口方确认"],
            ["真实投放与成交", "未执行", "没有使用任何真实账号、广告预算、支付账户或外发权限"],
            ["市场规模与销售结果", "不能下结论", "HS 03 数据只作贸易口径参考，不等于全部市场或可获得收入"],
        ],
        [1.55, 1.25, 3.7],
        8.6,
    )
    doc.add_heading("合作分工", level=1)
    add_table(
        doc,
        ["责任方", "主要责任", "首要交付"],
        [
            ["中国供应方", "提供真实商品、价格、库存、资质、产地、包装、保质期和温控要求", "可核实的商品资料包"],
            ["尼泊尔进口与供应链方", "许可、海关归类、税费、仓储、冷链、配送和召回", "清关与冷链书面确认"],
            ["线上销售团队", "内容、公开名单、沟通草稿、订单协同、复盘", "小范围验证与回流记录"],
            ["项目负责人", "批准报价、样品、预算、投放、特殊条款和扩大或暂停", "每周决策记录"],
        ],
        [1.3, 3.25, 1.95],
        8.6,
    )
    add_next_steps(doc, ["先读总方案完整版和第14份供应链清单", "补齐 P0 信息并请报关行逐品复核", "在加德满都谷地选 10-20 个真实商品候选", "人工审核后联系首批 15-25 个公开组织入口", "30 天后只按有效沟通、试单、履约和复购判断"])
    save(doc, "00_项目总览与阅读说明.docx")


def doc_market():
    sections = ["执行摘要", "宏观与水产供给", "进口结构", "城市与需求场景", "市场规模边界", "进入建议", "来源"]
    doc = new_doc("尼泊尔海鲜市场深度调研报告", "贸易、城市、消费场景、支付与进入策略", compact=False)
    add_toc(doc, sections)
    add_basis(doc)
    doc.add_heading("执行摘要", level=1)
    add_callout(doc, "结论", "尼泊尔存在进口海鲜与城市冷冻零售，但公开证据支持的是“加德满都谷地的小规模验证”，不是“全国大市场已经成熟”。", PALE_GREEN)
    add_table(
        doc,
        ["判断", "证据", "意义"],
        [
            ["进口市场存在但波动明显", "HS 03 进口额由 2021 年约 1,520 万美元降至 2022 年约 636 万美元", "用最新清关与销售数据校准，不能用单年外推"],
            ["本地淡水鱼供给占主导", "研究显示 2023 年鱼类供给约 11.37 万吨、养殖占比较高", "进口产品要做品类差异、规格稳定和方便处理"],
            ["鱼与鱼柳是进口主结构", "2022 年鲜冷鱼约 43%，鱼柳及鱼肉约 39%", "巴沙鱼柳等标准化商品更适合先跑"],
            ["支付基础可用", "二维码和电子钱包月度笔数都在 3,700 万笔以上", "独立页面可接本地支付，但必须由本地主体办理"],
            ["酒店餐饮是可触达场景", "旅游统计列出 1,578 家登记住宿单位、62,642 张床位", "先找厨师测试与稳定补货，不先做全国品牌广告"],
        ],
        [1.55, 2.55, 2.4],
        8.2,
    )
    doc.add_heading("宏观与水产供给", level=1)
    add_table(doc, ["指标", "数值", "时间", "来源", "状态", "对项目的意义"], [row[:6] for row in d.MARKET_FACTS], [1.1, 1.05, 0.8, 1.2, 1.05, 2.3], 7.6)
    doc.add_heading("进口结构", level=1)
    add_picture(doc, ASSETS / "import_sources.png", "图 1｜2022 年 HS 03 进口来源；中国当年金额极低，未单独显示。数据源：S09。")
    add_picture(doc, ASSETS / "import_categories.png", "图 2｜2022 年 HS 03 进口品类；只覆盖第 03 章，不覆盖所有加工海鲜。数据源：S09。")
    add_table(doc, ["来源国", "份额", "金额（百万美元）"], [[x[0], f"{x[1]}%", x[2]] for x in d.IMPORT_SOURCE_2022], [2.5, 1.5, 2.5], 9)
    add_table(doc, ["进口品类", "份额", "金额（百万美元）"], [[x[0], f"{x[1]}%", x[2]] for x in d.IMPORT_CATEGORY_2022], [3, 1.5, 2], 9)
    doc.add_heading("城市与需求场景", level=1)
    add_picture(doc, ASSETS / "city_scores.png", "图 3｜城市评分是工作排序，不是市场份额或销售预测。")
    add_table(doc, ["城市", "评分", "优先级", "建议", "理由", "可信度"], [x[:6] for x in d.CITY_SCORES], [1, 0.55, 0.85, 1.05, 2.55, 0.5], 7.7)
    add_table(doc, ["场景", "评分", "适合品类", "方向", "首要验证"], d.SCENARIO_SCORES, [1.5, 0.55, 2.1, 0.75, 1.6], 8.2)
    doc.add_heading("市场规模边界", level=1)
    add_callout(doc, "不能直接给出精确规模", "目前只有 HS 03 的历史进口口径、公开零售价和城市代理指标。它们不覆盖加工海鲜、非正规贸易、本地销售、损耗和可服务区域，因此不能直接写成 TAM、SAM 或 SOM。", PALE_RED)
    add_table(
        doc,
        ["层级", "当前可用口径", "还缺什么", "本轮处理"],
        [
            ["全国需求参考", "人口、收入、鱼类供给和 HS 03 进口", "最新贸易、加工海鲜、本地销售", "只作背景"],
            ["可服务市场", "加德满都谷地城市、酒店、餐饮、线上零售入口", "真实冷链半径、可售商品、批发价格", "待 30 天验证"],
            ["可获得份额", "公开组织名单与内容渠道", "回复、试单、复购、损耗、毛利", "90 天后计算"],
        ],
        [1.2, 2.2, 2.1, 1],
        8.4,
    )
    doc.add_heading("进入建议", level=1)
    add_step_list(
        doc,
        [
            "把 Kathmandu、Lalitpur、Bhaktapur 当作一个冷链区域，只做可追溯冷冻品。",
            "先试巴沙鱼柳、规格化虾、鱿鱼和即烹组合；三文鱼只做小量高价测试。",
            "先找酒店、亚洲餐厅、冷冻经销和精品零售，做样品、厨师测试、出成率与补货验证。",
            "消费者侧只做烹饪教育、询问收集和小范围配送，不做全国投放。",
            "补齐许可、海关编码、税费、标签、冷链和支付后再谈正式报价与扩大。",
        ],
    )
    add_open_items(doc)
    add_sources(doc, "S01,S02,S03,S04,S05,S06,S07,S08,S09,S10,S20,S21,S28,S29,S30,S31,S32,S33,S41,S42,S43,S44,S45,S46")
    save(doc, "01_尼泊尔海鲜市场深度调研报告.docx")


def doc_products():
    sections = ["评分方法", "首批候选", "全品类评分", "包装与场景", "暂缓方向", "验证计划", "来源"]
    doc = new_doc("尼泊尔全品类海鲜机会手册", "鱼、虾、蟹、贝、鱿鱼与即烹组合的试点顺序", compact=True)
    add_toc(doc, sections)
    doc.add_heading("评分方法", level=1)
    add_callout(doc, "口径", "评分只用于首批测试排序。公开价格和在售证据能说明“当地有人卖”，不能证明供应方的商品一定有利润或一定能成交。")
    add_table(doc, ["评分因素", "权重"], [[k, f"{v}%"] for k, v in d.PRODUCT_WEIGHTS.items()], [3.5, 3], 9)
    doc.add_heading("首批候选", level=1)
    first = [p for p in d.PRODUCTS if p["conclusion"] == "首批测试"]
    add_table(doc, ["商品方向", "形态", "包装", "B2B", "B2C", "总分", "依据"], [[p["name"], p["form"], p["pack"], p["b2b"], p["b2c"], p["total"], p["evidence"]] for p in first], [1.4, 0.9, 1.35, 0.45, 0.45, 0.5, 1.85], 7.6)
    doc.add_heading("全品类评分", level=1)
    add_table(
        doc,
        ["商品方向", "类别", "形态", "B2B", "B2C", "总分", "结论", "公开价格", "来源"],
        [[p["name"], p["category"], p["form"], p["b2b"], p["b2c"], p["total"], p["conclusion"], p["price"], p["source"]] for p in d.PRODUCTS],
        [1.25, 0.65, 0.9, 0.4, 0.4, 0.45, 0.9, 1.05, 0.5],
        6.9,
    )
    doc.add_heading("包装与场景", level=1)
    add_table(
        doc,
        ["场景", "建议形态", "建议包装", "成交重点", "不能先承诺"],
        [
            ["酒店与餐厅", "规格统一的鱼柳、虾、鱿鱼", "1.8 千克、5-10 千克或按真实供应规格", "出成率、稳定到货、替代规则", "长期价、账期与每日库存"],
            ["精品零售", "小包装冷冻品、即烹组合", "500 克至 1 千克", "标签清楚、烹饪简单、复购", "全国配送与无条件退换"],
            ["家庭聚餐", "去骨鱼柳、虾、鱿鱼圈、组合包", "500 克至 1 千克", "分量、做法、解冻指引", "营养或健康功效"],
            ["礼赠与高端体验", "三文鱼、软壳蟹等小量测试", "小包装", "真实产地、储运与稀缺性", "固定高端需求"],
        ],
        [1.15, 1.6, 1.4, 1.55, 1.3],
        8,
    )
    doc.add_heading("暂缓方向", level=1)
    delayed = [p for p in d.PRODUCTS if "不建议" in p["conclusion"] or "暂缓" in p["conclusion"]]
    add_table(doc, ["品类", "当前结论", "原因"], [[p["name"], p["conclusion"], p["evidence"]] for p in delayed], [1.7, 1.2, 3.6], 8.4)
    doc.add_heading("验证计划", level=1)
    add_step_list(doc, ["从供应方获得 10-20 个真实商品资料", "逐个确认海关编码、许可、标签、温控和落地成本", "选 5 个候选做厨师测试、家庭试吃与配送演练", "记录出成率、拒绝原因、退款、损耗和补货", "只保留能稳定履约且有复购信号的商品"])
    add_sources(doc, "S06,S07,S08,S09,S11,S12,S13,S29,S30,S31,S32,S33")
    save(doc, "02_尼泊尔全品类海鲜机会手册.docx")


def doc_priorities():
    sections = ["城市排序", "B2B群体", "B2C群体", "消费场景", "进入顺序", "验证方式", "来源"]
    doc = new_doc("城市、人群、场景优先级评分", "先在可控区域验证，再决定是否扩城", compact=True)
    add_toc(doc, sections)
    doc.add_heading("城市排序", level=1)
    add_picture(doc, ASSETS / "city_scores.png", "评分用于排工作顺序；不等于销售额预测。")
    add_table(doc, ["城市", "评分", "优先级", "动作", "理由", "可信度", "来源"], d.CITY_SCORES, [0.85, 0.45, 0.7, 0.85, 2.25, 0.55, 0.85], 7.2)
    doc.add_heading("B2B群体", level=1)
    add_table(doc, ["群体", "评分", "优先级", "需求逻辑", "首要验证", "来源"], d.B2B_SEGMENTS, [1.25, 0.45, 0.75, 2.2, 1.35, 0.5], 7.8)
    doc.add_heading("B2C群体", level=1)
    add_table(doc, ["群体", "评分", "优先级", "需求逻辑", "首要验证"], d.B2C_SEGMENTS, [1.55, 0.45, 0.75, 2.45, 1.3], 7.9)
    doc.add_heading("消费场景", level=1)
    add_table(doc, ["场景", "评分", "适合品类", "方向", "首要验证"], d.SCENARIO_SCORES, [1.55, 0.5, 2.05, 0.7, 1.7], 8.2)
    doc.add_heading("进入顺序", level=1)
    add_step_list(doc, ["加德满都谷地：餐厅、酒店、冷冻经销、精品零售", "加德满都谷地：家庭聚餐和年轻专业人士的小范围配送", "Pokhara：只有在谷地出现稳定补货后进入", "Bharatpur、Butwal、Biratnagar、Birgunj：先做名单和配送条件观察", "其他城市：不做先期库存和付费投放"])
    doc.add_heading("验证方式", level=1)
    add_table(
        doc,
        ["对象", "30天要看什么", "60天要看什么", "90天怎么判断"],
        [
            ["城市", "有效回复、可送地址、配送时效", "试单与拒绝原因", "补货、损耗和单位贡献"],
            ["B2B群体", "厨师或采购讨论", "样品、出成率、账期要求", "稳定补货与回款"],
            ["消费者群体", "询问、加购、价格反馈", "首单、配送、售后", "复购和推荐"],
            ["场景", "内容互动和真实需求", "同一场景的多次成交", "能否复制而不增加异常"],
        ],
        [1.1, 1.8, 1.8, 1.8],
        8.2,
    )
    add_sources(doc, "S01,S02,S03,S04,S05,S28,S34,S35,S36,S37,S38,S39,S40")
    save(doc, "03_城市_人群_场景优先级评分.docx")


def doc_prices():
    sections = ["价格结论", "公开样本", "价格带与定位", "竞争判断", "持续监测", "来源"]
    doc = new_doc("竞品与价格研究手册", "公开零售样本、价格带与首批定位", compact=True)
    add_toc(doc, sections)
    doc.add_heading("价格结论", level=1)
    add_callout(doc, "总判断", "公开价格显示巴沙鱼柳适合走量，虾覆盖中低到中高价带，鳟鱼、鱿鱼、组合包可做差异，三文鱼与龙虾只适合小量高价测试。", PALE_GREEN)
    add_table(
        doc,
        ["品类", "公开价格带", "当前定位", "说明"],
        [
            ["巴沙／Pangasius", "约 540-900 NPR／千克", "走量与入门", "公开样本较多，适合标准化测试"],
            ["虾", "约 1,000-2,200 NPR／千克；部分规格更高", "主力毛利", "规格、处理方式和包装影响很大"],
            ["鳟鱼", "约 1,700-1,750 NPR／千克", "中高端差异", "有本地供给，不能只按进口鱼比较"],
            ["鱿鱼", "约 3,190 NPR／千克的公开样本", "细分与餐饮", "样本少，先测厨师需求"],
            ["三文鱼", "约 8,900 NPR／千克", "高价形象", "只做小量，不能假设广泛接受"],
        ],
        [1.2, 1.8, 1.35, 2.15],
        8.3,
    )
    doc.add_heading("公开样本", level=1)
    add_table(doc, ["商品", "平台／企业", "类别", "规格", "售价NPR", "折算NPR/kg", "页面状态", "来源", "可信度"], d.PRICE_RECORDS, [1.25, 0.95, 0.75, 0.7, 0.6, 0.65, 1.15, 0.45, 0.45], 6.7)
    doc.add_heading("价格带与定位", level=1)
    add_table(
        doc,
        ["定位", "适合商品", "比较方式", "试点动作"],
        [
            ["入门", "巴沙鱼柳、鱼排", "每千克到手价、可食用比例、烹饪难度", "做 500 克和 1 千克包装测试"],
            ["主力", "规格化虾、鱿鱼、组合包", "规格、净重、冰衣、处理方式、出成率", "做家庭和餐饮两套包装"],
            ["中高端", "鳟鱼、软壳蟹", "产地、稳定性、口感、供应频率", "先做餐厅或精品零售"],
            ["高价形象", "三文鱼、龙虾", "真实产地、切割、储运与损耗", "限量预订，不先压库存"],
        ],
        [1, 1.65, 2.5, 1.35],
        8.3,
    )
    doc.add_heading("竞争判断", level=1)
    add_table(
        doc,
        ["竞争来源", "优势", "中国供应链需要回答的问题"],
        [
            ["印度", "距离、传统贸易和快速补货", "能否在稳定规格、加工深度或价格上形成明显优势"],
            ["越南", "鱼柳等品类已有进口基础", "中国商品是否只是同质替代，还是有更好组合"],
            ["本地淡水鱼", "新鲜、熟悉、价格和供应关系", "进口海鲜提供什么不可替代的场景"],
            ["现有冷冻零售商", "已有商品、配送和本地信任", "更适合合作、供货还是直接竞争"],
        ],
        [1.2, 2.25, 3.05],
        8.3,
    )
    doc.add_heading("持续监测", level=1)
    add_step_list(doc, ["每周记录同一商品的规格、价格、库存和配送范围", "价格变化与缺货分开记录", "比较净重、冰衣、处理方式和可食用比例", "供应方落地成本完成前不承诺低价", "每月按成交与复购重新定义主力价格带"])
    add_sources(doc, "S09,S29,S30,S31,S32,S33,S34,S35")
    save(doc, "04_竞品与价格研究手册.docx")


def doc_b2b_channels():
    sections = ["战略重点", "采购群体", "渠道组合", "成交路径", "人工确认点", "30天指标", "来源"]
    doc = new_doc("B2B线上销售全渠道方案", "从公开名单到样品、试单和补货的可控路径", compact=True)
    add_toc(doc, sections)
    doc.add_heading("战略重点", level=1)
    add_callout(doc, "主线", "首阶段把线上渠道当作“发现与沟通工具”，把样品、厨师测试、正式报价、合同与回款交给人工确认。目标是补货，不是名单数量。", PALE_GREEN)
    add_table(
        doc,
        ["优先层", "对象", "核心商品", "首要动作"],
        [
            ["P1", "高端酒店、度假村、亚洲餐厅与火锅门店", "规格化虾、鱿鱼、鱼柳、少量三文鱼", "厨师测试、出成率和稳定交货"],
            ["P1", "冷冻经销、批发与精品零售", "鱼柳、虾、鱿鱼、即烹组合", "覆盖区域、冷库、上架与退货条件"],
            ["P2", "宴会、旅游餐饮、云厨房", "大包装虾、鱼柳、鱿鱼", "档期、单次用量和回款"],
            ["暂缓", "团餐与宽泛企业采购", "待价格与招采验证", "不先投入大量时间"],
        ],
        [0.65, 2.2, 2.05, 1.6],
        8.3,
    )
    doc.add_heading("采购群体", level=1)
    add_table(doc, ["采购群体", "评分", "优先级", "需求逻辑", "首要验证", "来源"], d.B2B_SEGMENTS, [1.25, 0.45, 0.75, 2.15, 1.4, 0.5], 7.7)
    doc.add_heading("渠道组合", level=1)
    add_table(
        doc,
        ["渠道", "可用性", "适合对象", "可用信息", "推进方式", "AI作用", "风险", "指标", "来源"],
        d.B2B_CHANNELS,
        [0.8, 0.8, 1.2, 0.9, 1.3, 0.55, 1.05, 0.8, 0.5],
        6.25,
    )
    doc.add_heading("成交路径", level=1)
    add_step_list(
        doc,
        [
            "从协会、官网、地图和公开目录发现组织入口，保留原始链接。",
            "按城市、业态、菜单适配、冷链能力和公开信息完整度排序。",
            "AI只生成个性化联系草稿，人工核实后从公开商务入口发送。",
            "有效回复后先确认品类、规格、用量、现有供应、痛点和决策流程。",
            "供应链真值通过后才给正式报价；样品、账期、独家与合同都要批准。",
            "每次样品、试单、配送、异常、回款和补货都写回 CRM。",
            "30、60、90 天按试单和补货决定继续、调整或暂停。",
        ],
    )
    doc.add_heading("人工确认点", level=1)
    add_table(
        doc,
        ["事项", "AI可以做", "必须由人工做"],
        [
            ["名单", "整理公开组织、去重、评分和来源记录", "判断是否适合联系与是否存在关系冲突"],
            ["首轮联系", "起草个性化介绍和问题", "最终审核并发送"],
            ["商品建议", "按真实供应资料匹配候选", "决定最终商品、规格和替代方案"],
            ["报价", "读取已批准价格形成草稿", "确认币种、税费、有效期、库存、账期与签发"],
            ["样品与试单", "安排提醒和记录", "批准成本、发货、签收和异常处理"],
            ["停止联系", "识别明确拒绝并标记", "处理争议和特殊关系"],
        ],
        [1.15, 2.6, 2.75],
        8.3,
    )
    doc.add_heading("30天指标", level=1)
    add_table(
        doc,
        ["指标", "最低目标", "解释"],
        [
            ["高优先公开组织", "50 个", "有城市、类别和至少一个可复核来源；高优先尽量两处来源"],
            ["人工发送", "15-25 条", "个性化、使用公开商务入口，不群发"],
            ["有效沟通", "5-10 次", "讨论了需求、规格、采购或试样，不把已读当有效"],
            ["样品或试单", "3 个方向以上", "必须有真实供应与履约记录"],
            ["误承诺", "0", "不猜价格、库存、许可和送达"],
        ],
        [1.45, 1.05, 4],
        8.5,
    )
    add_sources(doc, "S28,S29,S30,S33,S34,S35,S36,S37,S38,S39,S40")
    save(doc, "05_B2B线上销售全渠道方案.docx")


def doc_b2b_leads():
    sections = ["使用边界", "公开组织名单", "优先级复核", "接触准备", "状态说明", "来源"]
    doc = new_doc("B2B潜在合作对象数据库", "公开可复核的组织入口，不代表采购意向", compact=True)
    add_toc(doc, sections)
    doc.add_heading("使用边界", level=1)
    add_callout(doc, "重要说明", "名单来自公开网页和协会入口，尚未联系，也没有证明对方愿意采购。只使用公开商务联系方式，不补猜个人身份、私人号码或采购权限。", PALE_GOLD)
    doc.add_heading("公开组织名单", level=1)
    add_table(
        doc,
        ["组织", "城市", "类别", "可能关系", "公开入口", "公开联系", "可信度", "建议动作", "来源"],
        d.LEADS,
        [1.25, 0.75, 1.05, 0.85, 0.8, 1.1, 0.5, 1.35, 0.45],
        6.4,
    )
    doc.add_heading("优先级复核", level=1)
    add_table(
        doc,
        ["检查", "通过标准", "不通过时"],
        [
            ["主体存在", "官网、协会、地图或公开目录能复核", "从高优先移出"],
            ["城市适配", "位于谷地或可由已确认冷链服务", "转观察"],
            ["品类适配", "菜单、经营范围或供应能力与首批商品有关", "不做强推"],
            ["公开入口", "有官网表单、企业邮箱、公开电话或公开社媒商务入口", "不寻找私人联系"],
            ["关系冲突", "与现有合作无冲突或已获内部确认", "人工判断"],
            ["最新状态", "联系前 48 小时内复核营业与公开资料", "暂缓联系"],
        ],
        [1.2, 3.1, 2.2],
        8.4,
    )
    doc.add_heading("接触准备", level=1)
    add_step_list(doc, ["查看官网、菜单或业务范围", "选择一个真实匹配商品和一个具体问题", "确认价格、库存和配送资料是否足够", "生成短而具体的介绍草稿", "人工审核后从公开入口联系", "记录回复、拒绝、下一步和日期"])
    doc.add_heading("状态说明", level=1)
    add_table(
        doc,
        ["状态", "含义", "不能代表什么"],
        [
            ["候选", "公开页面显示可能适配", "不代表已核实采购需要"],
            ["已核实", "主体、入口和适配性已复核", "不代表已联系"],
            ["已联系", "人工完成首次联系并留有记录", "不代表已沟通"],
            ["已沟通", "对方真实讨论了需求或条件", "不代表已试单"],
            ["已试单", "有真实订单与履约", "不代表稳定合作"],
            ["已补货", "出现真实再次采购", "仍需看毛利、回款和持续性"],
        ],
        [1.05, 2.9, 2.55],
        8.5,
    )
    add_sources(doc, "S28,S29,S30,S33,S34,S35,S36,S37,S38,S39,S40")
    save(doc, "06_B2B潜在合作对象数据库.docx")


def doc_b2b_sop():
    sections = ["工作原则", "每日流程", "联系节奏", "话术样例", "CRM记录", "停止规则", "来源"]
    doc = new_doc("B2B AI获客与跟进SOP", "AI负责整理和草拟，人工负责现实动作与承诺", compact=True)
    add_toc(doc, sections)
    doc.add_heading("工作原则", level=1)
    add_table(
        doc,
        ["原则", "执行要求"],
        [
            ["公开来源", "只使用官网、协会、地图、公开目录和授权后台"],
            ["真实供应", "正式联系前确认候选商品；正式报价前确认真实价格、库存、税费和配送"],
            ["个性化", "每条联系必须说明为什么适合该组织，不使用无差别群发"],
            ["人工发送", "AI只生成草稿；首次联系、报价、样品、电话、见面和合同都由人工确认"],
            ["尊重拒绝", "明确拒绝或要求停止后立即停止，不自动重启"],
            ["结果分层", "已读、回复、有效沟通、样品、试单和补货分开记录"],
        ],
        [1.25, 5.25],
        8.7,
    )
    doc.add_heading("每日流程", level=1)
    add_step_list(doc, ["复核供应链真值和阻断事项", "从高优先队列选 3-5 个组织", "检查公开页面是否仍有效", "AI按组织与商品生成联系草稿", "人工审核、修改并发送", "当天记录回复和下一步", "异常、拒绝或价格问题进入人工处理"])
    doc.add_heading("联系节奏", level=1)
    add_table(
        doc,
        ["时间", "动作", "目的", "停止条件"],
        [
            ["D0", "短介绍＋一个具体问题", "确认是否负责相关采购或是否愿意了解样品", "明确拒绝"],
            ["D3-D4", "补充真实商品图、规格或一页说明", "降低理解成本", "资料不足或已拒绝"],
            ["D7-D10", "围绕需求、出成率或配送做一次跟进", "推进样品或短会", "无回复且无新价值"],
            ["D14", "礼貌结束本轮联系", "保留未来公开联系可能", "结束后 90 天内不自动重启"],
            ["成交后", "按实际消耗周期提醒补货", "验证复购", "对方要求停止或履约异常"],
        ],
        [0.85, 2.1, 2.1, 1.45],
        8.3,
    )
    doc.add_heading("话术样例", level=1)
    add_callout(doc, "首次联系", "您好，我们正在评估向加德满都谷地稳定供应冷冻虾、鱼柳和鱿鱼的合作方式。看到贵方经营酒店／餐饮／食品零售，想先了解目前更看重规格稳定、到货时效，还是价格与最小订货量？如果方向合适，我们可以在资料确认后提供一页商品说明供参考。")
    add_callout(doc, "样品邀请", "根据您提到的使用场景，我们初步匹配了一个候选规格。正式安排前，我们还会确认净重、温控、到货时间和费用。若您愿意测试，请告知预计用量、收货区域和方便反馈的时间。")
    add_callout(doc, "结束本轮联系", "谢谢您阅读。本轮先不继续打扰。如果以后有冷冻鱼柳、虾或鱿鱼的采购需求，可以通过这个公开入口联系我们；届时会以最新库存和配送条件为准。")
    doc.add_heading("CRM记录", level=1)
    add_table(doc, ["每次联系后记录", "要求"], [["来源与原始链接", "能回到公开页面"], ["联系时间与渠道", "区分草稿、已发送与已回复"], ["需求与拒绝原因", "保留原意，不强行归类"], ["使用的价格版本", "避免过期报价"], ["下一步与日期", "活跃记录必须有"], ["停止联系标记", "立即阻断后续自动任务"]], [2.1, 4.4], 8.5)
    doc.add_heading("停止规则", level=1)
    add_table(doc, ["触发", "处理"], [["明确拒绝或要求停止", "立即结束并写入停止标记"], ["公开主体或入口失效", "停止联系，重新核实"], ["供应、价格、库存或配送缺失", "不报价，只说明待确认"], ["投诉、食品安全或冷链异常", "转人工，停止同批次销售"], ["连续两轮无有效沟通", "调整对象或商品，不增加频率"]], [2.5, 4], 8.5)
    add_sources(doc, "S18,S36,S37,S38,S39,S40")
    save(doc, "07_B2B_AI获客与跟进SOP.docx")


def doc_b2c_channels():
    sections = ["首阶段人群", "渠道组合", "成交与支付", "履约与售后", "内容到复购", "指标", "来源"]
    doc = new_doc("B2C线上销售全渠道方案", "加德满都谷地的小范围内容、下单、支付与配送", compact=True)
    add_toc(doc, sections)
    doc.add_heading("首阶段人群", level=1)
    add_table(doc, ["人群", "评分", "优先级", "需求逻辑", "首要验证"], d.B2C_SEGMENTS, [1.5, 0.45, 0.75, 2.5, 1.3], 7.9)
    doc.add_heading("渠道组合", level=1)
    add_table(
        doc,
        ["渠道", "可用性", "主要作用", "成交路径", "AI作用", "人工确认", "指标", "来源"],
        d.B2C_CHANNELS,
        [0.85, 1.1, 1.05, 1.25, 0.95, 1.05, 0.8, 0.45],
        6.8,
    )
    doc.add_heading("成交与支付", level=1)
    add_callout(doc, "建议路径", "TikTok／Meta／Google／Foodmandu 等负责发现，自有页面负责展示真实商品，WhatsApp／Viber／电话负责补充沟通，Fonepay QR、eSewa、Khalti 和货到付款负责结账。", PALE_GREEN)
    add_table(
        doc,
        ["环节", "首选", "备选", "上线条件"],
        [
            ["发现", "TikTok自然内容、Meta、Google Maps", "美食达人与本地社群", "真实商品和素材"],
            ["了解", "独立商品页或本地电商页面", "WhatsApp／Viber目录", "价格、库存、配送和标签已核"],
            ["沟通", "WhatsApp／Viber／电话", "网页表单", "隐私告知和人工接管"],
            ["收款", "Fonepay QR、eSewa、Khalti", "货到付款、卡", "本地主体、对账和退款流程"],
            ["配送", "谷地冷链或已验证同城配送", "自提", "温控、半径、时效和异常处理"],
        ],
        [1, 2.25, 1.55, 1.7],
        8.3,
    )
    doc.add_heading("履约与售后", level=1)
    add_step_list(doc, ["下单前显示可送区域、费用、截止时间和预计到达", "付款前再次确认商品、净重、价格和库存", "拣货时记录批次和包装状态", "配送过程保留温控与交接记录", "签收后提供解冻、储存和烹饪指引", "温控、破损、少件或延误立即转人工", "完成后按真实消耗周期提醒复购"])
    doc.add_heading("内容到复购", level=1)
    add_table(
        doc,
        ["阶段", "内容", "行动", "衡量"],
        [
            ["认识", "海鲜怎么选、怎么看净重、怎么解冻", "收藏或进入商品页", "完播、收藏、页面访问"],
            ["考虑", "15分钟做法、分量、真实包装与配送", "询问或加购", "有效询问、加购"],
            ["下单", "真实库存、价格、配送和售后", "付款或货到付款", "转化、支付成功"],
            ["体验", "收货检查、储存和烹饪", "完成订单", "准时率、异常率"],
            ["复购", "按场景提醒、同类替代和组合", "再次购买", "30／60／90天复购"],
        ],
        [0.85, 2.1, 1.7, 1.85],
        8.3,
    )
    doc.add_heading("指标", level=1)
    add_table(doc, ["类别", "首阶段指标"], [["内容", "完播、收藏、主页或商品页访问、有效询问"], ["商品页", "可售页面访问、加购、结账开始、退出原因"], ["支付", "成功率、货到付款占比、退款时长"], ["配送", "准时率、温控异常、破损、退回"], ["商业", "首单、毛利、复购、售后成本"]], [1.2, 5.3], 8.6)
    add_sources(doc, "S17,S18,S19,S20,S21,S27,S28,S34,S35")
    save(doc, "08_B2C线上销售全渠道方案.docx")


def doc_tiktok():
    sections = ["能力结论", "账号定位", "内容矩阵", "站外成交", "广告与直播闸门", "30天安排", "来源"]
    doc = new_doc("TikTok专项运营与AI内容方案", "把 TikTok 作为发现与教育入口，不假设尼泊尔 Shop 可用", compact=True)
    add_toc(doc, sections)
    doc.add_heading("能力结论", level=1)
    add_callout(doc, "核心判断", "尼泊尔可以访问 TikTok，普通内容可作为首阶段入口；TikTok Shop 不在当前官方市场清单，自助广告开户也未在官方名单中确认。站内支付、商品挂载、直播带货和达人联盟不能写成已开通。", PALE_GOLD)
    add_table(doc, ["能力", "状态", "证据或限制", "标记", "项目建议", "来源"], d.TIKTOK_CAPABILITIES, [1.15, 1.1, 2.2, 0.85, 1.6, 0.5], 7.6)
    doc.add_heading("账号定位", level=1)
    add_table(
        doc,
        ["账号", "作用", "内容", "限制"],
        [
            ["主账号", "项目与商品的正式说明", "真实商品、储运、做法、问答、合作故事", "没有真值不展示价格、库存、检测或购买反馈"],
            ["辅助内容账号", "测试本地语言、开头、节奏和生活场景", "做法、选购、火锅烧烤、家庭快手菜", "不冒充消费者，不编造体验"],
            ["员工或主厨合作", "提供专业解释与本地信任", "出成率、菜品、储存和烹饪", "需授权，不伪造身份或评价"],
        ],
        [1.25, 1.65, 2.25, 1.35],
        8.4,
    )
    doc.add_heading("内容矩阵", level=1)
    add_table(
        doc,
        ["层级", "目标", "示例", "下一步"],
        [
            ["认识", "让人看懂冷冻海鲜", "净重、规格、解冻、鱼柳与整鱼差异", "收藏或进入主页"],
            ["信任", "证明商品和履约真实", "真实包装、批次、仓储、配送与厨师测试", "查看商品页或询问"],
            ["考虑", "降低做法与价格理解难度", "15分钟做法、家庭分量、餐厅出成率", "进入站外页面"],
            ["复购", "帮助形成固定场景", "周末火锅、烧烤、快手菜与补货提醒", "再次购买"],
        ],
        [0.8, 1.5, 2.8, 1.4],
        8.4,
    )
    doc.add_heading("站外成交", level=1)
    add_step_list(doc, ["视频或 LIVE 提供清楚的站外入口", "落地页只显示真实可售商品、区域、库存、价格和时效", "通过 WhatsApp／Viber／电话补充问题", "通过本地二维码、电子钱包或货到付款结账", "订单进入供应链系统并锁定库存", "配送与售后结果回流到 CRM 和内容复盘"])
    doc.add_heading("广告与直播闸门", level=1)
    add_table(
        doc,
        ["能力", "上线前必须看到", "未通过时"],
        [
            ["LIVE", "实际账号显示 LIVE 权限，年龄、粉丝和地区要求满足", "只发短视频"],
            ["TikTok Ads", "实际商务后台可开户、可选择 Nepal 地区、可付款和可追踪", "不使用跨区绕行"],
            ["Shop／商品挂载", "官方市场清单出现 Nepal 且实际账号开通", "全部走站外"],
            ["付费转化", "Pixel 或 Events API、真实库存、配送、支付和退款完成", "只做自然内容"],
            ["达人合作", "身份、授权、报酬、内容与披露规则确认", "只做自有内容"],
        ],
        [1.1, 3.6, 1.8],
        8.2,
    )
    doc.add_heading("30天安排", level=1)
    add_table(doc, ["周", "重点", "产出", "判断"], [["第1周", "账号、语言、真实素材与站外路径检查", "账号说明、链接页、素材清单", "不能成交也能安全收集询问"], ["第2周", "三个内容角度各做两个版本", "6个短视频草稿", "比较完播、收藏和有效询问"], ["第3周", "保留最好角度，补充真实商品与做法", "4-6个改进版本", "是否出现商品页访问和真实询问"], ["第4周", "复盘内容到站外路径", "继续／调整／暂停结论", "不以单条播放量决定"]], [0.8, 2.15, 2.05, 1.5], 8.3)
    add_sources(doc, "S22,S23,S24,S25,S26,S27")
    save(doc, "09_TikTok专项运营与AI内容方案.docx")


def doc_tiktok_content():
    sections = ["内容原则", "选题库", "完整脚本样例", "审校清单", "复盘方式", "来源"]
    doc = new_doc("TikTok内容选题与脚本样例库", "12个选题与6个可改写的短视频脚本", compact=True)
    add_toc(doc, sections)
    doc.add_heading("内容原则", level=1)
    add_table(doc, ["原则", "要求"], [["真实", "只展示真实商品、包装、库存、配送和合作经历"], ["通俗", "先讲场景和做法，再讲规格和来源"], ["单一目标", "每条视频只推动一个动作：收藏、查看做法、进入商品页或询问"], ["站外成交", "不承诺 Shop；链接到可控页面或公开沟通入口"], ["双重审校", "尼泊尔语由母语者与食品安全人员审核"], ["不做功效", "不做未经证明的健康、治疗或绝对品质承诺"]], [1.05, 5.45], 8.6)
    doc.add_heading("选题库", level=1)
    add_table(doc, ["编号", "选题", "对象", "开头", "素材边界", "下一步", "指标"], d.CONTENT_IDEAS, [0.4, 1.05, 1.05, 1.5, 1.5, 1.4, 0.7], 6.9)
    doc.add_heading("完整脚本样例", level=1)
    scripts = [
        ("15分钟蒜香虾", "开头：冷冻虾不等于复杂晚餐，先看从解冻到出锅要多久。画面：真实包装、称重、冷水解冻、沥干、蒜香快炒。旁白：本条只演示通用做法，实际商品规格、价格和可送区域以页面当日信息为准。结尾：需要完整做法，可从主页进入食谱页。"),
        ("巴沙鱼柳三种做法", "开头：同一块去骨鱼柳，能不能同时适合孩子、上班族和餐厅？画面：清蒸、香煎、咖喱三段对比。旁白：重点比较处理时间、份量和出成，不做营养功效承诺。结尾：收藏这条，等真实商品上线后再查看规格。"),
        ("火锅组合怎么配", "开头：两到四个人的海鲜火锅，买多了浪费，买少了不够。画面：虾、鱼片、鱿鱼圈按人数摆盘。旁白：组合比例需要通过真实试吃确认，过敏提示和净重必须在包装上清楚显示。结尾：评论人数，领取通用配比参考。"),
        ("如何看虾的规格", "开头：21/25、31/40 到底是什么意思？画面：用相同盘子展示不同规格。旁白：数字通常与每磅数量有关，但购买时还要看处理方式、净重和冰衣。结尾：进入说明页查看规格比较。"),
        ("冷冻海鲜收货检查", "开头：收到冷冻海鲜，先别急着放进冰箱。画面：检查包装、标签、温度、破损和批次。旁白：如发现解冻、破损或信息不清，应立即停止食用并联系人工处理。结尾：收藏这份收货清单。"),
        ("餐厅厨师测试", "开头：餐厅选虾和鱼柳，为什么不能只看每千克价格？画面：称重、解冻、烹饪、成品称重。旁白：真正需要比较的是净重、冰衣、出成率、稳定性和到货。只有获得真实授权与数据后才发布结果。结尾：有采购需求可通过公开商务入口联系。"),
    ]
    for index, (title, body) in enumerate(scripts, 1):
        doc.add_heading(f"{index}. {title}", level=2)
        add_callout(doc, "脚本", body)
    doc.add_heading("审校清单", level=1)
    add_table(doc, ["检查", "通过标准"], [["商品", "画面、规格、包装与页面一致"], ["价格库存", "来自当天真值，或明确不展示"], ["语言", "中文／英文／尼泊尔语含义一致，母语者已审"], ["食品安全", "储存、解冻、过敏与售后表达准确"], ["授权", "人物、门店、音乐、图片和评价都有使用许可"], ["行动入口", "链接可用，页面显示可送区域和真实条件"]], [1.2, 5.3], 8.5)
    doc.add_heading("复盘方式", level=1)
    add_step_list(doc, ["同一选题只改一个核心变量", "至少比较开头、场景或时长中的一项", "优先看收藏、有效询问和商品页访问", "播放高但没有有效询问时不扩大", "出现误解、投诉或规则变化立即停止并修正"])
    add_sources(doc, "S22,S23,S24,S25,S26,S29,S30,S31")
    save(doc, "10_TikTok内容选题与脚本样例库.docx")


def doc_ai_architecture():
    sections = ["设计目标", "三个底座", "八个模块", "最小闭环", "人工确认", "验收标准", "来源"]
    doc = new_doc("AI销售系统总体架构", "先保证信息真实，再让 AI 提升研究、沟通与复盘效率", compact=False)
    add_toc(doc, sections)
    doc.add_heading("设计目标", level=1)
    add_callout(doc, "核心原则", "AI不能替代真实供应、许可、冷链和商务批准。系统价值在于把公开研究、真实商品、沟通记录和履约结果放到一条可追溯链路里。", PALE_GREEN)
    add_picture(doc, ASSETS / "ai_flow.png", "图 1｜研究、供应、AI草拟、人工确认与现实动作之间的关系。")
    doc.add_heading("三个底座", level=1)
    add_table(doc, ["底座", "作用", "来源", "异常处理", "责任"], d.AI_FOUNDATIONS, [1.25, 2.55, 1.2, 1.4, 0.7], 7.8)
    doc.add_heading("八个模块", level=1)
    rows = []
    for module in d.AI_MODULES:
        rows.append([module["name"], module["goal"], module["inputs"], module["outputs"], module["human"], module["fallback"], module["acceptance"]])
    add_table(doc, ["模块", "目标", "读取信息", "产出", "人工确认", "异常回退", "验收"], rows, [0.8, 1.4, 1.3, 1.2, 1.3, 1.25, 1.2], 6.5)
    doc.add_heading("最小闭环", level=1)
    add_step_list(
        doc,
        [
            "真值中心只放 10-20 个经过核实的商品和可送区域。",
            "市场情报模块筛出一个城市、两类采购群体和 20-50 个公开组织。",
            "产品匹配模块只根据真实供应资料给出候选。",
            "获客模块生成草稿，人工发送 15-25 条首轮联系。",
            "咨询与销售模块收集需求，正式报价和样品继续由人工确认。",
            "CRM记录有效沟通、拒绝、样品、试单、配送、回款和补货。",
            "数据分析模块在 30、60、90 天给出继续、调整或暂停建议。",
        ],
    )
    doc.add_heading("人工确认", level=1)
    add_table(
        doc,
        ["高风险动作", "为什么必须确认", "确认人"],
        [
            ["正式报价", "涉及价格、库存、税费、有效期和履约承诺", "销售与供应链负责人"],
            ["样品与试单", "涉及真实成本、食品安全和配送", "项目与供应链负责人"],
            ["平台发布与广告", "涉及公开承诺、账号规则和预算", "内容、合规与预算负责人"],
            ["退款、投诉与冷链异常", "涉及食安、赔付和声誉", "售后、供应链与项目负责人"],
            ["账期、独家、合同与大额订单", "涉及法律和回款风险", "项目决策组"],
            ["扩城、扩品与放量", "可能放大库存、履约和合规风险", "项目决策组"],
        ],
        [1.5, 3.1, 1.9],
        8.3,
    )
    doc.add_heading("验收标准", level=1)
    add_table(
        doc,
        ["部分", "最低验收"],
        [
            ["真值中心", "价格、库存、配送和合规缺失时能阻断输出"],
            ["问答", "常见问题准确率不低于 95%，误承诺为 0"],
            ["草稿审批", "未批准的首次联系、报价与发布不能发送"],
            ["CRM", "90%以上活跃记录有下一步与日期"],
            ["日志", "能看到来源、真值版本、草稿、批准、发送和结果"],
            ["回退", "平台、支付、库存、价格和冷链异常都有可执行处理"],
        ],
        [1.35, 5.15],
        8.5,
    )
    add_sources(doc, "S17,S18,S19,S20,S21,S22,S23,S24,S27")
    save(doc, "11_AI销售系统总体架构.docx")


def doc_ai_modules():
    sections = ["数据流总览", "模块说明", "权限与日志", "异常与回退", "对接顺序", "验收", "来源"]
    doc = new_doc("AI Agent功能与数据流设计", "八个 AI 协作模块如何读取、判断、输出和转人工", compact=True)
    add_toc(doc, sections)
    doc.add_heading("数据流总览", level=1)
    add_picture(doc, ASSETS / "ai_flow.png", "所有外部承诺都要经过人工确认；结果回流用于修正判断。")
    add_table(
        doc,
        ["起点", "经过", "终点", "回流"],
        [
            ["官方与公开网页", "市场情报、来源索引", "城市、品类、渠道候选", "价格与规则变化"],
            ["供应方与进口方资料", "真值中心、产品匹配", "可售商品与阻断清单", "库存、履约、退货"],
            ["公开组织与消费者询问", "获客、咨询、CRM", "草稿、报价准备与人工待办", "回复、拒绝、订单与补货"],
            ["内容与广告", "发布审批、转化追踪", "站外页面与订单", "询问、成交、成本与投诉"],
        ],
        [1.4, 2, 1.8, 1.3],
        8.2,
    )
    doc.add_heading("模块说明", level=1)
    for index, module in enumerate(d.AI_MODULES, 1):
        doc.add_heading(f"{index}. {module['name']}", level=2)
        add_table(
            doc,
            ["问题", "说明"],
            [
                ["目标", module["goal"]],
                ["读取信息", module["inputs"]],
                ["判断规则", module["rules"]],
                ["输出", module["outputs"]],
                ["何时触发", module["trigger"]],
                ["人工确认", module["human"]],
                ["记录", module["log"]],
                ["权限", module["permission"]],
                ["异常回退", module["fallback"]],
                ["验收", module["acceptance"]],
            ],
            [1.1, 5.4],
            8.25,
        )
    doc.add_heading("权限与日志", level=1)
    add_table(
        doc,
        ["角色", "可以做", "不能做", "必须记录"],
        [
            ["研究协作", "读取公开网页、形成候选和引用", "访问非公开资料、绕过登录与验证码", "来源、时间、判断和变化"],
            ["销售协作", "生成联系、回答与报价草稿", "直接发送、修改真值或承诺例外", "使用的真值版本和人工决定"],
            ["内容协作", "生成选题、脚本和版本", "发布、编造商品与评价", "素材来源、审校与发布结果"],
            ["分析协作", "汇总成交、复购、履约与成本", "用缺失数据做确定结论", "口径、缺失项和建议"],
            ["人工负责人", "批准、驳回、修改与现实执行", "删除审计记录", "决定、理由、时间和结果"],
        ],
        [1.1, 2.2, 1.75, 1.45],
        8.2,
    )
    doc.add_heading("异常与回退", level=1)
    add_table(doc, ["异常", "系统动作", "人工动作"], [["价格或库存过期", "停止报价并标红", "从供应方重新确认"], ["平台功能不可用", "切回站外或自然渠道", "复核账号和官方规则"], ["支付失败", "保留订单但不出库", "人工对账或换收款方式"], ["冷链异常", "停止同批次订单", "隔离、调查、退款或召回"], ["误答或误承诺", "终止自动回答并保留记录", "主动更正、处理影响并复盘"], ["数据冲突", "显示冲突并不自动合并", "决定采用哪一项并说明理由"]], [1.55, 2.5, 2.45], 8.3)
    doc.add_heading("对接顺序", level=1)
    add_step_list(doc, ["供应链真值与合规闸门", "来源索引与公开名单", "产品匹配和审批队列", "CRM与咨询转人工", "订单、库存、支付和配送", "内容生产与站外页面", "广告闸门与数据分析"])
    doc.add_heading("验收", level=1)
    add_callout(doc, "通过标准", "不是“模块能运行”，而是能在资料缺失、价格过期、库存不足、平台不可用、支付失败和冷链异常时正确停下，并把问题交给明确责任人。", PALE_GOLD)
    add_sources(doc, "S17,S18,S19,S20,S21,S22,S23,S24,S27")
    save(doc, "12_AI_Agent功能与数据流设计.docx")


def doc_crm():
    sections = ["漏斗口径", "核心信息项", "状态变化", "每日工作台", "复购与停止", "权限与隐私", "来源"]
    doc = new_doc("CRM信息项与销售漏斗设计", "让名单、沟通、样品、订单和补货保持清晰分层", compact=True)
    add_toc(doc, sections)
    doc.add_heading("漏斗口径", level=1)
    add_table(
        doc,
        ["状态", "进入条件", "退出到下一步的条件"],
        [
            ["候选", "有组织名称、城市、类别和公开来源", "主体、入口与适配性已复核"],
            ["已核实", "公开资料有效且可联系", "人工实际发送首次联系"],
            ["已联系", "有发送记录", "收到真实回复并讨论需求"],
            ["已沟通", "有需求、规格、用量或采购流程信息", "样品或报价获批准"],
            ["已试样", "真实样品已送达并有反馈", "真实试单成立"],
            ["已试单", "有订单、收款与履约记录", "再次采购或补货"],
            ["已补货", "有第二次真实采购", "持续复购或长期协议"],
            ["暂停／淘汰", "拒绝、无需求、负贡献、违规或异常", "人工决定是否在未来重启"],
        ],
        [1, 2.75, 2.75],
        8.4,
    )
    doc.add_heading("核心信息项", level=1)
    add_table(doc, ["分组", "信息项", "说明", "来源", "更新", "是否必填", "用途"], d.CRM_INFO_ITEMS, [0.75, 1.05, 1.7, 1, 0.7, 0.8, 1], 7.1)
    doc.add_heading("状态变化", level=1)
    add_callout(doc, "防止误报", "目录里存在不等于已核实，已发送不等于已沟通，寄出样品不等于订单，订单不等于补货。任何状态升级必须有时间和证据。")
    add_table(
        doc,
        ["动作", "需要证据", "负责人"],
        [["候选→已核实", "近期公开页面与入口", "市场人员"], ["已核实→已联系", "真实发送记录", "销售人员"], ["已联系→已沟通", "对方回复或会议记录", "销售人员"], ["已沟通→已试样", "样品批准、寄送和签收", "销售与供应链"], ["已试样→已试单", "订单、收款与履约", "销售、财务与履约"], ["已试单→已补货", "第二次真实订单", "CRM自动识别＋人工复核"]], [1.35, 3.5, 1.65], 8.3)
    doc.add_heading("每日工作台", level=1)
    add_table(doc, ["队列", "进入条件", "当日动作"], [["待核实", "公开来源不足或过期", "补证或降级"], ["待人工发送", "草稿完成且真值通过", "审核与发送"], ["待报价", "需求清楚但商务条件未批准", "补齐价格、库存、税费与有效期"], ["待样品", "样品方向确认", "批准成本、安排发货与反馈"], ["待售后", "温控、破损、少件、延误或退款", "优先人工处理"], ["待复购", "接近实际消耗周期", "个性化提醒"], ["停止联系", "明确拒绝或要求停止", "永久阻断自动跟进"]], [1.15, 2.4, 2.95], 8.3)
    doc.add_heading("复购与停止", level=1)
    add_step_list(doc, ["按真实消耗周期设置提醒，不按固定天数骚扰", "提醒前先检查库存、价格和配送", "补货以真实第二次订单为准", "明确拒绝、投诉未解决或许可不足时停止", "沉默对象在 90 天内不自动重启"])
    doc.add_heading("权限与隐私", level=1)
    add_table(doc, ["原则", "做法"], [["最少收集", "只保留完成沟通、订单、配送和售后所需信息"], ["公开与自愿", "B2B优先公开商务入口；消费者信息来自自愿下单或咨询"], ["分级权限", "研究、销售、财务、供应链和管理层只看所需部分"], ["停止权", "明确拒绝后立即停止后续联系"], ["可追溯", "保留状态变化、读取、修改、批准和发送记录"], ["删除与保留", "按当地法律和合同制定周期，过期后安全删除"]], [1.2, 5.3], 8.5)
    add_sources(doc, "S17,S18,S19")
    save(doc, "13_CRM信息项与销售漏斗设计.docx")


def doc_supply():
    sections = ["上线闸门", "供应链信息清单", "进口与标签", "系统对接", "职责边界", "报关行问题单", "来源"]
    doc = new_doc("供应链资料与数据接口清单", "没有真实供应、许可和冷链，就不进入正式销售", compact=True)
    add_toc(doc, sections)
    doc.add_heading("上线闸门", level=1)
    add_callout(doc, "P0原则", "商品、合规、价格、库存和履约中的任何 P0 信息不完整时，不生成正式报价、不公开上架、不收款。", PALE_RED)
    add_table(doc, ["阶段", "必须完成", "未完成时"], [["研究", "公开市场与候选商品", "可以继续研究"], ["对外介绍", "真实商品、素材和可送范围", "只做通用教育"], ["正式报价", "价格、库存、税费、有效期、最低订货量", "不得发出"], ["上架收款", "进口、标签、支付、库存、配送和售后", "不得上线"], ["放量", "稳定补货、复购、毛利和异常处理", "不扩大"]], [1.05, 3.45, 2], 8.4)
    doc.add_heading("供应链信息清单", level=1)
    add_table(doc, ["优先级", "类别", "需要提供", "最晚时间", "载体", "更新频率", "影响"], d.SUPPLY_INFO_ITEMS, [0.6, 0.8, 2.25, 1, 1, 0.9, 1.25], 7.1)
    doc.add_heading("进口与标签", level=1)
    add_table(
        doc,
        ["项目", "已确认", "待验证"],
        [
            ["进口许可", "食品进口需要走 DFTQC 许可路径；2026 年公告指向 Nepal National Single Window", "每个商品需要的具体资料与办理时间"],
            ["EXIM Code", "尼泊尔进口主体需要有效海关注册", "实际主体、有效期与授权范围"],
            ["海关归类", "活鲜、鲜冷、冷冻海鲜通常从第 03 章核；加工制备品可能在第 16 章", "每个商品准确编码、当期税率和附加税费"],
            ["标签", "包装食品至少要清楚显示批号、生产日期、有效期和净重／计量", "是否必须尼泊尔语、完整成分和特定格式"],
            ["过敏提示", "本轮未找到海鲜专门强制条文", "鱼类、甲壳类、贝类的当地强制要求"],
            ["温控", "需要按商品和进口许可要求保留储运记录", "尼泊尔针对各形态的法定阈值"],
            ["隐私", "收集个人资料要有明确目的、同意和保护", "订单页和电商规则的具体展示方式"],
        ],
        [1.1, 3, 2.4],
        8.15,
    )
    doc.add_heading("系统对接", level=1)
    add_table(
        doc,
        ["对接", "最少读取", "最少写回", "异常回退"],
        [
            ["商品与价格", "商品编号、名称、规格、净重、当前价格、有效期", "采用的价格版本", "过期即停止报价"],
            ["库存", "可售、锁定、缺货、补货日期", "订单锁定与释放", "不同步时转人工并关闭在线购买"],
            ["配送", "可送区域、费用、截止时间、预计到达", "实际揽收、送达与异常", "超区或超时不承诺"],
            ["订单", "订单、付款、退款与取消", "供应、发货与完成状态", "支付失败不出库"],
            ["批次与质量", "批号、日期、温控与质检", "签收温度、破损、隔离与召回", "异常批次立即停售"],
        ],
        [1.1, 2.1, 1.75, 1.55],
        8.15,
    )
    doc.add_heading("职责边界", level=1)
    add_table(doc, ["责任方", "负责", "不应替代"], [["中国供应方", "商品、价格、库存、资质、包装、保质期与温控要求", "尼泊尔进口许可与本地经营责任"], ["尼泊尔进口方／报关行", "海关归类、许可、报关、税费与口岸资料", "线上内容与沟通"], ["仓储与配送方", "收货、温控、库存、拣货、配送与异常", "食品与商务承诺"], ["线上销售团队", "商品页面、公开名单、沟通、订单协同与复盘", "凭空补齐合规、库存和配送"], ["项目负责人", "批准报价、样品、预算、例外和阶段决策", "删除审计记录"]], [1.55, 3.1, 1.85], 8.25)
    doc.add_heading("报关行问题单", level=1)
    add_step_list(doc, ["每个商品准确的海关编码是什么，落在第 03 章还是第 16 章？", "当前财年的关税、增值税、附加税和其他费用是什么？", "本项目的进口转售、仓储、分装与零售是否需要额外食品许可？", "是否需要卫生证、检验报告、原产地证、冷链证明或特定口岸文件？", "标签是否必须尼泊尔语，过敏提示和成分表有什么强制要求？", "不同形态的温控、保质期和抽检要求是什么？", "经第三地转运是否影响原产地与单证？", "首单前是否建议做样品清关或文件预审？"])
    add_sources(doc, "S10,S11,S12,S13,S14,S15,S16,S17,S18,S19")
    save(doc, "14_供应链资料与数据接口清单.docx")


def doc_roadmap():
    sections = ["总目标", "阶段路线", "周度安排", "角色分工", "例会口径", "90天交付", "来源"]
    doc = new_doc("30_60_90天执行计划", "从真值与合规，到试单、复购和阶段决策", compact=True)
    add_toc(doc, sections)
    doc.add_heading("总目标", level=1)
    add_picture(doc, ASSETS / "roadmap.png", "90 天目标是验证可复制场景，不是追求全国声量。")
    add_callout(doc, "北极星", "真实成交＋真实补货／复购＋可持续毛利＋稳定履约。播放、名单、已读、样品和首单都不能单独代表成功。", PALE_GREEN)
    doc.add_heading("阶段路线", level=1)
    add_table(doc, ["时间", "目标", "动作", "产出", "通过标准", "阻断", "责任"], d.ROADMAP, [0.65, 1.05, 1.75, 1.4, 1.35, 1.25, 1.05], 6.5)
    doc.add_heading("周度安排", level=1)
    add_table(
        doc,
        ["周", "供应链与合规", "市场与销售", "内容与系统", "决策"],
        [
            ["W1", "进口主体、商品与 P0 清单", "复核首城与首批群体", "建立真值和来源索引", "哪些可以继续研究"],
            ["W2", "海关编码、标签与冷链预审", "50 个高优先公开组织", "CRM与草稿审批", "哪些可以联系"],
            ["W3", "真实商品与样品边界", "人工联系 8-12 个组织", "三个内容角度", "回复与阻力"],
            ["W4", "配送演练准备", "累计 15-25 条人工联系", "复盘内容到询问", "30天继续或调整"],
            ["W5-W6", "价格、库存和订单对接", "推进样品与短会", "咨询转人工演示", "系统是否可靠"],
            ["W7-W8", "谷地真实履约", "3-5 个样品或试单", "真实商品内容", "60天继续或调整"],
            ["W9-W10", "损耗、退货与补货", "追踪复购，消费者小测", "站外路径与周报", "可复制场景"],
            ["W11-W12", "供应上限与扩城条件", "补货与回款复盘", "阶段分析", "继续／调整／暂停"],
        ],
        [0.75, 1.55, 1.55, 1.45, 1.2],
        7.7,
    )
    doc.add_heading("角色分工", level=1)
    add_table(doc, ["角色", "每日", "每周", "阶段决策"], [["项目负责人", "处理高风险批准", "主持复盘与阻断清单", "决定预算、扩城、扩品和暂停"], ["供应链负责人", "更新价格、库存与配送", "核合规、温控和异常", "确认供应能力"], ["B2B负责人", "核实、联系和记录", "复盘回复、样品和试单", "建议群体与商品"], ["内容负责人", "制作与审校", "复盘有效询问", "建议渠道与选题"], ["系统与数据负责人", "监测同步和错误", "出周报与缺失信息", "验证决策证据"], ["尼泊尔本地人员", "语言、现场和关系处理", "核平台与履约", "复核当地可行性"]], [1.3, 1.9, 1.8, 1.5], 8.1)
    doc.add_heading("例会口径", level=1)
    add_table(doc, ["必须回答", "证据"], [["本周新增了什么已确认事实？", "来源或真实业务记录"], ["哪些事项仍待验证？", "责任人和完成日期"], ["哪些组织从哪一状态进入下一状态？", "沟通、样品、订单或补货证据"], ["哪些商品出现价格、库存或冷链问题？", "真值与异常记录"], ["下周只验证哪三个问题？", "明确实验和停止条件"]], [2.75, 3.75], 8.4)
    doc.add_heading("90天交付", level=1)
    add_table(doc, ["交付", "通过标准"], [["可售商品池", "真实商品、合规、价格、库存和配送可追溯"], ["高优先组织池", "公开来源、联系状态和下一步清楚"], ["试单与补货证据", "订单、履约、回款与复购分层"], ["内容与站外路径", "不依赖 TikTok Shop 也能完成询问和下单"], ["AI协作系统", "缺失时会阻断，高风险动作需要人工批准"], ["阶段决策", "有继续、调整或暂停的明确理由"]], [2.15, 4.35], 8.4)
    add_sources(doc, "S11,S12,S14,S17,S18,S20,S22,S23,S24,S27,S28,S36,S37,S38")
    save(doc, "15_30_60_90天执行计划.docx")


def doc_risks():
    sections = ["风险矩阵", "扩大与停止", "预算闸门", "应急顺序", "决策记录", "来源"]
    doc = new_doc("风险_停止投入_扩大投入标准", "把合规、冷链、平台、价格和自动化风险放到预算之前", compact=True)
    add_toc(doc, sections)
    doc.add_heading("风险矩阵", level=1)
    add_table(doc, ["风险", "触发", "后果", "预防", "责任", "应急", "等级"], d.RISKS, [1.05, 1.35, 1.2, 1.55, 0.9, 1.35, 0.5], 6.6)
    doc.add_heading("扩大与停止", level=1)
    add_table(doc, ["判断", "触发标准", "动作"], d.SCALE_RULES, [1.1, 3.8, 1.6], 8.1)
    doc.add_heading("预算闸门", level=1)
    add_table(doc, ["投入类型", "计算方式", "前提", "核心指标", "停止线"], d.BUDGET_MODEL, [1.1, 1.3, 1.55, 1.25, 1.3], 7.5)
    doc.add_heading("应急顺序", level=1)
    add_step_list(doc, ["立即停止受影响商品、批次、页面或投放", "隔离库存与订单，保护食品安全和付款", "由人工联系受影响购买方或合作单位", "核实原因、范围、责任和证据", "按情况退款、召回、补发或停止合作", "修复真值、流程和权限后再评估恢复", "把事件写入阶段决策，不删除历史"])
    doc.add_heading("决策记录", level=1)
    add_table(doc, ["日期", "问题", "证据", "决定", "责任", "复查日"], [["待填写", "例如：是否扩大虾类测试", "试单、复购、毛利、冷链和投诉", "继续／调整／暂停", "项目决策组", "待填写"], ["待填写", "例如：是否进入 Pokhara", "谷地补货、当地冷链和名单验证", "继续／调整／暂停", "项目决策组", "待填写"]], [0.8, 1.8, 1.8, 0.9, 0.75, 0.45], 8.1)
    add_sources(doc, "S11,S12,S13,S14,S17,S18,S22,S23,S24,S25,S26,S27")
    save(doc, "16_风险_停止投入_扩大投入标准.docx")


def doc_sources_index():
    sections = ["来源统计", "全部来源", "证据使用规则", "动态复核", "当前缺口"]
    doc = new_doc("来源与引用索引", f"{len(d.SOURCES)}个公开来源的机构、时间、用途与可信度", compact=True)
    add_toc(doc, sections)
    doc.add_heading("来源统计", level=1)
    official = sum(1 for s in d.SOURCES if s["official"])
    platform = sum(1 for s in d.SOURCES if s["platform_official"])
    company = sum(1 for s in d.SOURCES if s["company"])
    other = len(d.SOURCES) - official - platform - company
    add_table(doc, ["类别", "数量", "主要用途"], [["尼泊尔官方与国际组织", official, "法律、统计、贸易、支付、城市与旅游"], ["平台官方", platform, "TikTok 与 Daraz 能力"], ["企业官网", company, "在售、价格、配送与零售入口"], ["协会、学术与数据库", other, "酒店、餐饮、商会、水产与贸易补充"]], [2, 1, 3.5], 8.5)
    add_sources(doc, None, "全部来源")
    doc.add_heading("证据使用规则", level=1)
    add_table(doc, ["来源类型", "可支持", "不能单独支持"], [["政府与监管", "法律入口、官方流程、统计与公告", "具体商品已获批或实际已办妥"], ["平台官方", "市场清单、功能条件和帮助说明", "某个实际账号已经开通"], ["企业官网", "公开在售、标价、业务与公开入口", "全市场份额、持续库存与采购意向"], ["协会与名录", "成员和组织发现", "成员愿意采购"], ["贸易再发布", "历史进口结构和比较", "最新市场规模、利润或未来销量"]], [1.25, 2.55, 2.7], 8.3)
    doc.add_heading("动态复核", level=1)
    add_table(doc, ["内容", "建议频率", "责任"], [["价格与库存", "每 1-2 周及报价前", "市场与供应链"], ["平台功能与地区", "每月及开户前", "平台负责人"], ["进口、关税与许可", "每批、每个商品及政策变化时", "进口方与报关行"], ["支付与配送", "接入前、每季度和异常后", "本地运营"], ["组织入口", "联系前 48 小时内", "B2B负责人"]], [2.3, 2.1, 2.1], 8.4)
    doc.add_heading("当前缺口", level=1)
    add_open_items(doc)
    save(doc, "17_来源与引用索引.docx")


def doc_codex_tasks():
    sections = ["执行原则", "任务总表", "依赖顺序", "验收方法", "测试场景", "交接资料"]
    doc = new_doc("后续Codex5.5执行任务包", "将研究结论转成可测试、可审计、可回退的系统任务", compact=True)
    add_toc(doc, sections)
    doc.add_heading("执行原则", level=1)
    add_table(doc, ["原则", "要求"], [["真值先行", "没有真实价格、库存、配送和合规时，只允许测试数据和阻断演示"], ["API优先", "优先官方接口、授权后台、数据导出和 webhook，不以屏幕模拟作为核心链路"], ["草稿优先", "联系、报价、发布和退款默认只生成草稿"], ["人工闸门", "现实动作必须有明确批准人、时间和记录"], ["最小权限", "研究、编辑、批准、发送、退款和管理分开"], ["可回退", "价格过期、缺货、平台不可用、支付失败和冷链异常都要安全停止"]], [1.25, 5.25], 8.6)
    doc.add_heading("任务总表", level=1)
    add_table(doc, ["编号", "任务", "目标", "对应资料", "验收", "依赖"], d.CODEX_TASKS, [0.45, 1.2, 2.1, 1.4, 1.9, 0.75], 7.1)
    doc.add_heading("依赖顺序", level=1)
    add_step_list(doc, ["T01 真值中心与 T02 来源索引", "T03 公开名单与 T04 产品匹配", "T05 审批队列与 T06 CRM", "T07 咨询转人工与 T11 订单库存", "T08 内容流程与 T09 站外成交页", "T10 本地支付与 T12 广告闸门", "T13 决策、T14 权限审计与 T15 回退演练"])
    doc.add_heading("验收方法", level=1)
    add_table(doc, ["层级", "证明"], [["单元", "每条规则对正常、缺失、过期和冲突信息都有结果"], ["集成", "真值、CRM、订单、支付和配送能按测试流程连通"], ["权限", "未授权角色不能报价、发送、投放、退款或修改真值"], ["审计", "来源、版本、草稿、批准、发送和结果可回看"], ["业务演练", "用合成对象完成候选→草稿→批准→试单→异常→回退"], ["真实小样", "只有 P0 资料完成后，才在一个城市和少量商品上运行"]], [1.05, 5.45], 8.4)
    doc.add_heading("测试场景", level=1)
    add_table(doc, ["场景", "期望结果"], [["价格缺失", "不生成正式报价，只列缺失信息"], ["库存为零", "关闭购买并提供人工处理，不承诺补货日期"], ["地址超区", "不接单，显示当前可送范围"], ["TikTok Shop 不可用", "站外页面仍能完成询问、付款与订单"], ["支付失败", "订单不出库，保留对账记录"], ["冷链温度异常", "隔离同批次并启动人工调查"], ["明确拒绝联系", "后续自动任务立即停止"], ["来源冲突", "显示冲突并等待人工决定"]], [2.2, 4.3], 8.5)
    doc.add_heading("交接资料", level=1)
    add_table(doc, ["资料", "责任方", "状态"], [["供应链 P0 资料", "供应方与进口方", "待提供"], ["当地合规与报关答复", "进口方、报关行、顾问", "待提供"], ["真实账号与授权", "本地运营", "待提供"], ["测试用合成数据", "系统团队", "可先建立"], ["验收人和决策人", "项目负责人", "待确认"], ["上线与回退负责人", "各业务负责人", "待确认"]], [2.6, 2.1, 1.8], 8.5)
    add_sources(doc, "S14,S17,S18,S19,S20,S21,S22,S23,S24,S27")
    save(doc, "18_后续Codex5.5执行任务包.docx")


def doc_source_readme():
    sections = ["资料用途", "来源分级", "引用方式", "更新规则", "已知限制"]
    doc = new_doc("来源使用说明", "如何复核、引用和更新本项目的公开研究来源", compact=True)
    add_toc(doc, sections)
    doc.add_heading("资料用途", level=1)
    add_callout(doc, "说明", "本文件夹只放可供合作讨论的来源说明。原始网页仍由来源机构维护，引用时应打开链接核对最新内容。")
    doc.add_heading("来源分级", level=1)
    add_table(doc, ["级别", "来源", "典型用途"], [["优先", "尼泊尔政府、监管、央行、统计、国际组织", "规则、流程、统计和公共基础设施"], ["平台官方", "TikTok、Daraz 等官方帮助页", "平台能力、市场清单和条件"], ["企业官网", "零售、配送、供应与公开商品页", "在售、标价、覆盖和入口"], ["补充", "协会、学术与贸易再发布", "组织发现、背景与历史结构"]], [1.05, 2.8, 2.65], 8.5)
    doc.add_heading("引用方式", level=1)
    add_step_list(doc, ["用 S01-S40 编号回到《17_来源与引用索引》", "打开原始链接并确认页面、日期和口径", "区分事实、工作计算和推测", "动态价格、库存、平台能力和税费必须在行动前复核", "企业页面只证明公开展示，不推断全市场"])
    doc.add_heading("更新规则", level=1)
    add_table(doc, ["内容", "频率", "变化后"], [["价格和库存", "每 1-2 周、报价前", "更新样本和价格版本"], ["平台市场与账号能力", "每月、开户前", "更新能力矩阵与替代路径"], ["法律、许可和税则", "每批、每商品、政策变化时", "由当地专业人员复核"], ["城市与组织入口", "每季度、联系前", "去除失效记录"], ["支付与配送", "接入前、异常后", "重跑全链路测试"]], [2.1, 2.1, 2.3], 8.5)
    doc.add_heading("已知限制", level=1)
    add_table(doc, ["限制", "影响"], [["2022 HS 03 是历史贸易口径", "不能代表当前全部海鲜市场"], ["企业网页会变", "价格、库存和配送只能作当日快照"], ["公开名单不是采购意向", "必须人工联系和验证"], ["尼泊尔当地法律与报关细节未逐商品确认", "不能发货、上架或做正式报价"], ["真实供应资料缺失", "商品评分是候选排序，不是利润结论"]], [2.6, 3.9], 8.5)
    add_sources(doc, None, "来源总表")
    save(doc, "00_来源使用说明.docx", "research_sources")


def doc_full():
    sections = [
        "总决策摘要",
        "研究边界与证据",
        "市场与贸易",
        "城市、人群与场景",
        "品类与价格",
        "B2B销售",
        "B2C销售",
        "TikTok专项",
        "支付、配送与合规",
        "AI销售系统",
        "CRM与供应链",
        "30／60／90天路线",
        "风险与投入标准",
        "执行任务",
        "来源与待验证事项",
    ]
    doc = new_doc("尼泊尔海鲜AI线上销售系统总方案", "完整版｜中国供应链进入尼泊尔的市场研究、渠道策略、AI系统与90天试点", compact=False)
    add_toc(doc, sections)
    doc.add_heading("总决策摘要", level=1)
    add_callout(doc, "主结论", "项目值得进入受控试点，但不适合立即做全国铺货。第一阶段只做加德满都谷地、冷冻品、真实供应、小范围采购与配送验证。", PALE_GREEN)
    add_table(
        doc,
        ["决策问题", "建议", "状态与边界"],
        [
            ["进入还是等待", "进入 90 天受控试点", "公开研究支持试点；真实商品与当地条件待核"],
            ["首个区域", "Kathmandu＋Lalitpur＋Bhaktapur 作为一个冷链单元", "工作判断，需真实配送演练"],
            ["第二城市", "Pokhara", "谷地出现稳定补货后才进入"],
            ["首批商品", "巴沙鱼柳、HLSO／IQF虾、鱿鱼筒／圈、火锅或烧烤组合", "需逐品确认成本、许可、标签和冷链"],
            ["高价测试", "三文鱼小量预订", "不能当主销量商品"],
            ["暂缓商品", "活鲜、冰鲜、整蟹／龙虾、贝类、章鱼、海参、鲍鱼", "冷链、价格、认知和证据不足"],
            ["B2B主线", "酒店、亚洲餐厅、冷冻经销、精品零售", "先样品、厨师测试、出成率和补货"],
            ["B2C主线", "谷地家庭聚餐、火锅烧烤、年轻即烹需求", "只在已确认可送区域小测"],
            ["TikTok定位", "发现、教育、询问与站外引导", "Shop 不在官方 Nepal 市场清单"],
            ["收款", "Fonepay QR、eSewa、Khalti、货到付款", "由尼泊尔本地主体接入与对账"],
            ["扩大条件", "至少 8 个试单单位，补货／复购不低于 30%，至少一个可复制场景", "内部试点闸门，不是市场事实"],
            ["停止条件", "重大食安／合规／冷链问题，或 50 个已核实对象两轮实验后仍少于 3 个试单", "立即停止或暂停"],
        ],
        [1.35, 3.15, 2],
        8.1,
    )
    doc.add_heading("研究边界与证据", level=1)
    add_basis(doc)
    add_table(
        doc,
        ["已经完成", "部分完成", "尚未执行"],
        [
            ["公开市场、贸易、城市、支付、平台、竞品、合规与协会研究", "城市、品类、渠道与系统方案已形成候选判断", "真实账号、投放、联系、样品、订单、支付和配送"],
            [f"{len(d.SOURCES)} 个公开来源索引", "25 个公开组织入口，尚未联系", "本地访谈、法律意见、报关答复和真实复购"],
            ["21 份合作讨论 DOCX", "商品评分缺少真实成本与库存", "正式市场规模、商业结果与全国扩张"],
        ],
        [2.25, 2.25, 2],
        8.3,
    )
    doc.add_heading("市场与贸易", level=1)
    add_table(doc, ["指标", "数值", "时间", "来源", "状态", "意义"], [row[:6] for row in d.MARKET_FACTS], [1.05, 1, 0.75, 1.15, 1, 2.3], 7.4)
    add_picture(doc, ASSETS / "import_sources.png", "图 1｜HS 03 进口来源结构（2022）。中国当年金额极低，需重新建立渠道与信任。")
    add_picture(doc, ASSETS / "import_categories.png", "图 2｜HS 03 进口品类结构（2022）。鱼和鱼柳占主导。")
    add_callout(doc, "规模边界", "2022 年 HS 03 进口约 636 万美元，只能作为历史贸易下限参考。它不等于全部海鲜消费，也不能直接转成可获得收入。", PALE_GOLD)
    doc.add_heading("城市、人群与场景", level=1)
    add_picture(doc, ASSETS / "city_scores.png", "图 3｜城市评分用于安排试点顺序，不代表城市销售规模。")
    add_table(doc, ["城市", "评分", "优先级", "动作", "理由", "可信度", "来源"], d.CITY_SCORES, [0.85, 0.45, 0.65, 0.85, 2.15, 0.55, 0.8], 7.1)
    doc.add_heading("B2B群体", level=2)
    add_table(doc, ["群体", "评分", "优先级", "需求逻辑", "首要验证", "来源"], d.B2B_SEGMENTS, [1.25, 0.45, 0.7, 2.1, 1.45, 0.55], 7.7)
    doc.add_heading("B2C群体", level=2)
    add_table(doc, ["群体", "评分", "优先级", "需求逻辑", "首要验证"], d.B2C_SEGMENTS, [1.55, 0.45, 0.75, 2.35, 1.4], 7.8)
    doc.add_heading("消费场景", level=2)
    add_table(doc, ["场景", "评分", "适合品类", "方向", "首要验证"], d.SCENARIO_SCORES, [1.5, 0.5, 2.1, 0.7, 1.7], 8.1)
    doc.add_heading("品类与价格", level=1)
    add_table(
        doc,
        ["商品方向", "类别", "形态", "B2B", "B2C", "评分", "结论", "公开价格", "依据"],
        [[p["name"], p["category"], p["form"], p["b2b"], p["b2c"], p["total"], p["conclusion"], p["price"], p["evidence"]] for p in d.PRODUCTS],
        [1.15, 0.6, 0.85, 0.4, 0.4, 0.4, 0.8, 1, 1.4],
        6.55,
    )
    doc.add_heading("公开价格样本", level=2)
    add_table(doc, ["商品", "平台／企业", "类别", "规格", "售价NPR", "折算NPR/kg", "页面状态", "来源", "可信度"], d.PRICE_RECORDS, [1.15, 0.9, 0.7, 0.65, 0.6, 0.65, 1.1, 0.45, 0.45], 6.45)
    add_callout(doc, "定价原则", "比较净重、冰衣、规格、处理方式、可食用比例、运输与损耗。供应方落地成本完成前，不以公开零售价反推利润。")
    doc.add_heading("B2B销售", level=1)
    add_callout(doc, "主线", "线上发现＋人工沟通＋真实样品＋可追溯履约＋补货判断。名单数量和样品数量不是最终结果。", PALE_GREEN)
    add_table(doc, ["渠道", "可用性", "适合对象", "推进方式", "AI作用", "风险", "指标", "来源"], [[r[0], r[1], r[2], r[4], r[5], r[6], r[7], r[8]] for r in d.B2B_CHANNELS], [0.8, 0.8, 1.25, 1.35, 0.6, 1.1, 0.8, 0.5], 6.45)
    doc.add_heading("公开组织入口", level=2)
    add_table(doc, ["组织", "城市", "类别", "可能关系", "公开入口", "公开联系", "可信度", "建议", "来源"], d.LEADS, [1.2, 0.75, 1, 0.85, 0.75, 1.05, 0.45, 1.3, 0.45], 6.2)
    add_step_list(doc, ["公开来源发现与复核", "产品匹配和草稿", "人工联系与需求确认", "正式报价和样品批准", "试单、履约与回款", "补货、调整或暂停"])
    doc.add_heading("B2C销售", level=1)
    add_table(doc, ["渠道", "可用性", "作用", "成交路径", "AI作用", "人工确认", "指标", "来源"], d.B2C_CHANNELS, [0.8, 1.05, 1.05, 1.3, 0.9, 1.05, 0.8, 0.45], 6.55)
    add_callout(doc, "建议路径", "内容或搜索→真实商品页→WhatsApp／Viber／电话→二维码／电子钱包／货到付款→谷地冷链→售后与复购。")
    doc.add_heading("TikTok专项", level=1)
    add_table(doc, ["能力", "状态", "证据或限制", "标记", "建议", "来源"], d.TIKTOK_CAPABILITIES, [1.1, 1.1, 2.15, 0.85, 1.65, 0.5], 7.4)
    add_table(doc, ["编号", "选题", "对象", "开头", "素材边界", "下一步", "指标"], d.CONTENT_IDEAS, [0.4, 1.05, 1.05, 1.5, 1.5, 1.35, 0.65], 6.75)
    add_callout(doc, "禁止误判", "平台可访问不等于 Shop、广告、LIVE、商品挂载或站内支付已开通。每项能力都要看官方清单和实际账号后台。", PALE_RED)
    doc.add_heading("支付、配送与合规", level=1)
    add_table(
        doc,
        ["主题", "已确认", "待验证", "行动"],
        [
            ["支付", "二维码与电子钱包在尼泊尔使用量高", "本地主体接入、费率、结算与退款", "Fonepay／eSewa／Khalti＋货到付款演练"],
            ["食品进口", "需要 DFTQC 进口许可路径和海关注册", "逐品资料、时间与口岸要求", "发货前由进口方和报关行书面确认"],
            ["海关编码与税费", "年度税则和第 03／16 章入口存在", "每个商品的准确编码和当期税费", "不做伪精确落地成本"],
            ["标签", "至少清楚显示批号、生产日期、有效期和净重", "语言、成分、过敏提示与具体格式", "逐品审稿"],
            ["冷链", "冷冻商品需要储运与异常记录", "法定阈值、实际路线、费用和服务范围", "谷地真实配送演练"],
            ["隐私", "个人资料收集需要目的、同意和保护", "订单页具体展示与保留周期", "最少收集、分级权限、停止联系"],
        ],
        [1.05, 2.05, 1.85, 1.55],
        8.05,
    )
    doc.add_heading("AI销售系统", level=1)
    add_picture(doc, ASSETS / "ai_flow.png", "图 4｜系统只负责整理、匹配、草拟和提醒；现实动作由人工确认。")
    add_table(doc, ["底座", "作用", "来源", "异常处理", "责任"], d.AI_FOUNDATIONS, [1.2, 2.55, 1.15, 1.4, 0.7], 7.7)
    add_table(doc, ["模块", "目标", "输入", "输出", "人工确认", "回退", "验收"], [[m["name"], m["goal"], m["inputs"], m["outputs"], m["human"], m["fallback"], m["acceptance"]] for m in d.AI_MODULES], [0.8, 1.35, 1.25, 1.15, 1.25, 1.25, 1.25], 6.35)
    doc.add_heading("CRM与供应链", level=1)
    add_callout(doc, "状态分层", "候选、已核实、已联系、已沟通、已试样、已试单和已补货必须分开记录；任何升级都要有真实证据。")
    add_table(doc, ["分组", "信息项", "说明", "来源", "更新", "必填", "用途"], d.CRM_INFO_ITEMS, [0.7, 1, 1.65, 0.95, 0.65, 0.75, 0.95], 6.85)
    add_table(doc, ["优先级", "类别", "需要提供", "最晚时间", "载体", "更新", "影响"], d.SUPPLY_INFO_ITEMS, [0.55, 0.75, 2.2, 0.95, 0.9, 0.8, 1.25], 6.9)
    doc.add_heading("30／60／90天路线", level=1)
    add_picture(doc, ASSETS / "roadmap.png", "图 5｜先验证真值和合规，再推进样品、试单与补货。")
    add_table(doc, ["时间", "目标", "动作", "产出", "通过标准", "阻断", "责任"], d.ROADMAP, [0.6, 0.95, 1.7, 1.3, 1.3, 1.2, 1], 6.3)
    doc.add_heading("风险与投入标准", level=1)
    add_table(doc, ["风险", "触发", "后果", "预防", "责任", "应急", "等级"], d.RISKS, [1, 1.3, 1.15, 1.5, 0.85, 1.25, 0.45], 6.35)
    add_table(doc, ["判断", "触发标准", "动作"], d.SCALE_RULES, [1.05, 3.85, 1.6], 8)
    add_table(doc, ["投入类型", "计算方式", "前提", "指标", "停止线"], d.BUDGET_MODEL, [1.05, 1.3, 1.5, 1.25, 1.4], 7.4)
    doc.add_heading("执行任务", level=1)
    add_table(doc, ["编号", "任务", "目标", "对应资料", "验收", "依赖"], d.CODEX_TASKS, [0.45, 1.15, 2.05, 1.35, 1.85, 0.75], 6.95)
    doc.add_heading("来源与待验证事项", level=1)
    add_open_items(doc)
    add_sources(doc, None, "完整来源索引")
    add_next_steps(doc, ["由中国供应方完成 P0 商品资料", "由尼泊尔进口方和报关行完成逐品预审", "确定谷地冷链、收款与售后责任", "选 10-20 个商品候选和 15-25 个公开组织做首轮验证", "30 天后按有效沟通、样品、试单和异常复盘", "90 天只在达到补货、毛利和履约闸门后扩大"])
    save(doc, "尼泊尔海鲜AI线上销售系统总方案_完整版.docx")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    create_charts()
    doc_overview()
    doc_market()
    doc_products()
    doc_priorities()
    doc_prices()
    doc_b2b_channels()
    doc_b2b_leads()
    doc_b2b_sop()
    doc_b2c_channels()
    doc_tiktok()
    doc_tiktok_content()
    doc_ai_architecture()
    doc_ai_modules()
    doc_crm()
    doc_supply()
    doc_roadmap()
    doc_risks()
    doc_sources_index()
    doc_codex_tasks()
    doc_full()
    doc_source_readme()
    print(f"Generated {len(list(OUT.rglob('*.docx')))} DOCX files in {OUT}")


if __name__ == "__main__":
    main()
