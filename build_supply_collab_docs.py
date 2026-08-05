from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import supply_collab_data as d


ROOT = Path("/Volumes/WD_BLACK/汾酒尼泊尔")
OUT = ROOT / "汾酒海鲜_尼泊尔线上销售_供应链协同与资料交付体系"
EVIDENCE = OUT / "evidence"
FONT = "Hiragino Sans GB"
BLUE = "1F4D78"
BLUE_2 = "2E74B5"
INK = "1F2937"
MUTED = "667085"
PALE = "EDF4FA"
PALE_GOLD = "FFF4CE"
PALE_RED = "FDECEC"
PALE_GREEN = "EAF6EF"
WHITE = "FFFFFF"


def clean(value) -> str:
    if value is None:
        return ""
    return (
        str(value)
        .replace("客户", "采购单位")
        .replace("字段", "信息项")
        .replace("用户", "线上运营团队")
        .replace("customer", "buyer")
        .replace("Customer", "Buyer")
        .replace("client", "partner")
        .replace("Client", "Partner")
    )


def set_run(run, size=10.5, bold=None, color=INK, italic=None):
    run.font.name = FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str, placeholder: str = "1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])


def add_hyperlink(paragraph, label: str, url: str):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([fonts, color, underline])
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = clean(label)
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_table_geometry(table, proportions: Sequence[float], table_width: int):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(proportions)
    widths = [max(300, round(table_width * part / total)) for part in proportions]
    widths[-1] += table_width - sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)


def new_doc(title: str, subtitle: str, *, landscape=False, status="部分完成") -> Document:
    doc = Document()
    sec = doc.sections[0]
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width = Inches(11)
        sec.page_height = Inches(8.5)
        sec.top_margin = Inches(0.55)
        sec.bottom_margin = Inches(0.55)
        sec.left_margin = Inches(0.55)
        sec.right_margin = Inches(0.55)
        doc._table_width = 14256
    else:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        doc._table_width = 9360
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.4)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE_2, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("汾酒＋海鲜尼泊尔项目｜合作执行资料"), size=8, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run(f"{d.ACCESS_DATE}  |  "), size=8, color=MUTED)
    add_field(footer, "PAGE")
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    cover = doc.add_paragraph()
    cover.paragraph_format.space_before = Pt(36 if landscape else 58)
    cover.paragraph_format.space_after = Pt(12)
    set_run(cover.add_run(clean(title)), size=23 if landscape else 26, bold=True, color=BLUE)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(20)
    set_run(sub.add_run(clean(subtitle)), size=12, color=MUTED)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(16)
    for label, value in (
        ("核验日期", d.ACCESS_DATE),
        ("当前状态", status),
        ("判断方式", "已确认 / 部分成立 / 待验证 / 推测"),
        ("版本", "v1.0 合作确认稿"),
    ):
        set_run(meta.add_run(f"{label}："), size=9, bold=True)
        set_run(meta.add_run(f"{value}\n"), size=9, color=MUTED)
    add_callout(
        doc,
        "使用提醒",
        "本文件用于责任确认、资料准备和上线判断。未取得真实资料、当地专业意见或平台后台结果的事项，均保持待确认，不作为正式销售承诺。",
        PALE_GOLD,
    )
    doc.add_page_break()
    return doc


def add_para(doc: Document, text: str, label: str | None = None, color=INK):
    p = doc.add_paragraph()
    if label:
        set_run(p.add_run(clean(label)), bold=True, color=BLUE)
    set_run(p.add_run(clean(text)), color=color)
    return p


def add_callout(doc: Document, label: str, text: str, fill=PALE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(9)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    set_run(p.add_run(f"{clean(label)}  "), size=10.5, bold=True, color=BLUE)
    set_run(p.add_run(clean(text)), size=10.5)
    return p


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Iterable],
    proportions: Sequence[float] | None = None,
    font_size=8.2,
    header_fill=BLUE,
):
    matrix = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(clean(header)), size=font_size, bold=True, color=WHITE)
    for row_index, row in enumerate(matrix):
        cells = table.add_row().cells
        for index in range(len(headers)):
            value = row[index] if index < len(row) else ""
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if isinstance(value, tuple) and len(value) == 2 and str(value[1]).startswith("http"):
                add_hyperlink(p, clean(value[0]), str(value[1]))
            else:
                set_run(p.add_run(clean(value)), size=font_size)
            if row_index % 2 == 1:
                shade(cells[index], "F8FAFC")
    if proportions is None:
        proportions = [1] * len(headers)
    set_table_geometry(table, proportions, doc._table_width)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_steps(doc: Document, items: Sequence[str]):
    add_table(
        doc,
        ["顺序", "行动"],
        [[str(i), item] for i, item in enumerate(items, 1)],
        [0.7, 5.8],
        9.1,
    )


def add_toc(doc: Document, sections: Sequence[str]):
    doc.add_heading("阅读目录", level=1)
    add_table(doc, ["章节", "内容"], [[str(i), name] for i, name in enumerate(sections, 1)], [0.7, 5.8], 9.2)
    p = doc.add_paragraph()
    set_run(p.add_run("可更新目录："), size=9, bold=True, color=MUTED)
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "在 Word/WPS 中更新目录")


def add_status_legend(doc: Document):
    doc.add_heading("状态说明", level=1)
    add_table(
        doc,
        ["状态", "含义", "使用要求"],
        [
            ("已确认", "有官方、平台官方、真实文件或实际记录支持", "可作为当前讨论依据；动态事项行动前再核"),
            ("部分成立", "方向有依据，但地区、账号、商品或执行条件未全部核实", "只用于受控试点"),
            ("待验证", "缺少真实资料、当地书面意见或后台测试", "完成验证前不作承诺"),
            ("推测", "根据现有资料形成的工作假设", "用小范围测试检验"),
            ("阻断", "缺少P0条件或存在法律、平台、质量、收款、履约风险", "不得进入对应动作"),
        ],
        [1.05, 2.75, 2.7],
        8.6,
    )


def add_sources(doc: Document, source_ids: Sequence[str] | None = None):
    doc.add_heading("来源与适用范围", level=1)
    wanted = set(source_ids or [source["id"] for source in d.SOURCES])
    rows = []
    for source in d.SOURCES:
        if source["id"] in wanted:
            rows.append(
                [
                    source["id"],
                    (source["title"], source["url"]),
                    source["institution"],
                    source["date"],
                    source["scope"],
                    source["status"],
                    d.ACCESS_DATE,
                ]
            )
    add_table(
        doc,
        ["编号", "官方来源", "发布机构", "页面日期", "适用范围", "状态", "查询日"],
        rows,
        [0.55, 1.75, 1.25, 1.05, 1.65, 0.7, 0.8],
        7.1,
    )


def add_signature(doc: Document, title="双方确认"):
    doc.add_heading(title, level=1)
    add_table(
        doc,
        ["确认方", "公司名称", "负责人", "签字或盖章", "日期"],
        [
            ("线上运营团队", "________________", "________________", "________________", "____年__月__日"),
            ("产品与履约团队", "________________", "________________", "________________", "____年__月__日"),
            ("尼泊尔当地责任主体", "________________", "________________", "________________", "____年__月__日"),
        ],
        [1.35, 1.45, 1.2, 1.45, 1.05],
        8.2,
    )


def save(doc: Document, filename: str, *, evidence=False):
    paragraph_texts = [clean(paragraph.text).strip() for paragraph in doc.paragraphs]
    if "结论" not in paragraph_texts:
        doc.add_heading("结论", level=1)
        add_callout(
            doc,
            "当前结论",
            "本文件的结构与执行口径已完成；凡依赖真实主体、商品、许可、价格、库存、平台后台或当地专业意见的事项，仍以待确认状态管理，未通过前不得进入对应动作。",
            PALE_GREEN,
        )
    if "待确认项" not in paragraph_texts:
        doc.add_heading("待确认项", level=1)
        add_table(
            doc,
            ["事项", "责任方", "完成证据", "当前状态"],
            [
                ("真实主体、商品与履约资料", "产品与履约团队", "有效文件及批准版本", "待提供"),
                ("尼泊尔当地法律与许可适用性", "尼泊尔当地责任主体", "书面意见或官方回执", "待验证"),
                ("平台实际账号与功能能力", "平台与广告账户主体", "后台截图、工单或测试记录", "待平台测试"),
            ],
            [2.25, 1.55, 1.8, 0.9],
            8.4,
        )
    if "下一步" not in paragraph_texts:
        doc.add_heading("下一步", level=1)
        next_step = doc.add_paragraph()
        next_step.paragraph_format.keep_together = True
        next_step.paragraph_format.space_after = Pt(2)
        set_run(
            next_step.add_run(
                "1. 各责任方按P0顺序补齐真实资料，并指定主负责人和备用负责人。\n"
                "2. 完成尼泊尔当地书面复核与实际平台后台测试，把通过版本录入事实源。\n"
                "3. 执行模拟测试；对应门槛全部通过后再做小范围真实验证和放量决定。"
            ),
            size=9.5,
        )
    target_dir = EVIDENCE if evidence else OUT
    target_dir.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = filename.removesuffix(".docx")
    doc.core_properties.subject = "汾酒与海鲜尼泊尔线上销售供应链协同与资料交付"
    doc.core_properties.author = "项目总控与交付工作组"
    doc.core_properties.comments = "合作确认稿；待确认事项不得外推为已完成。"
    target = target_dir / filename
    doc.save(target)
    return target


def role_name(code_list: str) -> str:
    return "、".join(d.ROLES.get(code.strip(), code.strip()) for code in code_list.split(","))


def delivery_rows(items, *, default_scope="逐商品；首批SKU优先", default_channels="全部相关线上渠道"):
    rows = []
    for code, category, name, detail, provider, reviewer, block in items:
        frequency = "首次上线前；变更时立即；每季度复核"
        validity = "以文件有效期或批准版本为准"
        if category == "价格商务":
            frequency = "首次报价前；每次价格变化；每周复核"
            validity = "以报价有效期为准"
        elif category == "库存履约":
            frequency = "实时或高频；每日汇总"
            validity = "超过当日更新时间视为过期"
        elif category == "售后":
            frequency = "上线前；规则变化时；每季度演练"
            validity = "以当前批准版本为准"
        elif category == "合规资料":
            frequency = "首次上线前；到期前60日复核"
        fmt = "填写本DOCX登记表；证书或原件另附PDF/JPG；如已有系统可另附CSV/JSON导出"
        first_due = "收到本清单后7个自然日内；P0最迟在任何公开销售前"
        acceptance = "名称、主体、适用商品、日期、版本一致；责任方签字；可回到原件"
        ai = "通过后可读；原始证件仅限授权人员"
        impact = f"缺失将导致{name}无法进入正式使用"
        change_sla = "一般变更生效前2个工作日；合规、质量、停售和召回立即通知"
        rows.append(
            [
                code,
                f"{category}\n{name}\n{detail}",
                f"提供：{role_name(provider)}\n复核：{role_name(reviewer)}",
                f"{fmt}\n范围：{default_scope}",
                f"首次：{first_due}\n更新：{frequency}\n有效：{validity}",
                acceptance,
                default_channels,
                ai,
                f"{impact}\n等级：{block}",
                f"{change_sla}\n主责：{role_name(provider)}\n备份：{role_name(reviewer)}",
            ]
        )
    return rows


def doc_00_overview():
    sections = ["本套资料解决什么问题", "当前结论", "文件使用顺序", "P0与P1", "项目角色", "执行原则", "待确认与下一步"]
    doc = new_doc("项目总览与使用说明", "双方分工、资料交付、上线准入、AI读取和异常处理的统一入口")
    add_toc(doc, sections)
    doc.add_heading("本套资料解决什么问题", level=1)
    add_callout(
        doc,
        "核心结论",
        "这不是一份普通清单，而是一套“谁提供、谁复核、何时更新、缺什么就停在哪里”的合作执行体系。它把汾酒与海鲜的内容、报价、上架、投流、成交和放量条件分别说明。",
        PALE_GREEN,
    )
    add_table(
        doc,
        ["问题", "本套资料的处理"],
        [
            ("责任说不清", "汾酒与海鲜分别建立RACI，每个动作只有一个最终负责人"),
            ("资料来了但不能用", "每项资料写明格式、适用商品、验收、有效期和缺失影响"),
            ("内容和销售混在一起", "区分内部内容、教育内容、询问、报价、上架、投流、成交和放量"),
            ("AI根据聊天猜信息", "AI只读通过审核的事实源；缺失、过期或冲突时转人工"),
            ("供应变化来不及同步", "价格、库存、配送、许可和召回分级设置SLA"),
            ("账号与复购归属争议", "提供选项、风险和推荐默认方案，保留双方签字决定"),
        ],
        [2.05, 4.45],
        8.7,
    )
    doc.add_heading("当前结论", level=1)
    add_table(
        doc,
        ["事项", "当前状态", "结论"],
        [
            ("责任与资料体系", "已完成", "本套DOCX已提供可填写、可签字、可复核的版本"),
            ("汾酒真实商品与本地主体资料", "待提供", "未补齐前不能正式销售"),
            ("海鲜真实SKU、价格、库存与冷链", "待提供", "未补齐前只做候选和资料准备"),
            ("TikTok酒类普通营销、导流与广告", "阻断", "官方规则仅允许有限地区和合资格企业例外；尼泊尔需当地书面意见与平台后台确认"),
            ("TikTok Shop尼泊尔", "阻断", "官方市场清单未列尼泊尔，不能默认站内上架、支付或联盟成交"),
            ("海鲜B2B/B2C", "部分成立", "公开研究和流程已完成；真实商品与履约条件通过后才能上线"),
            ("AI系统", "设计完成，连接待做", "数据项、同步、权限和回退规则已定义；尚未接入真实系统"),
        ],
        [1.85, 1.2, 3.45],
        8.4,
    )
    doc.add_heading("文件使用顺序", level=1)
    add_steps(
        doc,
        [
            "先读《01_双方职责与项目边界说明书》，确认线上运营团队不承担进口、仓储、配送、质量、退款资金和本地许可。",
            "用《02》《03》分别确认汾酒与海鲜每个动作的RACI。",
            "把《04》作为总资料清单，再用《05》《06》补齐两个品类的专属资料。",
            "用《07》至《12》填写商品、价格、库存、合规、内容事实和AI回复依据。",
            "用《14》判断当前能做到哪一步；没有达到的状态不得越级。",
            "用《16》《17》《22》《23》做SLA、异常、模拟测试和缺失阻断。",
            "最后签署《20》并用《19》记录所有未决选择。",
        ],
    )
    doc.add_heading("P0与P1", level=1)
    add_table(
        doc,
        ["优先级", "含义", "典型资料", "未完成后果"],
        [
            ("P0", "不提供就不能进入对应正式动作", "主体、授权、许可、商品、价格、库存、配送、收款、退款、质量责任", "阻断上架、投流或成交"),
            ("P1", "不提供就难以稳定运营和放量", "素材、案例、详细商务政策、日报周报、备份人员", "只允许小范围测试"),
            ("P2", "优化效率与表现", "更多语言、更多内容、长期分析", "不影响基础合规，但影响转化效率"),
        ],
        [0.8, 1.65, 2.55, 1.5],
        8.5,
    )
    doc.add_heading("项目角色", level=1)
    add_table(doc, ["代号", "角色", "主要责任"], [[code, name, role_summary(code)] for code, name in d.ROLES.items()], [0.7, 2.2, 3.6], 8.3)
    doc.add_heading("执行原则", level=1)
    add_table(
        doc,
        ["原则", "要求"],
        [
            ("真值先行", "任何内容、报价、库存、配送和售后答复只能读取已通过的当前版本"),
            ("一个最终负责人", "RACI每行只有一个A；多人执行不等于多人最终负责"),
            ("缺失即降级", "资料缺失、过期、冲突或无法追溯时，自动退回到更低准入状态"),
            ("现实动作人工批准", "正式发布、报价、投流、收款、退款、召回和大额政策必须人工批准"),
            ("本地责任不外移", "进口、许可、仓储、配送、质量、退款资金和票据由对应本地主体承担"),
            ("版本可追溯", "每次变更保留原值、新值、生效时间、影响范围、操作人和批准人"),
        ],
        [1.35, 5.15],
        8.7,
    )
    add_status_legend(doc)
    doc.add_heading("待确认与下一步", level=1)
    add_table(
        doc,
        ["现在由谁行动", "立即行动"],
        [
            ("线上运营团队", "发出《20_供应链正式资料提供及职责确认函》，安排首次会议并登记双方选择"),
            ("产品与履约团队", "在7个自然日内提交P0主体、授权、商品、价格、库存、配送、售后和质量资料"),
            ("尼泊尔进口与销售主体", "逐商品确认许可、税费、标签、收款、票据、退款和配送合法性"),
            ("平台账户负责人", "用实际账号测试企业授权、广告、LIVE、外链与Shop，不以其他国家结果替代"),
            ("AI技术执行方", "P0数据完成后建立事实源、权限、同步、告警和模拟验收"),
        ],
        [1.85, 4.65],
        8.4,
    )
    add_sources(doc, ["T01", "T02", "T03", "T05", "T06", "N01", "N02", "N03", "F01", "F02", "F04"])
    return save(doc, "00_项目总览与使用说明.docx")


def role_summary(code: str) -> str:
    return {
        "OP": "市场研究、渠道、内容、AI内容、线上询问、CRM和数据复盘；不承担进口与履约",
        "CN": "真实商品、来源、质量、价格、供货、包装和中国出口",
        "BR": "品牌授权、官方事实、视觉、素材权和真伪判断",
        "NI": "尼泊尔进口、海关、食品/酒类许可、标签与召回最终责任",
        "NS": "本地销售、收款、票据、退款资金、商务批准与销售责任",
        "WH": "收货、仓储、温控、库存、拣货、包装、配送协同和售后实物处理",
        "LC": "当地法律、许可、广告、电商、标签和年龄等书面复核",
        "PA": "平台账号、Business Center、广告账户、后台测试、申诉和权限",
        "PY": "商户入网、支付、结算、退款和对账",
        "LG": "运输、冷链、送达、签收和配送异常",
        "CR": "按批准方案制作和披露内容，不得绕过受管制商品限制",
        "AI": "事实源、同步、草稿、提醒、权限、日志和回退；不替代人工批准",
    }[code]


def doc_01_boundaries():
    sections = ["边界结论", "各方负责事项", "线上运营团队不承担事项", "共同事项", "决策与结果责任", "交接要求", "待确认"]
    doc = new_doc("双方职责与项目边界说明书", "把线上运营、产品供应、本地进口、收款、履约、平台和技术责任分开")
    add_toc(doc, sections)
    doc.add_heading("边界结论", level=1)
    add_callout(
        doc,
        "不可越过的边界",
        "线上运营团队负责线上市场与销售系统，不负责进口、报关、许可、法定标签、仓储、冷链、配送、产品质量、退款资金、召回、赔付和票据。上述事项必须由真实承担人书面确认。",
        PALE_RED,
    )
    doc.add_heading("各方负责事项", level=1)
    add_table(
        doc,
        ["责任方", "负责", "必须交付", "最终承担"],
        [
            ("线上运营团队", role_summary("OP"), "市场与渠道方案、内容、咨询和CRM记录、周报、人工审批清单", "线上运营动作的准确性和留痕"),
            ("中国产品供应方", role_summary("CN"), "商品、价格、供货、批次、检测、包装、素材和出口资料", "商品真实性、供货与源头质量"),
            ("品牌授权方", role_summary("BR"), "授权链、官方事实、视觉、素材权、防伪与品牌审核", "品牌权利和官方表述"),
            ("尼泊尔进口主体", role_summary("NI"), "主体、EXIM Code、进口许可、报关、税费、标签与召回文件", "进口与当地监管责任"),
            ("尼泊尔销售与收款主体", role_summary("NS"), "销售许可、收款、票据、价格批准、退款和赔付", "本地销售与资金责任"),
            ("仓储与履约团队", role_summary("WH"), "库存、批次、温控、出库、配送和售后实物记录", "履约与仓库操作"),
            ("尼泊尔合规顾问", role_summary("LC"), "书面意见、适用条文、所需许可、限制和复核日期", "专业意见范围内的复核"),
            ("平台与广告账户主体", role_summary("PA"), "账户资料、权限、后台能力、审核、驳回和申诉记录", "账户与平台操作"),
            ("AI系统与技术执行方", role_summary("AI"), "数据结构、权限、同步、告警、日志和安全回退", "系统按批准规则运行"),
        ],
        [1.45, 2.2, 1.9, 0.95],
        7.8,
    )
    doc.add_heading("线上运营团队不承担事项", level=1)
    add_table(
        doc,
        ["不承担", "实际责任方", "线上运营团队只做什么"],
        [
            ("进口、报关与税费", "尼泊尔进口主体、报关行", "根据已确认结果调整商品页和报价"),
            ("酒类与食品许可", "尼泊尔进口/销售主体、合规顾问", "登记状态并执行阻断"),
            ("法定包装与标签", "供应方、进口主体、合规顾问", "检查页面与实物是否一致"),
            ("产品质量与检测", "供应方、进口主体", "收集证据并转交"),
            ("仓储、冷链与配送", "仓储、物流和销售主体", "展示已批准范围和状态"),
            ("收款、票据与退款资金", "销售主体、支付机构", "引导合法路径并记录进度"),
            ("召回、赔付与销毁", "进口、销售、供应与履约责任方", "暂停内容、通知和CRM协同"),
        ],
        [2.05, 2.1, 2.35],
        8.4,
    )
    doc.add_heading("共同事项", level=1)
    add_table(
        doc,
        ["事项", "线上运营团队", "产品与履约团队", "共同验收"],
        [
            ("商品上架", "制作商品页、检查事实、记录版本", "提供并批准商品、许可、价格、库存、配送和售后", "准入表全部P0通过"),
            ("正式报价", "整理需求、生成草稿、发送批准版本", "确认价格、库存、税费、配送和有效期", "报价可回到批准记录"),
            ("内容发布", "制作、审校、披露和发布", "确认事实、素材权、库存与可售范围", "事实、法律与平台均通过"),
            ("异常处理", "停止自动回复、收集证据、通知相关方", "隔离、调查、退款、补发、召回或赔付", "处理结果与复盘归档"),
            ("放量", "复盘流量、询问、成交与复购", "确认供应、仓配、质量和资金能力", "达到《14》放量门槛"),
        ],
        [1.2, 1.85, 2.15, 1.3],
        8.2,
    )
    doc.add_heading("决策与结果责任", level=1)
    add_table(
        doc,
        ["层级", "谁决定", "谁执行", "谁对结果负责"],
        [
            ("市场与内容方向", "线上运营团队项目负责人", "内容、渠道与技术人员", "线上运营团队"),
            ("商品真实性与供货", "中国产品供应方负责人", "供应与质量人员", "中国产品供应方"),
            ("当地进口与许可", "尼泊尔进口主体负责人", "报关与合规人员", "尼泊尔进口主体"),
            ("销售、价格、收款与退款", "尼泊尔销售主体负责人", "商务、财务和售后人员", "尼泊尔销售主体"),
            ("仓库与配送", "仓储与履约负责人", "仓库和物流人员", "仓储与履约团队"),
            ("账号与平台操作", "账户主体负责人", "平台运营人员", "平台与广告账户主体"),
            ("AI自动化", "线上运营团队与技术负责人共同批准", "技术执行方", "规则批准方承担业务决定，技术方承担系统执行"),
        ],
        [1.6, 1.9, 1.45, 1.55],
        8.2,
    )
    doc.add_heading("交接要求", level=1)
    add_steps(
        doc,
        [
            "同一家公司承担多个角色时，也按角色分别签名，不把进口、销售、仓储和平台责任合并。",
            "主负责人离开或无法联系时，备用负责人应在SLA内接管。",
            "任何口头确认必须在1个工作日内写入对应DOCX并由责任方确认。",
            "合作终止时按《18》约定导出账号、内容、CRM、数据库和审批记录。",
        ],
    )
    doc.add_heading("待确认", level=1)
    add_table(
        doc,
        ["事项", "当前状态", "需要谁确认", "不确认的影响"],
        [
            ("各角色由哪家公司或人员承担", "待确认", "双方及尼泊尔当地团队", "RACI无法落到真实人员"),
            ("同一主体是否兼任进口、销售、收款和仓储", "待确认", "尼泊尔当地团队、合规顾问", "授权链和许可范围不清"),
            ("最终签字人与备用负责人", "待确认", "各责任方", "紧急事件无人接管"),
            ("90天试点预算与停止线", "待确认", "双方决策人", "无法启动现实动作"),
        ],
        [2.25, 1.1, 1.7, 1.45],
        8.4,
    )
    add_signature(doc)
    return save(doc, "01_双方职责与项目边界说明书.docx")


def doc_raci(filename: str, title: str, rows, product: str):
    sections = ["RACI说明", "角色对照", "责任矩阵", "阻断升级", "签字确认"]
    doc = new_doc(title, f"{product}从资料、内容、销售到履约和复盘的唯一最终负责人", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("RACI说明", level=1)
    add_table(
        doc,
        ["字母", "含义", "规则"],
        [
            ("R", "实际执行", "可多人"),
            ("A", "最终负责与批准", "每行只有一个"),
            ("C", "动作前必须协商", "未协商不得开始"),
            ("I", "动作后需要通知", "按SLA留痕"),
        ],
        [0.8, 2.2, 3.5],
        8.5,
    )
    doc.add_heading("角色对照", level=1)
    add_table(doc, ["代号", "角色"], [[code, name] for code, name in d.ROLES.items()], [0.8, 5.7], 8.3)
    doc.add_heading("责任矩阵", level=1)
    matrix = []
    for row in rows:
        code, action, r, a, c, i, pre, deliverable, sla, acceptance, block, escalate = row
        matrix.append(
            [
                code,
                action,
                r,
                a,
                c,
                i,
                pre,
                deliverable,
                sla,
                acceptance,
                block,
                escalate,
            ]
        )
    add_table(
        doc,
        ["编号", "动作", "R", "A", "C", "I", "前置条件", "交付物", "SLA", "验收", "阻断", "异常升级"],
        matrix,
        [0.45, 1.15, 0.65, 0.5, 0.75, 0.65, 1.25, 1.1, 0.85, 1.3, 0.45, 0.85],
        6.5,
    )
    doc.add_heading("阻断升级", level=1)
    add_table(
        doc,
        ["等级", "处理"],
        [
            ("P0", "停止对应发布、报价、上架、投流、收款或履约；立即通知A与升级对象"),
            ("P1", "允许受控测试，但不得放量；在下一个工作日内完成修正"),
            ("P2", "记录并进入周度优化，不影响已通过的基础动作"),
        ],
        [1, 5.5],
        8.5,
    )
    add_signature(doc, "RACI签字确认")
    return save(doc, filename)


def doc_04_master_delivery():
    sections = ["填写说明", "资料交付总表", "验收与退回", "P0立即阻断", "签字确认"]
    doc = new_doc("供应链资料交付总清单", "每项资料都写明提供、复核、格式、频率、有效期、AI读取和缺失影响", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("填写说明", level=1)
    add_callout(
        doc,
        "交付形式",
        "本轮最终文件全部为DOCX。请直接在本表填写；证书、许可、检测、授权和原件照片作为附件提供。若已有系统，可另附CSV或JSON导出，但不能只发聊天截图。",
        PALE_GREEN,
    )
    doc.add_heading("资料交付总表", level=1)
    add_table(
        doc,
        ["编号", "资料与具体内容", "提供/复核", "形式与范围", "首次/更新/有效", "验收", "渠道", "AI读取", "缺失影响/阻断", "变更SLA/主备"],
        delivery_rows(d.MASTER_DELIVERY_ITEMS),
        [0.45, 1.8, 1.15, 1.4, 1.35, 1.25, 0.9, 0.8, 1.25, 1.2],
        6.4,
    )
    doc.add_heading("验收与退回", level=1)
    add_table(
        doc,
        ["退回原因", "处理"],
        [
            ("看不出适用商品", "补充SKU编号和包装版本"),
            ("只有扫描件，没有登记信息", "补充文件编号、日期、有效期、主体和保管人"),
            ("与实物或其他资料冲突", "暂停使用，责任方书面说明并提交新版本"),
            ("已过期或无续期安排", "自动降级为阻断"),
            ("只在聊天中口头确认", "在本DOCX登记并由责任方确认后再使用"),
            ("价格、库存或配送没有更新时间", "视为过期，AI不得读取"),
        ],
        [2.5, 4.0],
        8.3,
    )
    doc.add_heading("P0立即阻断", level=1)
    add_callout(
        doc,
        "阻断规则",
        "无合法主体、无授权、来源不明、许可不明、标签未确认、实物不一致、无价格、无库存、无配送、无售后责任、收款主体不明、酒类平台资格未确认、海鲜食品状态不明或产品与履约团队拒绝承担质量责任时，不得正式成交。",
        PALE_RED,
    )
    add_signature(doc)
    return save(doc, "04_供应链资料交付总清单.docx")


def special_delivery_rows(items, product, channels):
    rows = []
    for code, name, provider, reviewer, block in items:
        fmt = "本DOCX填写；证明文件另附PDF/JPG；后台能力另附带日期截图"
        scope = "逐商品/逐账号/逐渠道"
        frequency = "上线前；变化立即；每月复核"
        validity = "以文件、后台或批准版本为准"
        first_due = "收到清单后7日内；P0在任何公开动作前"
        acceptance = "责任方签字；资料、实物或后台一致；可追溯"
        ai = "通过后可读；账号凭据与证件原件不向AI开放"
        missing = f"缺少“{name}”时，{product}对应动作保持阻断"
        rows.append(
            [
                code,
                name,
                f"提供：{role_name(provider)}\n复核：{role_name(reviewer)}",
                f"{fmt}\n范围：{scope}",
                f"首次：{first_due}\n更新：{frequency}\n有效：{validity}",
                acceptance,
                channels,
                ai,
                f"{missing}\n等级：{block}",
                f"一般变更提前2个工作日；合规、质量、停售立即\n主责：{role_name(provider)}\n备份：{role_name(reviewer)}",
            ]
        )
    return rows


def doc_special(filename, title, subtitle, items, product, channels, source_ids):
    sections = ["当前边界", "专属资料清单", "平台或当地验证动作", "阻断规则", "来源", "签字确认"]
    doc = new_doc(title, subtitle, landscape=True)
    add_toc(doc, sections)
    doc.add_heading("当前边界", level=1)
    if product == "汾酒":
        add_callout(
            doc,
            "当前判断",
            "TikTok普通上传、酒类营销、付费广告、直播、外链和站内成交必须分开。尼泊尔公开酒类广告受法律强限制；TikTok对酒类交易与营销只提供有限地区、合资格企业例外。未拿到当地书面意见和实际后台通过前，不做购买导流、广告或直播成交。",
            PALE_RED,
        )
    else:
        add_callout(
            doc,
            "当前判断",
            "海鲜B2B与B2C可以继续做资料和小范围验证，但真实商品、进口许可、标签、价格、库存、冷链、配送和售后任何一项P0缺失，都不能正式上架或收款。",
            PALE_RED,
        )
    doc.add_heading("专属资料清单", level=1)
    add_table(
        doc,
        ["编号", "资料/验证项", "提供/复核", "形式与范围", "首次/更新/有效", "验收", "渠道", "AI读取", "缺失影响/阻断", "变更SLA/主备"],
        special_delivery_rows(items, product, channels),
        [0.45, 1.45, 1.15, 1.4, 1.35, 1.2, 0.95, 0.75, 1.25, 1.15],
        6.45,
    )
    doc.add_heading("平台或当地验证动作", level=1)
    actions = (
        [
            "由尼泊尔合规顾问逐项出具：教育内容、品牌营销、付费广告、达人、LIVE、外链、线上销售、配送到家、赠送与抽奖的书面意见。",
            "由平台账户负责人用实际账号测试企业授权、酒类行业资格、自助广告、尼泊尔地域、外链、LIVE、Shop和站内支付。",
            "后台记录必须含账号、主体、日期、提交资料、状态、驳回原因、申诉和结果。",
            "其他国家可用不作为尼泊尔可用证据。",
        ]
        if product == "汾酒"
        else [
            "由尼泊尔进口主体和报关行逐商品确认海关编码、税费、DFTQC路径、单证和标签。",
            "由仓储与物流团队完成真实温控、拣货、配送、拒收、退款和召回演练。",
            "由当地语言人员复核英文和尼泊尔语名称、份量、保存、烹饪与过敏提示。",
            "每个销售渠道分别测试入驻、类目、结算、退货、物流和内容要求。",
        ]
    )
    add_steps(doc, actions)
    doc.add_heading("阻断规则", level=1)
    add_table(
        doc,
        ["情形", "允许", "禁止"],
        [
            ("资料不完整", "内部整理和补资料", "公开发布、报价、上架、成交"),
            ("素材完整但无现货", "制作内部样片", "写现货、时效或购买引导"),
            ("有库存但配送未定", "接收内部需求", "接受正式订单"),
            ("授权有但平台资格未定", "合规教育内容草稿", "广告、商品挂载、直播成交"),
            ("可测试但流程未稳定", "小范围人工测试", "加预算或扩大区域"),
        ],
        [2, 2.15, 2.35],
        8.2,
    )
    add_sources(doc, source_ids)
    add_signature(doc)
    return save(doc, filename)


def doc_07_product_template():
    sections = ["填写说明", "商品主数据模板", "版本与批准", "缺失处理", "签字确认"]
    doc = new_doc("产品主数据模板", "一行一个数据项，逐SKU填写，不允许用聊天记录代替", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("填写说明", level=1)
    add_callout(
        doc,
        "用法",
        "复制本表为每个SKU建立一份商品资料。酒类与海鲜共用基础部分，品类不适用的数据项写“不适用”，不能留空或凭估算补齐。",
        PALE_GREEN,
    )
    doc.add_heading("商品主数据模板", level=1)
    rows = []
    for group, item, instruction, example, required, frequency in d.PRODUCT_DATA_ITEMS:
        rows.append(
            [
                group,
                item,
                instruction,
                example,
                required,
                "中国产品供应方；当地名称由尼泊尔销售主体",
                "线上运营团队、尼泊尔进口主体",
                frequency,
                "审核通过后可读",
                "缺失则阻断相关内容、报价、上架或履约",
                "________________",
            ]
        )
    add_table(
        doc,
        ["分组", "数据项", "填写说明", "示例", "必填", "提供", "复核", "更新", "AI读取", "缺失影响", "填写值"],
        rows,
        [0.6, 1.15, 1.65, 1.2, 0.5, 1.2, 1.2, 0.75, 0.8, 1.25, 1.2],
        6.7,
    )
    doc.add_heading("版本与批准", level=1)
    add_table(
        doc,
        ["SKU编号", "当前版本", "生效时间", "提供负责人", "复核负责人", "最终批准", "状态"],
        [["________________", "v____", "____年__月__日 __:__", "________________", "________________", "________________", "待提供"]],
        [1.1, 0.8, 1.4, 1.2, 1.2, 1.2, 0.9],
        8.2,
    )
    doc.add_heading("缺失处理", level=1)
    add_table(
        doc,
        ["缺失类型", "系统处理"],
        [
            ("名称、规格、净重、酒精度或物种缺失", "不得制作销售内容或上架"),
            ("海关编码、许可或标签缺失", "不得进口或销售"),
            ("价格、库存或配送缺失", "不得报价、收款或承诺时效"),
            ("批次、保质期或温控缺失", "不得出库"),
            ("图片与实物不一致", "暂停内容和商品页"),
        ],
        [2.8, 3.7],
        8.4,
    )
    add_signature(doc, "商品资料确认")
    return save(doc, "07_产品主数据模板.docx")


def doc_08_price_template():
    sections = ["填写说明", "价格与商务政策", "变更记录", "报价批准", "签字确认"]
    doc = new_doc("价格与商务政策模板", "供货、B2B、B2C、样品、配送和合作佣金的批准版本", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("填写说明", level=1)
    add_callout(
        doc,
        "价格原则",
        "任何价格都必须同时写明币种、计价单位、税费、配送、有效期和批准人。没有有效期的价格视为不可用；AI不得从聊天记录引用临时报价。",
        PALE_RED,
    )
    doc.add_heading("价格与商务政策", level=1)
    rows = []
    for group, item, detail, provider, reviewer, block in d.PRICE_ITEMS:
        rows.append(
            [
                group,
                item,
                detail,
                role_name(provider),
                role_name(reviewer),
                "________________",
                "NPR / CNY / USD",
                "________________",
                "____年__月__日至____年__月__日",
                "待提供",
                block,
                "缺失或过期则停止正式报价",
            ]
        )
    add_table(
        doc,
        ["分组", "价格项", "说明", "提供", "复核", "SKU/范围", "币种", "金额/比例", "有效期", "状态", "阻断", "缺失影响"],
        rows,
        [0.6, 1.0, 1.55, 1.05, 1.05, 0.95, 0.65, 0.85, 1.2, 0.75, 0.5, 1.25],
        6.7,
    )
    doc.add_heading("变更记录", level=1)
    add_table(
        doc,
        ["变更时间", "SKU/范围", "原值", "新值", "生效时间", "影响渠道", "需要下架/改内容", "操作人", "批准人"],
        [["________________", "________________", "________________", "________________", "________________", "________________", "是 / 否", "________________", "________________"]],
        [1.1, 1, 0.95, 0.95, 1.1, 1.1, 1.35, 0.9, 0.9],
        7.7,
    )
    doc.add_heading("报价批准", level=1)
    add_table(
        doc,
        ["报价对象", "SKU与数量", "价格版本", "库存确认时间", "配送确认", "批准人", "报价有效期"],
        [["________________", "________________", "v____", "____年__月__日 __:__", "已确认 / 未确认", "________________", "________________"]],
        [1.2, 1.4, 0.8, 1.5, 1.2, 1.05, 1.1],
        8.0,
    )
    add_signature(doc)
    return save(doc, "08_价格与商务政策模板.docx")


def doc_09_inventory_template():
    sections = ["填写说明", "库存配送订单状态", "同步规则", "异常回退", "签字确认"]
    doc = new_doc("库存_配送_订单状态数据模板", "把在库、可售、锁定、在途、配送、订单和退款分开记录", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("填写说明", level=1)
    add_callout(
        doc,
        "实时性",
        "“实物库存”不等于“可售库存”。可售量必须扣除锁定、预留、临期、隔离和异常批次。库存和订单状态没有更新时间时，AI不得作销售承诺。",
        PALE_RED,
    )
    doc.add_heading("库存配送订单状态", level=1)
    rows = []
    for group, item, detail, provider, frequency in d.INVENTORY_ITEMS:
        rows.append(
            [
                group,
                item,
                detail,
                role_name(provider),
                frequency,
                "________________",
                "________________",
                "待提供",
                "是（通过后）",
                "缺失则转人工并关闭对应自动承诺",
                "________________",
            ]
        )
    add_table(
        doc,
        ["分组", "数据项", "填写说明", "更新方", "频率", "SKU/订单", "当前值", "状态", "AI读取", "缺失影响", "更新时间/人员"],
        rows,
        [0.55, 1.15, 1.65, 1.15, 0.85, 1.05, 1.0, 0.75, 0.75, 1.4, 1.25],
        6.7,
    )
    doc.add_heading("同步规则", level=1)
    add_table(
        doc,
        ["优先级", "同步内容", "目标频率", "超时处理"],
        [
            ("P0", "可售库存、缺货、紧急价格、配送异常、订单、停售、许可失效、召回", "实时或15分钟内", "立即停止自动承诺"),
            ("P1", "入库、出库、在途、临期、配送能力、售后、退款", "每日/节点后15分钟", "通知负责人并转人工"),
            ("P2", "商品资料、活动、达人政策、素材、FAQ", "每周或变更时", "保留旧版本并标过期"),
        ],
        [0.8, 3.35, 1.25, 1.1],
        8.2,
    )
    doc.add_heading("异常回退", level=1)
    add_table(
        doc,
        ["异常", "系统动作", "人工动作"],
        [
            ("库存冲突", "采用更保守值并告警", "仓库15分钟内确认"),
            ("价格过期", "停止正式报价", "销售主体重新批准"),
            ("配送状态停滞", "显示处理中，不承诺到达", "物流与仓库调查"),
            ("支付失败", "不出库、不扣减最终库存", "财务对账"),
            ("批次隔离或召回", "关闭相关SKU和内容", "进口与质量负责人处置"),
        ],
        [1.6, 2.45, 2.45],
        8.4,
    )
    add_signature(doc)
    return save(doc, "09_库存_配送_订单状态数据模板.docx")


def doc_10_compliance_index():
    sections = ["填写说明", "合规与授权索引", "公开与访问权限", "到期提醒", "签字确认"]
    doc = new_doc("合规资质与授权文件索引", "每份文件关联具体商品、主体、国家、渠道、有效期和原件保管人", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("填写说明", level=1)
    add_callout(
        doc,
        "证据规则",
        "证书扫描件可以作为附件，但必须在本表登记。没有适用商品、主体、日期、有效期和原件保管人的文件，不进入可用事实源。",
        PALE_RED,
    )
    doc.add_heading("合规与授权索引", level=1)
    rows = []
    for group, name, scope, provider, reviewer in d.COMPLIANCE_ITEMS:
        rows.append(
            [
                group,
                name,
                scope,
                role_name(provider),
                role_name(reviewer),
                "________________",
                "________________",
                "____年__月__日",
                "____年__月__日",
                "尼泊尔/相关国家",
                "公开 / 仅平台 / 仅内部",
                "________________",
                "待提供",
            ]
        )
    add_table(
        doc,
        ["类别", "文件", "适用商品/范围", "提供", "复核", "文件编号", "机构", "出具日", "到期日", "国家", "可见范围", "路径/保管人", "状态"],
        rows,
        [0.6, 1.2, 1.25, 0.9, 0.9, 0.9, 1.1, 0.8, 0.8, 0.75, 1.05, 1.15, 0.7],
        6.25,
    )
    doc.add_heading("公开与访问权限", level=1)
    add_table(
        doc,
        ["资料类型", "商品页可见", "可给采购单位", "可给平台", "AI读取", "说明"],
        [
            ("主体公开登记", "摘要", "按需要", "是", "摘要", "隐藏证件号码和非必要个人资料"),
            ("品牌授权", "授权身份摘要", "按需要", "是", "摘要", "完整文件限授权人员"),
            ("许可与检测", "编号/摘要", "按业务需要", "是", "结构化摘要", "原件不直接开放"),
            ("价格与库存", "按渠道", "按权限", "按需要", "当前版本", "不得跨渠道泄露底价"),
            ("账号凭据", "否", "否", "否", "否", "只在密码管理系统保存"),
        ],
        [1.6, 1.0, 1.2, 1.0, 1.0, 1.7],
        8.1,
    )
    doc.add_heading("到期提醒", level=1)
    add_table(
        doc,
        ["提醒时点", "动作"],
        [
            ("到期前90日", "确认是否续期、责任人和所需资料"),
            ("到期前60日", "发出正式提醒；如无计划，标记高风险"),
            ("到期前30日", "停止新增长期承诺和跨期投放"),
            ("到期前7日", "准备下架、停投和通知"),
            ("到期日", "自动降级为阻断，直到新文件通过"),
        ],
        [1.4, 5.1],
        8.5,
    )
    add_signature(doc)
    return save(doc, "10_合规资质与授权文件索引.docx")


def doc_11_content_facts():
    sections = ["内容事实库", "禁止表达库", "AI画面与实拍", "发布审批", "来源"]
    doc = new_doc("内容事实库与禁止表达库", "每个宣传点有来源；没有依据的说法不进入脚本、字幕、画面或回复", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("内容事实库", level=1)
    add_table(
        doc,
        ["类别", "允许使用的事实", "证明", "最终确认", "优先级", "具体内容", "来源编号/路径", "版本", "状态"],
        [[a, b, c, role_name(owner), priority, "________________", "________________", "v____", "待提供"] for a, b, c, owner, priority in d.CONTENT_FACTS],
        [0.8, 1.6, 1.4, 1.05, 0.65, 1.7, 1.35, 0.65, 0.75],
        6.8,
    )
    doc.add_heading("禁止表达库", level=1)
    add_table(
        doc,
        ["风险表达", "不能这样写", "建议处理", "适用范围", "批准例外"],
        [[a, b, c, "内容、广告、达人、直播、商品页、AI回复", "无；如法律或平台有变化须重新审"] for a, b, c in d.FORBIDDEN_EXPRESSIONS],
        [1.35, 2.05, 1.9, 1.4, 1.1],
        7.6,
    )
    doc.add_heading("AI画面与实拍", level=1)
    add_table(
        doc,
        ["类型", "要求"],
        [
            ("真实商品", "瓶体、包装、规格、标签、颜色和数量必须与当前商品一致"),
            ("仓库或门店", "优先实拍；AI示意不得伪装成真实地点"),
            ("人物", "有授权；酒类画面使用合规成年人；不表现过量饮酒"),
            ("食品状态", "不得用AI画面替代解冻前后、裹冰、净重和送达状态证据"),
            ("标记", "涉及真实感较强的AI生成或编辑时，按平台要求标记"),
        ],
        [1.45, 5.05],
        8.5,
    )
    doc.add_heading("发布审批", level=1)
    add_steps(doc, ["脚本引用事实库版本", "供应与品牌事实审核", "尼泊尔当地法律审核", "平台规则与披露检查", "库存、价格和配送再核", "批准后发布并记录URL", "变化时下架或更新"])
    add_sources(doc, ["T01", "T02", "T03", "T07", "N01", "F01"])
    return save(doc, "11_内容事实库与禁止表达库.docx")


def doc_12_faq():
    sections = ["AI答复原则", "FAQ与人工转接", "强制人工场景", "会话记录", "签字确认"]
    doc = new_doc("AI客服FAQ与人工转接规则", "AI只回答已批准事实；价格、质量、退款、酒类购买和异常必须按规则转人工", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("AI答复原则", level=1)
    add_callout(
        doc,
        "默认行为",
        "资料缺失、过期、冲突或不在批准范围时，AI只说明“需要人工确认”，并收集最少必要信息。不得自行承诺价格、库存、配送、质量、退款、许可或平台能力。",
        PALE_RED,
    )
    doc.add_heading("FAQ与人工转接", level=1)
    add_table(
        doc,
        ["主题", "AI权限", "答复依据", "转人工条件", "人工负责人", "标准答复要点", "记录要求"],
        [[a, b, c, e, role_name(owner), "仅使用当前批准版本；不确定时说明待确认", "时间、问题、引用版本、答复、接管人"] for a, b, c, e, owner in d.FAQ_ITEMS],
        [1.0, 0.85, 1.55, 1.55, 1.0, 1.65, 1.4],
        6.8,
    )
    doc.add_heading("强制人工场景", level=1)
    add_table(
        doc,
        ["场景", "最长接管时间", "自动动作"],
        [
            ("正式报价、账期、独家、返点、大额订单", "15分钟确认，4小时给结果", "创建审批任务"),
            ("酒类购买、外链、验龄、配送与退款", "15分钟", "停止自动销售引导"),
            ("食品安全、过敏、融化、变质、批次", "立即", "提示停止食用、保留证据"),
            ("疑似假货、重大舆情、监管或平台通知", "立即", "冻结相关自动回复"),
            ("支付失败、重复扣款、退款失败", "15分钟", "停止出库并创建财务任务"),
        ],
        [2.9, 1.35, 2.25],
        8.4,
    )
    doc.add_heading("会话记录", level=1)
    add_table(
        doc,
        ["会话编号", "渠道", "咨询方类型", "SKU/订单", "问题", "引用版本", "AI答复", "接管人", "结果"],
        [["________________", "________________", "采购单位 / 购买者", "________________", "________________", "v____", "________________", "________________", "________________"]],
        [1.0, 0.9, 1.15, 1.0, 1.4, 0.8, 1.5, 1.0, 1.0],
        7.6,
    )
    add_signature(doc)
    return save(doc, "12_AI客服FAQ与人工转接规则.docx")


def doc_13_ai_interface():
    sections = ["设计目标", "十二张事实表", "同步优先级", "Excel起步", "数据库过渡", "API成熟", "JSON示例", "权限与历史", "验收"]
    doc = new_doc("AI数据接口与同步规范", "让AI只读产品与履约团队确认的事实，不从聊天记录猜商品状态")
    add_toc(doc, sections)
    doc.add_heading("设计目标", level=1)
    add_callout(
        doc,
        "唯一事实源",
        "商品、价格、库存、配送、订单、售后、合规、内容和禁止表达必须各有唯一当前版本。聊天、邮件和会议只能触发更新，不能直接成为AI答复依据。",
        PALE_GREEN,
    )
    doc.add_heading("十二张事实表", level=1)
    tables = [
        ("产品主表", "sku_id", "名称、规格、净重、批次规则、储存、有效期、销售状态", "CN", "变更时", "缺失则不上架"),
        ("价格表", "price_id", "sku_id、渠道、币种、金额、税、配送、有效期、批准人", "NS", "变化/每周", "过期则不报价"),
        ("库存表", "inventory_id", "sku_id、仓库、可售、锁定、预留、在途、临期、更新时间", "WH", "实时/高频", "不同步则关闭承诺"),
        ("配送规则表", "delivery_rule_id", "区域、费用、截单、时效、温控、上限、例外", "WH", "每日/变化", "超区转人工"),
        ("订单状态表", "order_id", "商品、数量、付款、库存锁定、出库、配送、签收", "NS", "实时", "支付失败不出库"),
        ("售后规则表", "after_sale_rule_id", "场景、证据、时限、退款、补发、赔付、升级", "NS", "季度/变化", "未批准不成交"),
        ("合规文件索引", "document_id", "主体、编号、机构、日期、有效期、适用商品、访问权限", "LC", "到期/变化", "失效则降级"),
        ("内容事实库", "fact_id", "批准表述、来源、适用商品、渠道、版本、批准人", "BR/CN", "变化时", "无来源不发布"),
        ("禁止表达库", "restriction_id", "风险表达、适用范围、原因、替代处理", "LC/OP", "政策变化", "命中则阻断"),
        ("FAQ知识库", "faq_id", "问题、批准答复、引用版本、接管条件", "OP", "每周/变化", "未知则转人工"),
        ("真实案例库", "case_id", "合作类型、商品、结果、证据、公开授权", "CN/NS", "发生时", "无授权不公开"),
        ("变更记录表", "change_id", "原值、新值、生效、影响、操作人、批准人、通知", "AI", "每次变化", "记录不可删除"),
    ]
    add_table(doc, ["表名", "关键编号", "主要数据项", "责任", "更新", "失效处理"], tables, [1.35, 1.2, 2.35, 0.75, 0.95, 1.15], 7.5)
    doc.add_heading("同步优先级", level=1)
    add_table(
        doc,
        ["优先级", "内容", "频率", "失败回退"],
        [
            ("P0", "库存、缺货、紧急价格、配送异常、订单、停售、许可失效、召回", "实时或15分钟内", "停止自动承诺并转人工"),
            ("P1", "入库、出库、在途、临期、配送能力、售后、退款", "每日/节点后", "标记延迟并通知"),
            ("P2", "商品资料、活动、达人政策、合规文件、素材、FAQ", "每周/变更时", "保留旧版并标过期"),
        ],
        [0.85, 3.35, 1.2, 1.1],
        8.2,
    )
    doc.add_heading("Excel起步", level=1)
    add_table(
        doc,
        ["适用条件", "做法", "成本", "风险"],
        [
            ("SKU少、单仓、人工更新", "以本套DOCX作为合作确认源；由技术人员转入受控表格；价格和库存设置版本与锁定", "低", "多人同时改、权限弱、同步延迟"),
        ],
        [1.55, 3.0, 0.7, 1.25],
        8.4,
    )
    doc.add_heading("数据库过渡", level=1)
    add_table(
        doc,
        ["适用条件", "做法", "成本", "风险"],
        [
            ("多SKU、多渠道、多仓、每日订单", "建立商品、价格、库存、订单、文件和审批关系；按角色授权；保留历史", "中", "需要数据迁移、接口和运维"),
        ],
        [1.55, 3.0, 0.7, 1.25],
        8.4,
    )
    doc.add_heading("API成熟", level=1)
    add_table(
        doc,
        ["适用条件", "做法", "成本", "风险"],
        [
            ("仓库、收款、配送和平台均有稳定接口", "事件驱动同步；签名校验；幂等更新；监控延迟；自动回退", "高", "外部接口变化、权限和故障传播"),
        ],
        [1.55, 3.0, 0.7, 1.25],
        8.4,
    )
    doc.add_heading("JSON示例", level=1)
    add_callout(
        doc,
        "示例",
        '{ "sku_id": "SF-SHR-001", "data_version": "v3", "price_npr": 0, "price_status": "待提供", "sellable_stock": 0, "inventory_status": "待提供", "delivery_status": "待提供", "approved_by": "", "updated_at": "2026-07-27T18:00:00+05:45" }',
        PALE,
    )
    add_para(doc, "示例中的0不是可销售价格或库存；状态仍为“待提供”，系统必须阻断报价和成交。")
    doc.add_heading("权限与历史", level=1)
    add_table(
        doc,
        ["角色", "可读", "可改", "可批准"],
        [
            ("供应负责人", "商品与供货", "商品草稿", "商品真实性"),
            ("仓库负责人", "库存与订单", "库存、批次、配送状态", "出库与隔离"),
            ("销售主体", "价格、订单、售后", "价格与规则草稿", "价格、退款、特殊政策"),
            ("线上运营团队", "批准后的事实", "内容、FAQ和CRM草稿", "线上发布与一般回复"),
            ("AI系统", "授权范围", "只写建议、日志和状态", "无现实动作批准权"),
            ("系统管理员", "系统配置", "权限与接口", "不替代业务批准"),
        ],
        [1.35, 1.7, 2.0, 1.45],
        8.3,
    )
    doc.add_heading("验收", level=1)
    add_steps(doc, ["价格过期自动停止报价", "库存为零自动关闭购买", "合规文件到期自动降级", "未知问题自动转人工", "历史版本可回看", "未授权人员不能改价、发布或退款", "24项模拟测试全部留痕"])
    return save(doc, "13_AI数据接口与同步规范.docx")


def doc_14_launch_gate():
    sections = ["准入原则", "状态与条件", "汾酒特别闸门", "海鲜特别闸门", "P0阻断", "当前判定", "签字确认"]
    doc = new_doc("上线准入与产品状态检查表", "每个商品和渠道从资料收集到放量的逐级准入", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("准入原则", level=1)
    add_callout(doc, "禁止越级", "可发布教育内容不等于可广告；可报价不等于可收款；可上架不等于可投流；可成交不等于可放量。", PALE_RED)
    doc.add_heading("状态与条件", level=1)
    add_table(
        doc,
        ["状态", "最低条件", "允许", "禁止", "商品/渠道", "责任人", "证据", "当前选择"],
        [[state, condition, allowed, forbidden, "________________", "________________", "________________", "□"] for state, condition, allowed, forbidden in d.LAUNCH_STATES],
        [1.15, 2.15, 1.35, 1.7, 1.1, 1.0, 1.3, 0.55],
        7.0,
    )
    doc.add_heading("汾酒特别闸门", level=1)
    add_table(
        doc,
        ["能力", "官方与当前判断", "状态", "解除条件"],
        [
            ("普通教育内容", "可制作；发布前仍需品牌事实、当地法律和平台资格复核", "待验证", "书面审核与实际账号通过"),
            ("酒类营销与购买导流", "TikTok默认限制受管制商品营销与导流；只允许有限地区和合资格企业例外", "阻断", "平台企业授权＋当地书面意见"),
            ("付费广告", "TikTok酒类广告要求批准市场、许可与年龄限制；尼泊尔自助开户清单未列", "阻断", "平台代表或后台书面通过＋当地法律允许"),
            ("LIVE", "一般LIVE需18+且商业内容要披露；酒类交易和导流不能据此推定", "阻断成交", "账号、内容与当地意见分别通过"),
            ("Shop/商品挂载/站内支付", "官方市场清单未列尼泊尔", "阻断", "官方市场清单或实际后台变化并重新审核"),
            ("站外销售", "TikTok将外链、联系信息和引导站外互动计入交易/导流；酒类需例外", "阻断", "平台与当地两条线均书面通过"),
        ],
        [1.45, 3.55, 1.0, 2.0],
        7.6,
    )
    doc.add_heading("海鲜特别闸门", level=1)
    add_table(
        doc,
        ["阶段", "必须有", "未完成处理"],
        [
            ("对外展示", "真实商品、图片、规格、净重、储存和可送范围", "只做通用教育"),
            ("正式报价", "批准价格、库存、最低订货量、税费、配送和有效期", "不得发送"),
            ("上架收款", "进口、标签、许可、支付、库存、配送和售后", "不得上线"),
            ("真实配送", "温控、包装、路线、签收、拒收和退款演练", "仅内部测试"),
            ("放量", "稳定补货、复购、毛利、投诉和召回机制", "维持小范围"),
        ],
        [1.35, 3.35, 1.8],
        8.1,
    )
    doc.add_heading("P0阻断", level=1)
    add_callout(
        doc,
        "直接阻断",
        "无合法主体、无授权、来源不明、合规不明、标签未确认、产品与资料不一致、无价格、无库存、无配送、无售后责任、无退款规则、收款主体不明、酒类平台资格未确认、海鲜食品状态不明或拒绝承担质量责任。",
        PALE_RED,
    )
    doc.add_heading("当前判定", level=1)
    add_table(
        doc,
        ["项目", "当前最高状态", "原因"],
        [
            ("汾酒TikTok", "可以制作内部内容", "真实SKU、授权、当地广告法律意见和平台受管制行业资格尚未完整提供"),
            ("海鲜B2B/B2C", "资料不完整 / 可以制作内部内容", "真实SKU、逐商品合规、价格、库存、冷链、收款和售后仍待提供"),
            ("AI销售系统", "内部测试", "结构已设计，尚未接入真实数据与账号"),
        ],
        [1.65, 2.1, 2.75],
        8.4,
    )
    add_sources(doc, ["T01", "T02", "T03", "T04", "T05", "T06", "T08", "N01", "N02", "N03", "F01", "F02", "F04"])
    add_signature(doc, "上线状态批准")
    return save(doc, "14_上线准入与产品状态检查表.docx")


def doc_15_process():
    sections = ["流程总图", "逐步责任", "失败与退回", "修改和批准权限", "留痕", "签字确认"]
    doc = new_doc("资料交付_审核_上架流程", "从产品与履约团队提交，到上线使用、到期提醒和变更归档")
    add_toc(doc, sections)
    doc.add_heading("流程总图", level=1)
    add_steps(doc, [step[1] for step in d.PROCESS_STEPS])
    doc.add_heading("逐步责任", level=1)
    add_table(
        doc,
        ["步骤", "动作", "输入", "输出", "执行/审批", "时限"],
        [[number, action, input_, output, role_name(owner), sla] for number, action, input_, output, owner, sla in d.PROCESS_STEPS],
        [0.55, 1.45, 1.3, 1.35, 1.3, 0.95],
        7.8,
    )
    doc.add_heading("失败与退回", level=1)
    add_table(
        doc,
        ["失败点", "处理", "恢复条件"],
        [
            ("格式不完整", "退回并列出缺失项；不进入真实性检查", "按模板重新提交"),
            ("真实性无法确认", "标记待验证；限制访问；通知责任方", "补原件、登记或第三方证明"),
            ("商品无法匹配", "不进入事实库", "补SKU编号和适用范围"),
            ("文件过期", "相关商品或动作降级", "新文件通过"),
            ("当地确认不通过", "停止上架、投流或成交", "修改后重新取得书面意见"),
            ("上线后发生变化", "评估影响、暂停相关内容/商品、归档旧版", "新版本完成全链路审核"),
        ],
        [1.7, 2.8, 2.0],
        8.4,
    )
    doc.add_heading("修改和批准权限", level=1)
    add_table(
        doc,
        ["内容", "可提交", "可修改草稿", "最终批准"],
        [
            ("商品与质量", "中国产品供应方", "供应负责人", "中国产品供应方负责人"),
            ("进口、许可与标签", "尼泊尔进口主体/合规顾问", "进口与合规人员", "尼泊尔进口主体负责人"),
            ("价格、收款与售后", "销售主体/仓储团队", "商务与售后人员", "尼泊尔销售主体负责人"),
            ("库存、批次与配送", "仓储/物流", "仓库和物流人员", "仓储与履约负责人"),
            ("内容与FAQ", "线上运营团队/AI技术执行方", "内容和运营人员", "线上运营团队负责人＋必要的品牌/合规审核"),
            ("平台账号和广告", "平台账户人员", "账户管理员", "平台与广告账户主体负责人"),
        ],
        [1.65, 1.6, 1.45, 1.8],
        8.2,
    )
    doc.add_heading("留痕", level=1)
    add_table(
        doc,
        ["必须记录", "最低内容"],
        [
            ("提交", "提交人、时间、文件名、版本、适用商品"),
            ("审核", "审核人、意见、通过/退回、日期"),
            ("批准", "批准人、范围、生效时间、有效期"),
            ("使用", "使用渠道、内容/报价/订单编号、引用版本"),
            ("变更", "原值、新值、影响、通知、下架或改内容"),
            ("异常", "触发、证据、隔离、退款/补发/召回、复盘"),
        ],
        [1.35, 5.15],
        8.5,
    )
    add_signature(doc)
    return save(doc, "15_资料交付_审核_上架流程.docx")


def doc_16_sla():
    sections = ["SLA说明", "响应时限", "变更通知", "超时处理", "签字确认"]
    doc = new_doc("SLA响应时限与变更通知机制", "价格、库存、订单、售后、质量、许可和包装变化的明确时限", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("SLA说明", level=1)
    add_callout(
        doc,
        "开始计时",
        "以系统、邮件或双方认可的工作群收到完整信息为开始。聊天中的口头回复应在1个工作日内补录到本DOCX或系统；高风险事件不等补录才处理。",
        PALE_GREEN,
    )
    doc.add_heading("响应时限", level=1)
    add_table(
        doc,
        ["事项", "目标时限", "主责", "超时升级", "备用负责人", "确认方式", "当前接受"],
        [[a, b, role_name(c), d_, "________________", "时间戳＋处理记录", "□同意 / □调整"] for a, b, c, d_ in d.SLA_ITEMS],
        [1.4, 1.25, 1.15, 1.65, 1.1, 1.25, 1.05],
        7.1,
    )
    doc.add_heading("变更通知", level=1)
    add_table(
        doc,
        ["必须记录", "填写"],
        [
            ("变更内容", "________________"),
            ("原值", "________________"),
            ("新值", "________________"),
            ("生效时间", "________________"),
            ("影响SKU", "________________"),
            ("影响平台与渠道", "________________"),
            ("影响采购单位或购买者", "________________"),
            ("是否下架/停投/改内容/通知", "________________"),
            ("操作人、批准人和备用人员", "________________"),
        ],
        [2.35, 4.15],
        8.4,
    )
    doc.add_heading("超时处理", level=1)
    add_table(
        doc,
        ["超时类型", "自动处理", "人工处理"],
        [
            ("价格或库存确认超时", "停止报价与下单", "销售和仓库负责人确认"),
            ("许可或授权通知超时", "相关商品降级为阻断", "合规与主体负责人处理"),
            ("订单或配送更新超时", "标记异常，不承诺送达", "仓库和物流调查"),
            ("售后或退款超时", "升级项目与财务负责人", "说明原因和新时限"),
            ("重大质量或召回超时", "立即停卖、停投并扩大隔离", "决策组接管"),
        ],
        [1.85, 2.25, 2.4],
        8.3,
    )
    add_signature(doc, "SLA确认")
    return save(doc, "16_SLA响应时限与变更通知机制.docx")


def doc_17_incidents():
    sections = ["应急原则", "异常矩阵", "AI与人工边界", "证据和对外说明", "复盘", "签字确认"]
    doc = new_doc("异常_客诉_召回应急机制", "汾酒与海鲜的触发、人工接管、负责人、停售、停投、通知、赔付和复盘", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("应急原则", level=1)
    add_steps(doc, ["先保护人身与食品安全", "停止受影响商品、批次、内容或广告", "隔离库存与订单", "保留照片、视频、温控、批次、付款和会话证据", "由最终负责人决定退款、补发、召回或赔付", "统一人工说明", "修复事实源、流程与权限后再恢复"])
    doc.add_heading("异常矩阵", level=1)
    rows = []
    for project, incident, trigger, ai_allowed, first, final, action, notify_all in d.INCIDENTS:
        rows.append(
            [
                project,
                incident,
                trigger,
                ai_allowed,
                role_name(first),
                role_name(final),
                action,
                "涉及安全、质量、许可、资金、舆情时立即",
                "证据齐全前保持暂停",
                notify_all,
                "按已批准规则；例外需最终负责人",
                "事件结束后2个工作日",
            ]
        )
    add_table(
        doc,
        ["项目", "异常", "触发", "AI可处理", "第一负责人", "最终负责", "立即动作", "转人工", "下架/停投", "通知全部", "退款/赔付", "复盘"],
        rows,
        [0.6, 1.15, 1.4, 0.7, 1.0, 1.0, 1.55, 1.15, 1.0, 0.75, 1.15, 0.85],
        6.25,
    )
    doc.add_heading("AI与人工边界", level=1)
    add_table(
        doc,
        ["AI可以", "AI不可以"],
        [
            ("确认已登记、收集订单/批次/照片/视频、说明人工正在处理、提供已批准的安全提示", "判断食品是否安全、确定假货、批准退款、承诺赔付、公开归责、决定召回"),
        ],
        [3.25, 3.25],
        8.4,
    )
    doc.add_heading("证据和对外说明", level=1)
    add_table(
        doc,
        ["项目", "最低证据", "说明原则"],
        [
            ("订单与付款", "订单编号、金额、时间、票据", "不公开个人资料"),
            ("商品与批次", "SKU、批次、标签、包装、照片", "只说已核实事实"),
            ("配送与温控", "取件、路线、时间、温度、签收", "不在调查前推责"),
            ("平台与内容", "URL、版本、审核、驳回、申诉", "暂停争议内容"),
            ("退款与赔付", "批准、执行、到账和实物处置", "明确时限与联系人"),
        ],
        [1.4, 2.7, 2.4],
        8.4,
    )
    doc.add_heading("复盘", level=1)
    add_table(
        doc,
        ["事件编号", "根因", "哪些商品/渠道受影响", "处理结果", "流程修改", "负责人", "恢复批准"],
        [["________________", "________________", "________________", "________________", "________________", "________________", "________________"]],
        [1.0, 1.25, 1.45, 1.2, 1.2, 0.9, 1.0],
        8.0,
    )
    add_signature(doc)
    return save(doc, "17_异常_客诉_召回应急机制.docx")


def doc_18_ownership():
    sections = ["使用说明", "归属确认事项", "推荐默认原则", "终止与导出", "签字确认"]
    doc = new_doc("账号_客户_数据_内容归属确认事项", "不替双方擅自决定：提供选项、风险、推荐默认方案和签字栏", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("使用说明", level=1)
    add_callout(
        doc,
        "未决事项",
        "本文件不代替合同。表中推荐方案用于减少账号失控、数据不能导出、隐私责任错位和复购争议；双方可调整，但必须写明选择、负责人和生效日期。",
        PALE_GOLD,
    )
    doc.add_heading("归属确认事项", level=1)
    add_table(
        doc,
        ["事项", "推荐默认方案", "其他选项", "主要风险", "当前状态", "双方选择", "责任人", "生效日"],
        [[a, b, c, d_, e, "________________", "________________", "________________"] for a, b, c, d_, e in d.OWNERSHIP_ITEMS],
        [1.05, 2.15, 1.45, 1.55, 0.85, 1.2, 0.9, 0.85],
        6.8,
    )
    doc.add_heading("推荐默认原则", level=1)
    add_table(
        doc,
        ["原则", "说明"],
        [
            ("合法主体持有高风险资产", "广告、收款、票据和本地销售相关账户由实际承担法律责任的主体持有"),
            ("线上运营团队保留工作连续性", "获得完成运营所需的管理员、数据查看和导出权限"),
            ("项目专用恢复方式", "邮箱、手机号、双重验证和备用管理员不得依赖单一个人"),
            ("可完整导出", "商品、CRM、订单、内容、审批、广告和日志按约定格式导出"),
            ("数据按目的使用", "购买者资料由本地销售主体依法保管，线上运营团队只在约定目的和期限内使用"),
            ("品牌与原创分开", "品牌商标和官方素材归品牌方；原创内容与项目配置按委托协议确定"),
        ],
        [1.65, 4.85],
        8.5,
    )
    doc.add_heading("终止与导出", level=1)
    add_table(
        doc,
        ["事项", "建议时限", "验收"],
        [
            ("停止新增外联、投放和内容", "终止通知生效时", "无新任务"),
            ("导出CRM、商品、订单、审批和日志", "5个工作日", "文件可打开、数量可核"),
            ("移交账号、域名、像素与管理员", "5个工作日", "新管理员登录成功"),
            ("移除不再需要的权限", "移交完成后1个工作日", "权限清单确认"),
            ("按约定删除或保留个人资料", "法律与合同期限内", "删除证明或保留依据"),
            ("商品、品牌和供应资料下架", "授权终止时", "公开渠道不再展示"),
        ],
        [2.55, 1.35, 2.6],
        8.4,
    )
    add_signature(doc, "归属与终止条款确认")
    return save(doc, "18_账号_业务关系_数据_内容归属确认事项.docx")


def doc_19_decisions():
    sections = ["使用说明", "待确认决策", "优先顺序", "决策记录", "签字确认"]
    doc = new_doc("双方待确认决策清单", "把供应、品牌、尼泊尔当地、平台后台和双方商业选择分开", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("使用说明", level=1)
    add_callout(doc, "状态", "未完成证据栏前，所有决策保持待确认。口头同意不视为完成；涉及账号、收益、个人资料和授权的事项必须签字。", PALE_RED)
    doc.add_heading("待确认决策", level=1)
    add_table(
        doc,
        ["编号", "决策", "确认方", "需要决定", "完成证据", "优先级", "目标日期", "当前状态", "最终选择"],
        [[code, decision, role_name(owner), detail, evidence, priority, "____年__月__日", "待确认", "________________"] for code, decision, owner, detail, evidence, priority in d.DECISION_ITEMS],
        [0.55, 1.35, 1.2, 1.75, 1.25, 0.65, 1.0, 0.85, 1.4],
        6.9,
    )
    doc.add_heading("优先顺序", level=1)
    add_steps(doc, ["先完成合法主体、授权、许可和商品身份", "再完成价格、库存、配送、收款和售后", "再完成平台账号与内容/广告能力", "再签账号、数据、复购和费用条款", "最后批准90天试点和真实预算"])
    doc.add_heading("决策记录", level=1)
    add_table(
        doc,
        ["决策编号", "日期", "选择", "理由", "影响", "批准人", "复查日"],
        [["________________", "________________", "________________", "________________", "________________", "________________", "________________"]],
        [1.0, 1.0, 1.45, 1.65, 1.45, 1.0, 0.95],
        7.9,
    )
    add_signature(doc)
    return save(doc, "19_双方待确认决策清单.docx")


def doc_20_confirmation_letter():
    sections = ["项目背景", "职责确认", "需提交的P0资料", "汾酒专属资料", "海鲜专属资料", "交付与更新", "缺失影响", "订单与售后", "归属待确认", "回复与签字"]
    doc = new_doc("供应链正式资料提供及职责确认函", "可直接发给产品、进口、销售、仓储与履约团队填写和签字")
    add_toc(doc, sections)
    doc.add_heading("项目背景", level=1)
    add_para(
        doc,
        "为推进汾酒与海鲜产品在尼泊尔的线上展示、询问承接、B2B/B2C销售验证和AI辅助运营，双方需要先把产品、授权、许可、价格、库存、配送、收款和售后责任书面确认。本函不表示任何商品已经取得许可或已经具备正式销售条件。",
    )
    doc.add_heading("职责确认", level=1)
    add_table(
        doc,
        ["责任方", "负责事项", "不转移给对方的责任"],
        [
            ("线上运营团队", "市场研究、渠道、内容、AI内容、线上询问、CRM和复盘", "不承担进口、许可、仓储、配送、质量、退款资金、召回、赔付和票据"),
            ("产品与履约团队", "真实商品、授权、许可、价格、库存、进口、仓储、配送、收款、售后和质量责任", "不能把实物与当地责任交给线上内容人员承担"),
            ("共同事项", "商品上架、正式报价、内容发布、异常通知、阶段决策", "任何共同动作仍需明确一个最终负责人"),
        ],
        [1.6, 3.0, 1.9],
        8.5,
    )
    doc.add_heading("需提交的P0资料", level=1)
    p0 = [item for item in d.MASTER_DELIVERY_ITEMS if item[-1] == "P0"]
    add_table(
        doc,
        ["编号", "资料", "提供方", "截止时间", "责任人", "当前状态"],
        [[code, f"{category}：{name}", role_name(provider), "收到本函后7个自然日", "________________", "待提供"] for code, category, name, _, provider, _, _ in p0],
        [0.65, 2.25, 1.45, 1.2, 1.2, 0.85],
        7.5,
    )
    doc.add_heading("汾酒专属资料", level=1)
    add_table(
        doc,
        ["编号", "需要提供或验证", "责任方", "截止时间", "状态"],
        [[code, name, role_name(provider), "公开动作前", "待提供/待验证"] for code, name, provider, _, _ in d.FENJIU_SPECIAL_ITEMS],
        [0.7, 3.0, 1.4, 1.0, 1.1],
        8.0,
    )
    doc.add_heading("海鲜专属资料", level=1)
    add_table(
        doc,
        ["编号", "需要提供或验证", "责任方", "截止时间", "状态"],
        [[code, name, role_name(provider), "首批商品审核前", "待提供/待验证"] for code, name, provider, _, _ in d.SEAFOOD_SPECIAL_ITEMS],
        [0.7, 3.0, 1.4, 1.0, 1.1],
        8.0,
    )
    doc.add_heading("交付与更新", level=1)
    add_table(
        doc,
        ["要求", "说明"],
        [
            ("交付形式", "填写本套DOCX；证书、许可、检测和授权另附PDF/JPG；不得只发聊天截图"),
            ("适用范围", "每份资料写明适用SKU、批次、主体、国家、渠道和有效期"),
            ("库存与订单", "实时或15分钟内更新"),
            ("价格", "变化前更新并写明生效时间；旧版本自动失效"),
            ("许可与授权", "到期前60日通知"),
            ("包装、停产、缺货", "一般提前2个工作日；已知后不得延迟"),
            ("质量、停售、召回", "立即通知并启动应急"),
        ],
        [1.45, 5.05],
        8.5,
    )
    doc.add_heading("缺失影响", level=1)
    add_callout(
        doc,
        "阻断",
        "资料缺失不会由AI或线上运营团队补猜。缺主体、授权、许可、商品身份、价格、库存、配送、收款、售后或质量责任，将分别阻断发布、报价、上架、投流、成交或放量。",
        PALE_RED,
    )
    doc.add_heading("订单与售后", level=1)
    add_table(
        doc,
        ["事项", "由谁执行", "需要在本函确认"],
        [
            ("订单确认、收款和票据", "尼泊尔销售与收款主体", "主体、时限、联系方式"),
            ("库存锁定、出库和配送", "仓储与履约团队", "仓库、范围、时效、费用、上限"),
            ("退款、拒收、补发和赔付", "尼泊尔销售主体＋对应责任方", "规则、时限、资金来源、上限"),
            ("质量调查和召回", "产品供应方＋尼泊尔进口主体", "第一负责人、最终负责人、批次追踪"),
        ],
        [2.1, 2.1, 2.3],
        8.3,
    )
    doc.add_heading("归属待确认", level=1)
    add_table(
        doc,
        ["事项", "双方选择"],
        [
            ("TikTok账号、Business Center、广告账户、邮箱和手机号", "________________"),
            ("独立站、域名、CRM、数据库、广告数据和像素", "________________"),
            ("B2B采购单位资料、B2C购买者资料和复购收益", "________________"),
            ("达人资源、内容版权、AI工作流和合作终止导出", "________________"),
            ("广告费、达人费、软件费和其他预算", "________________"),
        ],
        [3.8, 2.7],
        8.4,
    )
    doc.add_heading("回复与签字", level=1)
    add_table(
        doc,
        ["收函方", "公司", "总负责人", "资料联系人", "首次提交日期", "确认意见"],
        [["产品与履约团队", "________________", "________________", "________________", "____年__月__日", "□同意 □需调整"]],
        [1.4, 1.3, 1.1, 1.1, 1.1, 1.0],
        8.2,
    )
    add_signature(doc, "正式确认")
    return save(doc, "20_供应链正式资料提供及职责确认函.docx")


def doc_21_meeting():
    sections = ["会议目标", "参会角色", "会前准备", "会议议程", "必须现场决定", "会议记录", "会后任务"]
    doc = new_doc("供应链首次沟通会议清单", "90分钟完成角色、P0资料、责任人、时限、阻断和下一步确认")
    add_toc(doc, sections)
    doc.add_heading("会议目标", level=1)
    add_callout(doc, "会议结果", "会议结束时，不要求所有资料已经齐全，但必须明确：谁提供、提供什么、何时交、谁复核、缺失会停在哪里。", PALE_GREEN)
    doc.add_heading("参会角色", level=1)
    add_table(doc, ["角色", "姓名/公司", "是否到会"], [[name, "________________", "□"] for name in d.ROLES.values()], [2.8, 2.6, 1.1], 8.5)
    doc.add_heading("会前准备", level=1)
    add_table(
        doc,
        ["准备方", "带到会议"],
        [
            ("线上运营团队", "现有方案、资料包、当前缺口、账号清单、试点建议"),
            ("中国产品供应方", "首批候选SKU、规格、图片、价格、库存、检测、授权和出口资料"),
            ("尼泊尔进口/销售主体", "公司、EXIM Code、许可、收款、票据、仓库、配送和售后资料"),
            ("合规顾问", "酒类广告/销售与海鲜进口/标签的初步问题单"),
            ("平台账户负责人", "实际账号、Business Center和后台截图"),
        ],
        [1.9, 4.6],
        8.4,
    )
    doc.add_heading("会议议程", level=1)
    add_table(
        doc,
        ["时间", "议题", "输出"],
        [
            ("0-10分钟", "项目目标与不可越过的责任边界", "边界确认"),
            ("10-25分钟", "12个角色由谁承担", "角色与主备负责人"),
            ("25-45分钟", "汾酒P0与TikTok酒类闸门", "书面验证任务"),
            ("45-65分钟", "海鲜P0、B2B/B2C与冷链", "首批SKU资料任务"),
            ("65-75分钟", "价格、库存、订单、退款和SLA", "更新与应急时限"),
            ("75-85分钟", "账号、数据、复购和费用", "待签选择"),
            ("85-90分钟", "复述任务、截止时间和下一次会议", "会议行动表"),
        ],
        [1.15, 2.7, 2.65],
        8.5,
    )
    doc.add_heading("必须现场决定", level=1)
    add_table(
        doc,
        ["问题", "决定"],
        [
            ("谁是尼泊尔进口、销售、收款、仓储和售后的最终负责人？", "________________"),
            ("汾酒和海鲜首批分别准备哪些SKU？", "________________"),
            ("P0资料最晚何时提交？", "________________"),
            ("价格、库存和配送由什么方式更新？", "________________"),
            ("重大质量、许可或平台问题通知谁？", "________________"),
            ("下一次复核会议何时进行？", "________________"),
        ],
        [4.35, 2.15],
        8.4,
    )
    doc.add_heading("会议记录", level=1)
    add_table(
        doc,
        ["编号", "讨论", "结论", "状态", "负责人", "截止日", "证据"],
        [[str(i), "________________", "________________", "待确认", "________________", "________________", "________________"] for i in range(1, 11)],
        [0.5, 1.5, 1.5, 0.75, 0.95, 0.85, 1.0],
        8.0,
    )
    doc.add_heading("会后任务", level=1)
    add_steps(doc, ["24小时内发送会议记录", "3个工作日内提交首批主体与授权资料", "7个自然日内提交P0商品、价格、库存、配送和售后资料", "当地合规与平台测试按责任表完成", "P0通过后再安排模拟测试"])
    add_signature(doc, "会议纪要确认")
    return save(doc, "21_供应链首次沟通会议清单.docx")


def doc_22_tests():
    sections = ["测试原则", "模拟测试", "通过条件", "问题单", "最终验收"]
    doc = new_doc("上线前模拟测试与验收表", "用合成数据和受控账号证明阻断、审批、同步、履约和回退", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("测试原则", level=1)
    add_callout(
        doc,
        "先模拟后真实",
        "没有P0资料时可以用明确标记的合成数据测试系统，但不能把测试结果当真实商品、真实账号或真实销售能力。真实小样只在主体、许可和履约通过后进行。",
        PALE_GOLD,
    )
    doc.add_heading("模拟测试", level=1)
    add_table(
        doc,
        ["编号", "场景", "输入", "期望结果", "责任", "等级", "实际结果", "证据路径", "结论"],
        [[code, scene, input_, expected, role_name(owner), level, "________________", "________________", "□通过 □失败"] for code, scene, input_, expected, owner, level in d.TEST_ITEMS],
        [0.55, 1.15, 1.45, 1.75, 1.1, 0.6, 1.4, 1.35, 1.0],
        6.8,
    )
    doc.add_heading("通过条件", level=1)
    add_table(
        doc,
        ["等级", "通过要求"],
        [
            ("P0", "全部通过；任何失败都阻断对应上线动作"),
            ("P1", "可带整改计划进入有限测试，但不得放量"),
            ("P2", "记录并进入下一迭代"),
        ],
        [1, 5.5],
        8.5,
    )
    doc.add_heading("问题单", level=1)
    add_table(
        doc,
        ["问题编号", "关联测试", "现象", "影响", "负责人", "修复日期", "复测结果"],
        [["________________", "________________", "________________", "________________", "________________", "________________", "________________"]],
        [1.0, 1.0, 1.55, 1.35, 1.0, 1.05, 1.2],
        7.8,
    )
    doc.add_heading("最终验收", level=1)
    add_table(
        doc,
        ["验收项", "结论", "验收人", "日期"],
        [
            ("所有P0测试通过", "□是 □否", "________________", "________________"),
            ("所有真实主体与许可通过", "□是 □否", "________________", "________________"),
            ("商品、价格、库存、配送和售后当前版本有效", "□是 □否", "________________", "________________"),
            ("平台与当地限制已按实际账号和书面意见确认", "□是 □否", "________________", "________________"),
            ("回退与应急已演练", "□是 □否", "________________", "________________"),
        ],
        [3.2, 1.0, 1.25, 1.05],
        8.3,
    )
    add_signature(doc, "上线验收批准")
    return save(doc, "22_上线前模拟测试与验收表.docx")


def doc_23_missing():
    sections = ["判断方法", "资料缺失影响", "组合阻断", "补齐顺序", "签字确认"]
    doc = new_doc("资料缺失影响与阻断对照表", "明确缺什么、停在哪里、由谁补、补齐后如何恢复", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("判断方法", level=1)
    add_callout(doc, "原则", "每一项缺失都必须有现实后果。不得用“后续补充”掩盖对发布、报价、上架、投流、收款、配送、退款或召回的影响。", PALE_RED)
    doc.add_heading("资料缺失影响", level=1)
    add_table(
        doc,
        ["缺失资料", "直接影响", "被阻断动作", "等级", "补充方", "恢复证据", "当前状态"],
        [[item, impact, blocked, level, role_name(owner), "资料通过审核并更新准入状态", "待提供"] for item, impact, blocked, level, owner in d.MISSING_IMPACTS],
        [1.35, 2.1, 1.55, 0.65, 1.2, 1.55, 0.9],
        7.0,
    )
    doc.add_heading("组合阻断", level=1)
    add_table(
        doc,
        ["组合", "当前可做", "不能做"],
        [
            ("有素材，无现货", "内部样片", "写现货、报价、上架、成交"),
            ("有商品，无价格", "商品资料审核", "正式报价"),
            ("有库存，无配送", "内部需求收集", "接单"),
            ("有授权，无平台资格", "内部内容准备", "广告、挂载、直播成交"),
            ("有许可，无售后", "流程准备", "正式成交"),
            ("系统可用，数据靠临时聊天", "内部演示", "放量"),
        ],
        [2.0, 2.15, 2.35],
        8.4,
    )
    doc.add_heading("补齐顺序", level=1)
    add_steps(doc, ["合法主体与授权", "商品身份与逐项合规", "价格、库存、配送、收款和售后", "平台账号与内容能力", "AI事实源与权限", "模拟测试", "真实小样", "放量决策"])
    add_signature(doc)
    return save(doc, "23_资料缺失影响与阻断对照表.docx")


def doc_24_codex():
    sections = ["执行边界", "任务总表", "依赖顺序", "测试与验收", "上线条件", "交接"]
    doc = new_doc("后续Codex5.5执行任务包", "把已确认资料转换成可测试、可审计、可回退的AI销售协同系统", landscape=True)
    add_toc(doc, sections)
    doc.add_heading("执行边界", level=1)
    add_callout(
        doc,
        "不得代替现实责任",
        "Codex只构建事实源、接口、审批、日志、告警、内容检查、CRM和测试。它不得生成虚假许可、授权、商品、价格、库存、案例或平台能力，也不得在未经批准时发送、发布、报价、退款或召回。",
        PALE_RED,
    )
    doc.add_heading("任务总表", level=1)
    add_table(
        doc,
        ["编号", "任务", "目标", "依赖", "验收", "状态", "负责人", "证据"],
        [[code, task, goal, dependency, acceptance, "待执行", "AI系统与技术执行方", "________________"] for code, task, goal, dependency, acceptance in d.CODEX_TASKS],
        [0.55, 1.35, 2.0, 1.25, 1.75, 0.8, 1.25, 1.25],
        6.9,
    )
    doc.add_heading("依赖顺序", level=1)
    add_steps(doc, ["C01-C04：事实源、提交、匹配、到期", "C05-C07：价格、库存、配送", "C08-C10：AI咨询、内容检查、TikTok酒类闸门", "C11-C13：订单、召回、CRM", "C14：权限与审计", "C15：模拟验收与回退"])
    doc.add_heading("测试与验收", level=1)
    add_table(
        doc,
        ["层级", "证明"],
        [
            ("规则", "正常、缺失、过期、冲突均有明确结果"),
            ("集成", "商品、价格、库存、订单、支付、配送和售后能连通"),
            ("权限", "未授权角色不能改价、发布、退款或修改事实"),
            ("审计", "来源、版本、草稿、批准、发送和结果可回看"),
            ("业务演练", "24项模拟测试有证据"),
            ("真实小样", "只有P0通过后在少量SKU和一个城市运行"),
        ],
        [1.15, 5.35],
        8.4,
    )
    doc.add_heading("上线条件", level=1)
    add_table(
        doc,
        ["条件", "当前状态"],
        [
            ("P0主体、授权、商品、合规、价格、库存、配送、收款、售后完整", "待提供"),
            ("尼泊尔当地书面意见完成", "待验证"),
            ("平台实际账号能力完成", "待平台测试"),
            ("全部P0模拟测试通过", "待执行"),
            ("真实小样无重大合规、质量、支付或履约问题", "未执行"),
        ],
        [4.6, 1.9],
        8.5,
    )
    doc.add_heading("交接", level=1)
    add_table(
        doc,
        ["交接资料", "提供方", "状态"],
        [
            ("本套DOCX与已填写版本", "项目负责人", "已生成/待填写"),
            ("证据附件", "各责任方", "待提供"),
            ("账号与权限清单", "平台和系统负责人", "待确认"),
            ("测试数据与测试报告", "技术执行方", "待执行"),
            ("上线、回退与应急负责人", "各业务负责人", "待确认"),
        ],
        [3.0, 1.8, 1.7],
        8.5,
    )
    return save(doc, "24_后续Codex5.5执行任务包.docx")


def doc_master_total():
    sections = [
        "执行摘要",
        "双方责任边界",
        "当前共同结论",
        "汾酒执行门槛",
        "海鲜执行门槛",
        "资料交付重点",
        "AI协同规则",
        "30/60/90天推进",
        "重大风险与应急",
        "双方待决定事项",
        "来源与下一步",
    ]
    doc = new_doc(
        "汾酒＋海鲜尼泊尔线上销售\n双方职责与供应链交付总方案",
        "合作方阅读版｜责任、资料、平台、合规、履约、AI协同与上线闸门一体化总册",
        status="部分完成：体系已建立，真实资料与外部资格待确认",
    )
    add_toc(doc, sections)
    doc.add_heading("执行摘要", level=1)
    add_callout(
        doc,
        "主结论",
        "本项目已经形成可供双方确认的责任与资料体系，但尚不具备直接宣布正式销售、正式投放或全国放量的条件。优先工作是补齐真实主体、首批商品、许可、价格、库存、配送、售后和平台后台证据，再在加德满都谷地开展受控验证。",
        PALE_GREEN,
    )
    add_table(
        doc,
        ["项目", "当前可做", "当前不可做", "下一道门槛"],
        [
            ("汾酒", "责任确认、资料准备、内部内容草拟、后台资格核验", "默认开展酒类营销、导流、广告、直播成交或Shop售酒", "当地书面意见＋平台实际账号批准＋完整销售链路"),
            ("海鲜", "首批商品筛选、资料收集、B2B名单研究、内部商品页和咨询演练", "无真实库存、冷链、许可和售后时公开报价或接单", "逐商品合规＋实时价格库存＋冷链与售后演练"),
            ("AI系统", "建立事实源、审批、权限、日志、提醒和模拟测试", "自行生成事实、自动承诺、未经批准发布或收款", "真实资料接入＋24项测试通过＋人工批准"),
        ],
        [1.0, 2.05, 2.15, 2.1],
        8.2,
    )
    add_status_legend(doc)

    doc.add_heading("双方责任边界", level=1)
    add_table(
        doc,
        ["责任方", "主要承担", "明确不承担"],
        [
            (
                "线上运营团队",
                "市场研究、内容草拟、渠道运营、公开商务线索整理、咨询分流、CRM、数据复盘和AI工作流协调",
                "进口、报关、许可、仓储、冷链、配送、收款、开票、退款资金、召回、质量赔付和当地最终法律责任",
            ),
            (
                "产品与履约团队",
                "真实商品、授权、出口与进口资料、价格、库存、仓储、配送、收款、票据、售后、质量和召回",
                "未经审核代替线上运营团队发布、修改运营数据或对外使用未批准内容",
            ),
            (
                "双方共同",
                "试点范围、预算、账号归属、资料用途、复购归因、内容权利、批准机制、停线和放量决定",
                "以口头约定替代P0书面确认，或把单方工作记录当作双方正式决定",
            ),
        ],
        [1.2, 2.65, 2.65],
        8.5,
    )
    add_para(
        doc,
        "完整逐项责任与主责、最终负责、会签、知会、时效和验收标准，分别见02、03号RACI责任矩阵。",
    )

    doc.add_heading("当前共同结论", level=1)
    add_table(
        doc,
        ["判断", "状态", "管理方式"],
        [
            ("先以加德满都谷地为首个受控验证范围", "建议采用", "双方确认后写入90天试点批准单"),
            ("现有研究资料可作为计划依据", "部分成立", "动态法律、平台、价格、库存和商品事实行动前再核"),
            ("汾酒TikTok可以默认投流、挂载或直播成交", "不成立", "获得当地书面意见和实际平台批准前保持阻断"),
            ("海鲜可以在缺少真实SKU和冷链数据时正式报价", "不成立", "内部研究与草拟可以继续，公开承诺保持阻断"),
            ("AI可以提高资料检查与响应效率", "部分成立", "AI只读已批准事实，价格、质量、退款和酒类咨询转人工"),
            ("文档完成即代表业务已经执行", "不成立", "把资料完成、测试通过、真实试单和放量决定分开记录"),
        ],
        [3.2, 1.0, 2.3],
        8.4,
    )

    doc.add_heading("汾酒执行门槛", level=1)
    add_table(
        doc,
        ["能力层", "至少需要", "未通过时"],
        [
            ("普通教育内容", "真实品牌和商品事实、素材权、商业披露、当地传播边界审核", "只保留内部草稿"),
            ("受管制行业企业能力", "合法主体、品牌授权、酒类许可、TikTok后台批准", "普通账号不得默认开展营销或导流"),
            ("广告", "当地书面意见、可开户地区与行业批准、年龄定向、落地页和预算批准", "不得投流"),
            ("LIVE", "内容、年龄、商业披露、受管制商品与导流边界均通过", "不得直播推广或成交"),
            ("外链或联系入口", "平台规则允许、当地销售路径合法、入口与披露经过批准", "不得以链接、联系方式或暗示绕过限制"),
            ("Shop与站内支付", "尼泊尔市场和酒类类目实际可用、主体和履约资格通过", "不得承诺站内售酒"),
        ],
        [1.25, 3.25, 2.0],
        8.2,
    )
    add_callout(
        doc,
        "保守处理",
        "TikTok当前公开规则把酒类及其营销、交易和导流列入受管制范围，并只在部分地区、部分已验证主体中设有限例外；公开Shop市场清单也未列出尼泊尔。因此本方案以实际账号回执为准，不把公开内容能力外推为售酒能力。",
        PALE_RED,
    )

    doc.add_heading("海鲜执行门槛", level=1)
    add_table(
        doc,
        ["路线", "核心动作", "必须具备", "停止条件"],
        [
            ("B2B", "公开单位研究、人工批准联系、试样、阶梯报价、订单与复购", "真实商品、阶梯价、样品规则、进口路径、库存、冷链、账期与售后", "价格过期、库存不明、冷链或许可异常"),
            ("B2C", "内容、咨询、商品页、小范围投流、订单、配送与复购", "逐商品准入、B2C价、实时库存、配送范围、收款、退款与质量处理", "超卖、超区、支付失败、冷链或质量异常"),
            ("共同底座", "商品编号、批次、文件关联、事实库、权限、日志和召回", "一个实物版本一个编号，价格与库存有时间戳，批次可追溯", "资料冲突、过期、无法匹配或越权"),
        ],
        [0.9, 2.15, 2.55, 1.35],
        7.9,
    )

    doc.add_heading("资料交付重点", level=1)
    add_table(
        doc,
        ["优先级", "必须先交付", "完成证据", "被阻断动作"],
        [
            ("P0-主体", "品牌授权、进口主体、销售与收款主体、税务与许可", "有效文件＋审核记录", "公开推广、进口、收款、销售"),
            ("P0-商品", "逐商品身份、规格、标签、检测、海关编码、批次和储存要求", "实物与文件一致", "上架、报价、AI答复、履约"),
            ("P0-交易", "批准价格、实时库存、配送、收款、票据、退款和质量责任", "批准版本＋演练记录", "报价、接单、出库和售后"),
            ("P0-平台", "账号归属、行业资格、广告、LIVE、外链和Shop后台结果", "后台回执或工单", "对应平台动作"),
            ("P1-增长", "内容计划、达人、联盟、广告测试、复购和归因规则", "双方批准单", "放量与长期结算"),
        ],
        [0.95, 2.75, 1.65, 1.5],
        8.2,
    )
    add_para(doc, "完整交付信息项、提交方、审核方、格式、频率、有效期、渠道、AI读取规则、缺失影响、SLA和备用负责人，见04、05、06号清单。")

    doc.add_heading("AI协同规则", level=1)
    add_table(
        doc,
        ["环节", "系统可做", "必须人工决定"],
        [
            ("资料", "登记、编号、格式检查、商品匹配、到期提醒、版本归档", "真实性、法律适用性和最终批准"),
            ("内容", "依据事实库生成草稿并检查冲突、禁用表述和素材版本", "品牌、法律、平台与发布批准"),
            ("咨询", "回答已批准的一般问题并记录来源", "价格例外、酒类销售、质量、过敏、退款和重大投诉"),
            ("交易", "同步状态、发现冲突、提醒和形成审核单", "报价、收款、退款、召回、赔付和放量"),
            ("治理", "权限限制、操作日志、数据导出和回退", "账号归属、资料用途、保留期限与终止安排"),
        ],
        [1.0, 3.0, 2.5],
        8.2,
    )

    doc.add_heading("30/60/90天推进", level=1)
    add_table(
        doc,
        ["阶段", "目标", "关键交付", "通过条件"],
        [
            ("0—30天", "证据与责任闭环", "主体、首批商品、RACI、平台测试、当地意见、事实源、内部内容与24项测试准备", "所有P0缺口有负责人、时限和阻断"),
            ("31—60天", "受控真实小样", "少量商品、一个城市、B2B试样或合规B2C试单、履约与售后全程留痕", "无重大合规、质量、支付和履约事故"),
            ("61—90天", "复购与阶段决策", "补货、毛利、复购、投诉、配送和内容效果复盘", "双方书面决定扩大、调整或停止"),
        ],
        [1.0, 1.55, 2.8, 1.6],
        8.1,
    )

    doc.add_heading("重大风险与应急", level=1)
    add_table(
        doc,
        ["风险", "立即动作", "最终负责", "恢复条件"],
        [
            ("许可、授权或平台资格失效", "停发、停投、下架、停止收款", "当地销售主体/账号主体", "新证据审核通过"),
            ("价格或库存错误", "停止报价与接单，锁定受影响订单", "销售主体/仓库", "修正并完成影响排查"),
            ("冷链、食品安全或批次异常", "隔离、停售、停止配送，必要时召回", "进口主体/质量负责人", "调查、处置和复核通过"),
            ("酒类年龄或导流问题", "拒绝交易、停止相关内容并保留记录", "当地销售主体", "规则与流程重新通过"),
            ("账号限制或数据异常", "停止新增动作、保护权限和导出记录", "账号主体/系统负责人", "申诉或修复完成并复测"),
        ],
        [1.8, 2.55, 1.45, 1.55],
        8.2,
    )

    doc.add_heading("双方待决定事项", level=1)
    add_table(
        doc,
        ["主题", "必须形成的书面决定", "优先级"],
        [
            ("主体与许可", "进口、销售、收款、税务、票据与质量责任主体", "P0"),
            ("首批商品", "汾酒与海鲜首批商品、批次、数量和试点城市", "P0"),
            ("账号与资料", "账号、Business Center、CRM、域名、资料使用与移交", "P0"),
            ("费用", "广告、达人、软件、样品、配送和超支批准", "P1"),
            ("收益", "复购与再次采购归因期限、计算和结算", "P1"),
            ("知识成果", "内容、素材、数据、AI工作流的使用、导出与终止安排", "P1"),
        ],
        [1.25, 4.45, 0.8],
        8.4,
    )
    add_signature(doc, "总方案确认")
    add_sources(doc)
    return save(doc, "汾酒海鲜尼泊尔线上销售_双方职责与供应链交付总方案.docx")


def doc_evidence_existing():
    sections = ["复核范围", "现有资料索引", "可复用结论", "本轮统一口径", "复核限制"]
    doc = new_doc(
        "现有项目资料与复用索引",
        "说明本轮方案依据了哪些工作区资料，以及哪些内容仍不能视为现实结果",
        landscape=True,
    )
    add_toc(doc, sections)
    doc.add_heading("复核范围", level=1)
    add_callout(
        doc,
        "已完成",
        "已对工作区内汾酒TikTok方案、海鲜总方案、AI系统、供应链资料和90天计划进行结构与关键结论复核，并把可复用部分纳入本套文件。",
        PALE_GREEN,
    )
    doc.add_heading("现有资料索引", level=1)
    rows = []
    for relative_path, purpose, reuse in d.EXISTING_FILES:
        path = ROOT / relative_path
        if path.exists():
            modified = datetime.fromtimestamp(path.stat().st_mtime).strftime("%Y-%m-%d %H:%M")
            status = "已找到"
        else:
            modified = "—"
            status = "未找到"
        rows.append([relative_path, purpose, reuse, status, modified])
    add_table(
        doc,
        ["相对路径", "主要用途", "本轮处理", "状态", "修改时间"],
        rows,
        [2.7, 1.85, 2.35, 0.75, 1.0],
        7.0,
    )
    doc.add_heading("可复用结论", level=1)
    add_table(
        doc,
        ["主题", "可复用内容", "使用限制"],
        [
            ("责任边界", "线上运营与当地进口、销售、收款、履约责任分开", "需双方签字后才成为正式合作口径"),
            ("首城策略", "优先在加德满都谷地受控验证", "仍需确认真实仓库、配送和预算"),
            ("AI架构", "事实源、人工批准、权限、日志与回退", "现有文档不代表真实接口已经接通"),
            ("海鲜路线", "B2B与B2C分开、先少量商品验证", "首批商品、价格、库存和许可仍待提供"),
            ("汾酒路线", "先合规与平台能力，再决定内容、广告或成交", "动态平台与当地法律必须行动前复核"),
        ],
        [1.2, 3.1, 2.2],
        8.1,
    )
    doc.add_heading("本轮统一口径", level=1)
    add_table(
        doc,
        ["容易混淆的说法", "本轮统一表达"],
        [
            ("资料已完成＝项目已执行", "资料完成、模拟测试、真实试单、正式验收和商业结果分别记录"),
            ("TikTok可发内容＝可以销售酒类", "普通内容、行业资格、广告、LIVE、外链、Shop和成交逐层核验"),
            ("公开名单＝真实意向", "公开研究名单只作为候选，不代表已经联系、回复或成交"),
            ("AI方案＝系统已接通", "当前为可执行设计；真实数据源、账号与权限接通后再验收"),
            ("供应链清单＝真实供应能力", "清单用于索取资料；通过审核的真实版本才进入事实源"),
        ],
        [2.45, 4.05],
        8.3,
    )
    doc.add_heading("复核限制", level=1)
    add_para(doc, "本索引证明工作区文件存在及其可复用范围，不证明文件内提及的商品、价格、库存、平台资格、咨询、订单、收入或合作关系已经真实发生。")
    return save(doc, "00_现有项目资料与复用索引.docx", evidence=True)


def doc_evidence_official():
    sections = ["检索方法", "平台官方依据", "尼泊尔官方依据", "适用限制", "行动前复核"]
    doc = new_doc(
        "平台与尼泊尔官方依据",
        "动态规则、法律与许可入口的官方来源索引｜查询日期与适用范围同时记录",
        landscape=True,
    )
    add_toc(doc, sections)
    doc.add_heading("检索方法", level=1)
    add_callout(
        doc,
        "来源原则",
        "优先采用TikTok官方页面、尼泊尔法律委员会、税务局、海关和食品技术与质量控制部门资料。页面结论只支持其明确范围；实际账号能力、逐商品许可与线上销售方式仍需后台测试或当地书面意见。",
        PALE,
    )
    doc.add_heading("平台官方依据", level=1)
    platform_ids = [source["id"] for source in d.SOURCES if source["id"].startswith("T")]
    add_sources(doc, platform_ids)
    doc.add_heading("尼泊尔官方依据", level=1)
    nepal_ids = [source["id"] for source in d.SOURCES if source["id"].startswith(("N", "F"))]
    add_sources(doc, nepal_ids)
    doc.add_heading("适用限制", level=1)
    add_table(
        doc,
        ["主题", "官方页面可支持", "仍需现实确认"],
        [
            ("TikTok酒类", "公开规则中的受管制范围、广告通用要求、商业披露、LIVE和Shop市场清单", "尼泊尔实际账号、行业授权、广告开户、外链、LIVE和Shop能力"),
            ("尼泊尔酒类", "酒类许可、销售时段、年龄、票据、消费税许可和广告限制入口", "网络销售、配送到家、具体内容形式和经营地点适用性"),
            ("海鲜进口", "食品进口申请、EXIM Code、标签和海关税则入口", "逐商品海关编码、许可、检测、标签语言和批次要求"),
            ("电子商务与隐私", "本地法律入口", "具体经营模式、资料收集目的、保存期限与跨境处理"),
            ("抽奖活动", "当地抽奖许可法律入口", "活动结构、奖品、平台和税务适用性"),
        ],
        [1.25, 3.0, 2.25],
        8.0,
    )
    doc.add_heading("行动前复核", level=1)
    add_steps(
        doc,
        [
            "由当地合规顾问逐项形成书面意见，并注明适用主体、商品、渠道、城市和有效日期。",
            "由平台账户主体在真实后台提交或测试，保存回执、工单、截图和限制原因。",
            "由进口主体按首批商品取得逐项预审、许可、海关编码和标签结论。",
            "把新结论更新到事实源，旧版本自动失效；对应测试通过后才开放权限。",
        ],
    )
    return save(doc, "01_平台与尼泊尔官方依据.docx", evidence=True)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    EVIDENCE.mkdir(parents=True, exist_ok=True)
    for apple_double in OUT.rglob("._*"):
        if apple_double.is_file():
            apple_double.unlink()

    generated = [
        doc_00_overview(),
        doc_01_boundaries(),
        doc_raci("02_汾酒项目RACI责任矩阵.docx", "汾酒项目RACI责任矩阵", d.FENJIU_RACI, "汾酒"),
        doc_raci("03_海鲜项目RACI责任矩阵.docx", "海鲜项目RACI责任矩阵", d.SEAFOOD_RACI, "海鲜"),
        doc_04_master_delivery(),
        doc_special(
            "05_汾酒TikTok专属资料清单.docx",
            "汾酒TikTok专属资料清单",
            "普通内容、受管制行业能力、广告、LIVE、外链与Shop逐层确认",
            d.FENJIU_SPECIAL_ITEMS,
            "汾酒",
            "TikTok普通内容、受管制行业企业能力、广告、LIVE、外链、Shop",
            ["T01", "T02", "T03", "T04", "T05", "T06", "T07", "T08", "N01", "N02", "N03", "N04", "N05"],
        ),
        doc_special(
            "06_海鲜B2B_B2C专属资料清单.docx",
            "海鲜B2B与B2C专属资料清单",
            "询价、报价、样品、上架、投流、收款、履约、售后与复购分别管理",
            d.SEAFOOD_SPECIAL_ITEMS,
            "海鲜",
            "TikTok、Meta、YouTube、Daraz、独立站、WhatsApp、Viber、Messenger、Google、达人、直播、联盟",
            ["F01", "F02", "F03", "F04", "F05", "F06", "N06"],
        ),
        doc_07_product_template(),
        doc_08_price_template(),
        doc_09_inventory_template(),
        doc_10_compliance_index(),
        doc_11_content_facts(),
        doc_12_faq(),
        doc_13_ai_interface(),
        doc_14_launch_gate(),
        doc_15_process(),
        doc_16_sla(),
        doc_17_incidents(),
        doc_18_ownership(),
        doc_19_decisions(),
        doc_20_confirmation_letter(),
        doc_21_meeting(),
        doc_22_tests(),
        doc_23_missing(),
        doc_24_codex(),
        doc_master_total(),
        doc_evidence_existing(),
        doc_evidence_official(),
    ]
    for apple_double in OUT.rglob("._*"):
        if apple_double.is_file():
            apple_double.unlink()
    print(f"generated={len(generated)}")
    for path in generated:
        print(path)


if __name__ == "__main__":
    main()
