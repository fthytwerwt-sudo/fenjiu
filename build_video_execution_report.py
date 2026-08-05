#!/usr/bin/env python3
"""Build the Chinese DOCX execution report after final video QC."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
ASCII_FONT = "Calibri"
EAST_ASIA_FONT = "Noto Sans CJK SC"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "555555"
LIGHT_GRAY = "F2F4F7"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", required=True, type=Path)
    return parser.parse_args()


def set_run_font(
    run: Any,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
    ascii_font: str = ASCII_FONT,
) -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table: Any, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    set_run_font(run, size=9, color=GRAY)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = ASCII_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    style_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in style_tokens.items():
        style = document.styles[name]
        style.font.name = ASCII_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("汾酒尼泊尔30秒视频 · 执行报告"), size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_run_font(footer.add_run("第 "), size=9, color=GRAY)
    add_page_field(footer)
    set_run_font(footer.add_run(" 页"), size=9, color=GRAY)


def add_label_paragraph(document: Document, label: str, value: str, *, after: float = 3) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    set_run_font(paragraph.add_run(label), size=10.5, bold=True)
    set_run_font(paragraph.add_run(value), size=10.5)


def set_table_text(table: Any, header: list[str], rows: list[list[str]], widths: list[int]) -> None:
    for column, text in enumerate(header):
        cell = table.rows[0].cells[column]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(text), size=9.5, bold=True)
        shade_cell(cell, LIGHT_GRAY)
    repeat_header(table.rows[0])

    for values in rows:
        cells = table.add_row().cells
        for column, text in enumerate(values):
            cells[column].text = ""
            paragraph = cells[column].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if column in {0, 2, 3, 4} else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_after = Pt(0)
            set_run_font(paragraph.add_run(text), size=9)
    set_table_geometry(table, widths)


def build_report(manifest_path: Path) -> Path:
    output_dir = manifest_path.parent
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    qc_path = output_dir / "06_qc" / "qc_summary.json"
    if not qc_path.is_file():
        raise SystemExit(f"QC summary missing: {qc_path}")
    qc = json.loads(qc_path.read_text(encoding="utf-8"))

    document = Document()
    configure_styles(document)
    configure_page(document)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("视频生成执行报告"), size=23, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(
        subtitle.add_run("汾酒尼泊尔30秒竖屏宣传视频"),
        size=14,
        color=GRAY,
    )
    add_label_paragraph(document, "执行日期：", "2026年7月16日")
    add_label_paragraph(document, "运行编号：", manifest["run_id"])
    add_label_paragraph(document, "分镜来源：", manifest["storyboard_source"])
    add_label_paragraph(document, "输出目录：", str(output_dir), after=10)

    document.add_heading("一、执行结论", level=1)
    lead = document.add_paragraph()
    lead.paragraph_format.space_after = Pt(8)
    set_run_font(
        lead.add_run("生成与剪辑已完成，当前状态：待业务人工复核。"),
        bold=True,
        color=DARK_BLUE,
    )
    set_run_font(
        lead.add_run(
            " 7个计划镜头均已由HappyHorse 1.1系列生成并下载，已合成为30秒、1080×1920、H.264、AAC双声道的竖屏成片和无字幕干净版。因项目未提供官方产品标准照，终版片尾不展示未经确认的具体SKU，改用尼泊尔夕景和后期真实文字落版。"
        )
    )
    add_label_paragraph(document, "技术验收：", qc["final_status"]["technical_validation"])
    add_label_paragraph(document, "内容验收：", qc["final_status"]["content_validation"])
    add_label_paragraph(document, "无水印检查：", qc["final_status"]["watermark_validation"])

    document.add_heading("二、镜头生成情况", level=1)
    shot_rows = []
    for shot in qc["shots"]:
        shot_rows.append(
            [
                str(shot["number"]),
                shot["model"],
                str(shot["generation_count"]),
                "是" if shot["reference_used"] else "否",
                shot["status"],
                shot["note"],
            ]
        )
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    set_table_text(
        table,
        ["镜头", "模型", "生成次数", "参考图", "结果", "说明"],
        shot_rows,
        [600, 1900, 850, 850, 1000, 4160],
    )

    document.add_heading("三、最终文件与技术参数", level=1)
    file_rows = []
    for item in qc["final_files"]:
        file_rows.append(
            [
                item["name"],
                item["duration"],
                item["resolution"],
                item["codec"],
                item["audio"],
                item["size"],
            ]
        )
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    set_table_text(
        table,
        ["文件", "时长", "分辨率", "视频", "音频", "大小"],
        file_rows,
        [3000, 950, 1350, 1150, 1500, 1410],
    )
    for item in qc["final_files"]:
        add_label_paragraph(document, f"{item['name']}：", item["path"], after=3)

    document.add_heading("四、质量检查", level=1)
    for label, value in qc["quality_checks"].items():
        add_label_paragraph(document, f"{label}：", value, after=5)

    document.add_heading("五、API调用记录", level=1)
    api = qc["api_summary"]
    api_rows = [
        ["正式任务总数", str(api["total_tasks"])],
        ["成功任务数", str(api["successful_tasks"])],
        ["失败任务数", str(api["failed_tasks"])],
        ["技术重试", str(api["technical_retries"])],
        ["质量重做", str(api["quality_retries"])],
        ["成功生成总秒数", f"{api['generated_seconds']}秒"],
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_text(table, ["项目", "结果"], api_rows, [2700, 6660])
    add_label_paragraph(
        document,
        "安全说明：",
        "报告、请求记录和响应日志均未写入或显示完整DASHSCOPE_API_KEY。",
        after=8,
    )

    document.add_heading("六、需要人工观看确认的重点", level=1)
    for item in qc["manual_review"]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        set_run_font(paragraph.add_run(f"{item['label']}："), bold=True, color=DARK_BLUE)
        set_run_font(paragraph.add_run(item["detail"]))

    document.add_heading("七、保留的技术记录", level=1)
    records = [
        ("分镜裁切", "01_storyboard_crops/"),
        ("产品参考", "02_product_assets/"),
        ("提示词和镜头配置", "shot_manifest.json"),
        ("任务ID与脱敏响应", "04_tasks/"),
        ("原始镜头", "05_raw_clips/"),
        ("质量检查", "06_qc/"),
        ("标准化片段与字幕图层", "07_edit/"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_text(table, ["记录", "相对路径"], [list(item) for item in records], [2700, 6660])

    final_dir = output_dir / "final"
    final_dir.mkdir(parents=True, exist_ok=True)
    output_path = final_dir / "视频生成执行报告.docx"
    document.core_properties.title = "汾酒尼泊尔30秒视频生成执行报告"
    document.core_properties.subject = "HappyHorse 1.1视频生成与质量验收"
    document.core_properties.author = "项目执行组"
    document.save(output_path)
    return output_path


def main() -> int:
    args = parse_args()
    output = build_report(args.manifest.resolve())
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
