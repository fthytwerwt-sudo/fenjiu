from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document

import supply_collab_data as d


ROOT = Path("/Volumes/WD_BLACK/汾酒尼泊尔")
OUT = ROOT / "汾酒海鲜_尼泊尔线上销售_供应链协同与资料交付体系"
REPORT = ROOT / "_qa_supply_collab_structure.json"
EXPECTED_ROOT = {
    *(f"{number:02d}_" for number in range(25)),
}
BANNED = ("客户", "字段", "用户", "customer", "client")
REQUIRED_HEADINGS = ("阅读目录", "结论", "待确认项", "下一步")


def all_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def check_raci(rows, label: str, failures: list[str]) -> None:
    for row in rows:
        code, action, responsible, accountable, *_ = row
        if not responsible.strip():
            failures.append(f"{label} {code} 缺少R")
        if "," in accountable:
            failures.append(f"{label} {code} 的A不是唯一责任方：{accountable}")
        if accountable not in d.ROLES:
            failures.append(f"{label} {code} 的A无法映射：{accountable}")
        if not action.strip():
            failures.append(f"{label} {code} 缺少任务名称")


def main() -> None:
    failures: list[str] = []
    details: list[dict] = []
    files = sorted(OUT.rglob("*"))
    docx_files = [path for path in files if path.is_file() and path.suffix.lower() == ".docx"]
    unexpected = [
        str(path.relative_to(OUT))
        for path in files
        if path.is_file() and path.suffix.lower() != ".docx"
    ]
    if unexpected:
        failures.append(f"交付目录存在非DOCX文件：{unexpected}")
    if len(docx_files) != 28:
        failures.append(f"DOCX数量应为28，实际为{len(docx_files)}")
    root_docx = [path for path in docx_files if path.parent == OUT]
    numbered = [path for path in root_docx if re.match(r"^\d{2}_", path.name)]
    if len(numbered) != 25:
        failures.append(f"编号DOCX应为25，实际为{len(numbered)}")
    for prefix in EXPECTED_ROOT:
        if not any(path.name.startswith(prefix) for path in numbered):
            failures.append(f"缺少编号前缀：{prefix}")
    if len([path for path in docx_files if path.parent == OUT / "evidence"]) != 2:
        failures.append("evidence目录应有2份DOCX")

    for path in docx_files:
        try:
            with zipfile.ZipFile(path) as archive:
                bad_member = archive.testzip()
                if bad_member:
                    failures.append(f"{path.name} 压缩包损坏：{bad_member}")
            document = Document(path)
            text = all_text(document)
        except Exception as exc:
            failures.append(f"{path.name} 无法打开：{exc}")
            continue
        lower_text = text.lower()
        for banned in BANNED:
            if banned.lower() in lower_text:
                failures.append(f"{path.name} 出现禁用用词：{banned}")
        for heading in REQUIRED_HEADINGS:
            if heading not in text:
                failures.append(f"{path.name} 缺少章节：{heading}")
        if not document.tables:
            failures.append(f"{path.name} 没有表格")
        if len(text.strip()) < 600:
            failures.append(f"{path.name} 内容过少：{len(text.strip())}字符")
        details.append(
            {
                "file": str(path.relative_to(OUT)),
                "bytes": path.stat().st_size,
                "paragraphs": len(document.paragraphs),
                "tables": len(document.tables),
                "characters": len(text),
            }
        )

    check_raci(d.FENJIU_RACI, "汾酒", failures)
    check_raci(d.SEAFOOD_RACI, "海鲜", failures)
    report = {
        "status": "passed" if not failures else "failed",
        "docx_count": len(docx_files),
        "numbered_count": len(numbered),
        "evidence_count": len([path for path in docx_files if path.parent == OUT / "evidence"]),
        "failures": failures,
        "documents": details,
    }
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    if failures:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
