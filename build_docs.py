from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Volumes/WD_BLACK/汾酒尼泊尔")
OUT = ROOT / "outputs" / "20260713_fenjiu_nepal"
ACCESS_DATE = "2026-07-13"
FONT = "Noto Sans CJK SC"
ACCENT = "1F4E79"
ACCENT_2 = "2F75B5"
INK = "1F2937"
MUTED = "667085"
PALE = "EAF2F8"
PALE_GOLD = "FFF4CE"
PALE_RED = "FDECEC"


def load_json(name: str, default: Any) -> Any:
    p = ROOT / name
    if not p.exists():
        return default
    return json.loads(p.read_text(encoding="utf-8"))


root = load_json("research_root.json", {})
culture = load_json("research_culture_compliance.json", {})
execution = load_json("research_execution.json", {})
channels = load_json("research_channels.json", {})


def as_list(value: Any) -> list:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, dict):
        return list(value.values())
    return [value]


def humanize(value: Any) -> str:
    """Render nested research structures as operator-readable text."""
    if value is None:
        return ""
    if isinstance(value, list):
        return "；".join(part for part in (humanize(item) for item in value) if part)
    if isinstance(value, dict):
        labels = {
            "sec": "秒",
            "visual": "画面",
            "line_en": "旁白",
            "line": "台词",
            "action": "动作",
            "shot": "镜头",
        }
        parts = []
        for key, item in value.items():
            rendered = humanize(item)
            if rendered:
                parts.append(f"{labels.get(key, key)}：{rendered}")
        return "｜".join(parts)
    return str(value)


def shown(value: Any) -> str:
    rendered = humanize(value)
    return rendered if rendered else "NEEDS_VERIFY：未取得公开或内部确认"


def first_present(obj: dict, names: Iterable[str], default=None):
    for name in names:
        if name in obj and obj[name] not in (None, "", []):
            return obj[name]
    return default


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths: list[float]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            w = widths[idx]
            cell.width = Inches(w)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(w * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=10.5, bold=None, color=INK, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_pr.extend([r_fonts, color, underline])
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])


def new_doc(title: str, subtitle: str, kind: str = "report", confidentiality: str = "内部执行资料") -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5 if kind == "guide" else 6)
    normal.paragraph_format.line_spacing = 1.18 if kind == "guide" else 1.2
    for name, size, color, before, after in (
        ("Heading 1", 16, ACCENT, 16, 7),
        ("Heading 2", 13, ACCENT_2, 12, 5),
        ("Heading 3", 11.5, "365F91", 8, 4),
    ):
        s = styles[name]
        s.font.name = FONT
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        s = styles[name]
        s.font.name = FONT
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        s.font.size = Pt(10.5)
        s.paragraph_format.space_after = Pt(3)
        s.paragraph_format.line_spacing = 1.15

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(f"汾酒尼泊尔｜{confidentiality}")
    set_run_font(r, size=8, color=MUTED)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("2026-07-13  |  ")
    set_run_font(r, size=8, color=MUTED)
    add_field(fp, "PAGE")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36 if kind == "report" else 18)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(title)
    set_run_font(r, size=25 if kind == "report" else 21, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    set_run_font(r, size=12, color=MUTED)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(18)
    for label, value in (("访问日期", ACCESS_DATE), ("判断体系", "KNOWN / INFERRED / COMPUTED / NEEDS_VERIFY / BLOCKED"), ("版本", "v1.0 执行基线")):
        rr = meta.add_run(f"{label}：")
        set_run_font(rr, size=9, bold=True, color=INK)
        rr = meta.add_run(f"{value}    ")
        set_run_font(rr, size=9, color=MUTED)
    return doc


def add_para(doc: Document, text: str, bold_prefix: str | None = None, style=None, color=INK):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=color)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=color)
    else:
        r = p.add_run(str(text))
        set_run_font(r, color=color)
    return p


def add_bullets(doc: Document, items: Iterable[Any]):
    for item in items:
        if isinstance(item, dict):
            text = "；".join(f"{k}：{v}" for k, v in item.items() if v not in (None, "", []))
        else:
            text = str(item)
        add_para(doc, text, style="List Bullet")


def add_numbered(doc: Document, items: Iterable[Any]):
    for item in items:
        add_para(doc, str(item), style="List Number")


def add_callout(doc: Document, label: str, text: str, fill=PALE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.5, bold=True, color=ACCENT)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[Any]], widths: list[float] | None = None, font_size=8.6):
    rows = [list(r) for r in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, ACCENT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(str(h))
        set_run_font(rr, size=font_size, bold=True, color="FFFFFF")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i in range(len(headers)):
            value = row[i] if i < len(row) else ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if isinstance(value, tuple) and len(value) == 2 and str(value[1]).startswith("http"):
                add_hyperlink(p, str(value[0]), str(value[1]))
            else:
                rr = p.add_run("" if value is None else str(value))
                set_run_font(rr, size=font_size, color=INK)
            if ridx % 2 == 1:
                set_cell_shading(cells[i], "F8FAFC")
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    total = sum(widths)
    widths = [w * 6.5 / total for w in widths]
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_sources(doc: Document, sources: list[dict], title="来源与证据"):
    if not sources:
        return
    doc.add_heading(title, level=1)
    add_callout(doc, "证据说明", "来源记录包含原始链接、时间、来源等级、可信度和对应结论。网页变化快，正式执行前复核访问状态。")
    for s in sources:
        p = doc.add_paragraph(style="List Number")
        title_text = str(s.get("title") or s.get("name") or s.get("id") or "未命名来源")
        url = s.get("url") or s.get("link")
        if url:
            add_hyperlink(p, title_text, url)
        else:
            r = p.add_run(title_text)
            set_run_font(r, size=9)
        meta = f"｜{s.get('published_or_updated') or s.get('publication_or_update_date') or s.get('published') or '日期未标'}｜访问 {s.get('accessed') or s.get('access_date') or ACCESS_DATE}｜{s.get('source_type') or s.get('type') or 'source'}｜{s.get('grade') or s.get('level') or '未分级'}｜{s.get('confidence') or s.get('credibility') or '待判定'}"
        r = p.add_run(meta)
        set_run_font(r, size=8.5, color=MUTED)
        claim = s.get("conclusion") or s.get("claim") or s.get("use")
        if claim:
            r = p.add_run(f"\n对应结论：{claim}")
            set_run_font(r, size=8.5, color=INK)


def save(doc: Document, filename: str):
    OUT.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = filename.rsplit(".", 1)[0]
    doc.core_properties.subject = "汾酒尼泊尔B2B、B2C市场研究与90天执行系统"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Evidence-first; unresolved inputs remain marked NEEDS_VERIFY or BLOCKED."
    doc.save(OUT / filename)


def extract_sources() -> list[dict]:
    all_sources = []
    for obj in (root, culture, execution, channels):
        for key in ("sources", "evidence_sources", "source_index", "evidence_registry"):
            all_sources.extend(as_list(obj.get(key)))
    seen = set()
    out = []
    for s in all_sources:
        if not isinstance(s, dict):
            continue
        token = s.get("url") or s.get("link") or s.get("id") or json.dumps(s, sort_keys=True, ensure_ascii=False)
        if token in seen:
            continue
        seen.add(token)
        out.append(s)
    return out


ALL_SOURCES = extract_sources()


def extract_leads() -> list[dict]:
    raw = first_present(channels, ["leads", "business_leads", "channel_leads", "records"], [])
    if isinstance(raw, dict):
        out = []
        for category, items in raw.items():
            for item in as_list(items):
                if isinstance(item, dict):
                    item = {"category_group": category, **item}
                    out.append(item)
        return out
    return [x for x in as_list(raw) if isinstance(x, dict)]


LEADS = extract_leads()


def doc_reading_order():
    doc = new_doc("项目执行说明与阅读顺序", "从证据、渠道数据库到90天现场执行的总入口", "guide")
    add_callout(doc, "本轮主结论", "首城锁定 Kathmandu Valley；B2B为主、B2C反馈为辅；社交公开酒类推广暂按方案C阻断。最重要的成功指标是试销终端补货，而不是线索量或播放量。")
    doc.add_heading("如何阅读", level=1)
    add_table(doc, ["顺序", "文件", "用途", "主要读者"], [
        [1, "00_待合作方补充的内部数据清单.docx", "先补SKU、价格、库存、利润和主体资料", "国代、品牌方"],
        [2, "01_执行前影响面与阻断检查.docx", "先看哪些能做、哪些必须停", "项目负责人、国代合规负责人"],
        [3, "02_汾酒尼泊尔深度市场调研报告.docx", "城市、渠道、竞争与进入策略", "决策人、销售负责人"],
        [4, "03/04 文化报告与桥接矩阵", "理解人群差异，避免错误文化植入", "内容、培训、销售"],
        [5, "05/06 执行手册与试销机制", "把线索推进到上架、补货", "本地执行人员"],
        [6, "07/07A 社交策略与合规清单", "明确方案C及切换条件", "内容、法务、平台账户负责人"],
        [7, "08/09 AI SOP与外联模板", "批量整理，人工只做现实动作", "运营、AI协作人员"],
        [8, "10-16 Excel + 17 JSON", "日常CRM、评分、脚本、来源、队列", "执行人员、复盘人员"],
        [9, "18_Claude复核任务包.docx", "Claude不可用时的独立复核入口", "外部复核人员"],
        [10, "19_最终执行报告.docx", "总览结论、缺口、未来7天", "所有关键人"],
    ], [0.45, 2.15, 2.35, 1.55])
    doc.add_heading("事实层与判断标签", level=1)
    add_table(doc, ["标签", "含义", "使用规则"], [
        ["KNOWN｜已确认事实", "有来源或用户明确确认", "可作为当前事实引用，但动态网页仍需执行前复查"],
        ["INFERRED｜基于证据推论", "从多项事实推导", "必须说明推导链，不冒充统计事实"],
        ["COMPUTED｜经过计算", "按公开评分/公式得到", "同时保留权重、输入与置信度"],
        ["NEEDS_VERIFY｜需要当地确认", "缺本地语言、商务、现场或平台书面确认", "进入人工队列"],
        ["BLOCKED｜暂时阻断", "不满足条件不得执行", "解除阻断要有证据"],
    ], [1.55, 2.05, 2.9])
    doc.add_heading("交付边界", level=1)
    add_bullets(doc, [
        "本包不研究中国出口、进口清关、税费、仓库和重建国代体系。",
        "国代仅作为供货、仓配、资质/平台主体、合规审核与合同履约方，不被当作渠道名单来源。",
        "SKU、供货价、库存、利润、预算、样品和团队配置未提供时，商业吸引力和ROI不做伪精确计算。",
        "Nepali文案均为草稿，公开前必须母语者审核。",
        "企业线索来自公开来源；没有公开决策人时不编造个人信息。",
    ])
    save(doc, "00_项目执行说明与阅读顺序.docx")


def doc_internal_inputs():
    doc = new_doc("待合作方补充的内部数据清单", "不阻断公开研究，但会阻断报价、试销利润与平台上线", "guide")
    fields = [
        "具体SKU", "酒精度", "容量", "包装和礼盒形式", "国代供货价", "建议零售价", "经销商建议利润", "酒店餐饮建议利润",
        "酒类门店建议利润", "线上平台建议利润", "当前库存", "配送区域", "配送时效", "最低订单数量", "可使用品牌素材",
        "可提供样品数量", "线下销售人员数量", "可投入预算", "可使用的当地公司主体", "TikTok/Facebook/Instagram/YouTube账号情况",
        "平台书面许可或广告账户资质"
    ]
    add_callout(doc, "状态与回复格式", "以上均为 NEEDS_VERIFY。缺失不会阻断线索研究，但会阻断正式报价、利润评分、上架承诺和公开投放。请按每个SKU一行回复，并注明价格币种/税费/有效期、库存批次、配送城市/时效/最低订单；平台资质须附书面文件或账户截图。", PALE_GOLD)
    add_table(doc, ["内部数据", "当前状态", "需要谁提供", "最晚节点", "不提供的影响"], [
        [f, "NEEDS_VERIFY", "国代/品牌方", "首次对外报价或送样前", "不能给出正式商务条件或商业吸引力分"] for f in fields
    ], [1.7, 1.1, 1.2, 1.25, 1.25], 8.3)
    save(doc, "00_待合作方补充的内部数据清单.docx")


def doc_impact():
    doc = new_doc("执行前影响面与阻断检查", "目录、商业数据、平台法律、文化敏感性与来源质量检查", "guide")
    checks = [
        ("同名历史文件", "KNOWN", "当前项目目录无历史交付文件，仅有研究与构建中间件", "新建版本化输出目录，不覆盖"),
        ("关键内部数据", "NEEDS_VERIFY", "SKU、价格、利润、库存、样品、预算、配送SLA缺失", "继续研究；报价和ROI阻断"),
        ("尼泊尔酒类公开广告", "BLOCKED", "Public Health Service Act 2018 s45覆盖alcohol advertising", "默认方案C；取得当地书面法律意见后再评估"),
        ("TikTok平台访问", "KNOWN", "平台在尼泊尔可访问；但平台可用不等于酒类推广可用", "执行前复查平台状态"),
        ("Nepali语言", "NEEDS_VERIFY", "当前仅能生成草稿", "母语者+合规负责人双审"),
        ("政治/宗教/民族表达", "NEEDS_VERIFY", "尼泊尔内部文化差异显著", "不使用国家关系、宗教仪式或民族符号做未经审核的促销"),
        ("渠道来源合法性", "KNOWN", "限定公开网页、地图、官网、社媒和合法名录", "不绕过登录/验证码，不收集非公开个人信息"),
        ("传统酒=接受白酒", "INFERRED风险", "传统饮酒文化只能证明酒精消费语境，不能证明汾酒需求", "所有搭配和口感主张必须品鉴"),
        ("中尼关系=喜欢中国品牌", "INFERRED风险", "机构和旅游往来不等于消费者偏好", "文化只做辅助桥接"),
        ("单一低质来源", "NEEDS_VERIFY", "长尾线索可能仅有一个公开来源", "高优先级线索两源交叉，未达标不进入人工"),
    ]
    add_table(doc, ["检查项", "状态", "证据/现状", "处理"], checks, [1.25, 1.05, 2.6, 1.6], 8.2)
    doc.add_heading("允许继续的范围", level=1)
    add_bullets(doc, ["公开市场、城市、文化、竞争和渠道发现", "线索去重、公开信息验证和评分", "商务话术草稿、CRM和试销SOP", "持牌场所内的品鉴/培训方案设计（实际执行须当地确认）"])
    doc.add_heading("必须停止的动作", level=1)
    add_bullets(doc, ["未获得书面确认前的公开酒类广告、付费投放和网红推广", "未经母语审核的Nepali公开文案", "虚构价格、利润、库存、联系人或平台许可", "自动拨号、冒充真人、群发骚扰或收集非公开个人信息"])
    add_sources(doc, [s for s in ALL_SOURCES if str(s.get("id", "")).endswith(("001", "002", "003", "004", "005")) or "Public Health" in str(s.get("title", ""))][:10])
    save(doc, "01_执行前影响面与阻断检查.docx")


def doc_market():
    doc = new_doc("汾酒尼泊尔深度市场调研报告", "城市优先级、竞争格局、渠道入口与可验证的市场进入假设", "report")
    c = root.get("core_conclusions", {})
    add_callout(doc, "决策摘要", "先集中 Kathmandu Valley，用持牌零售/电商、酒店餐饮宴会、中餐/亚洲餐与中尼商务网络做B2B试销；B2C只承担反馈与动销验证。Pokhara在加德满都形成补货后再进入。")
    doc.add_heading("1. 市场进入逻辑", level=1)
    add_bullets(doc, [
        "[KNOWN] 国代已有合法供货主体、库存、仓配与履约；本项目不重新找国代。",
        "[KNOWN] 公开市场存在国内主流威士忌、进口威士忌、伏特加/朗姆、啤酒/葡萄酒/RTD和传统酒等多层替代。",
        "[INFERRED] 汾酒不宜用‘更高端的中国烈酒’正面对打；第一阶段卖点是可解释的小杯体验、食物场景、礼赠与可量化试销支持。",
        "[BLOCKED] 未确认SKU/价格/利润前，不对消费者可接受价、渠道利润或销量做确定预测。",
    ])
    doc.add_heading("2. 城市评分与优先级", level=1)
    city_rows = []
    city_source = channels.get("city_scoring") or root.get("city_scores", [])
    for x in city_source:
        city_rows.append([x.get("city"), x.get("total"), x.get("confidence"), x.get("status") or x.get("label"), x.get("entry_condition"), x.get("caution") or x.get("note")])
    add_table(doc, ["城市", "总分/100", "置信度", "标签", "进入条件", "说明"], city_rows, [1.15, 0.65, 0.65, 0.75, 1.7, 1.6], 7.8)
    component_labels = {
        "consumption_capacity": ("消费能力", 15), "hospitality_density": ("酒店餐饮密度", 15),
        "imported_spirits_base": ("进口烈酒基础", 15), "business_tourism_banquets": ("商务/旅游/宴会", 10),
        "distributor_retail_resources": ("经销/零售资源", 10), "china_asia_dining": ("中国/亚洲餐饮", 10),
        "digital_reach": ("数字触达", 10), "competitive_gap": ("竞争空档", 10),
        "execution_delivery": ("执行与配送", 5),
    }
    for city in city_source:
        doc.add_heading(f"{city.get('city')}｜九维评分明细", level=2)
        details = []
        for key, (label, weight) in component_labels.items():
            details.append([label, f"{city.get('score_components', {}).get(key, 0)}/{weight}", "COMPUTED：公开代理分；不是酒类销量或汾酒接受度"])
        add_table(doc, ["维度", "得分/权重", "解释"], details, [1.55, 1.0, 3.95], 8.0)
        evidence_rows = []
        for item in as_list(city.get("raw_evidence")):
            if isinstance(item, dict):
                evidence_rows.append([
                    item.get("metric"), shown(item.get("value")), item.get("unit"),
                    item.get("status"), humanize(item.get("source_id") or item.get("source_ids")),
                ])
        add_table(doc, ["公开证据代理", "值", "单位", "状态", "来源"], evidence_rows, [2.55, 0.85, 1.05, 0.9, 1.15], 7.8)
    score_text = "、".join(f"{x.get('city')} {x.get('total')}" for x in city_source)
    add_callout(doc, "评分解释", f"城市分是公开证据代理的 COMPUTED 决策工具，不是城市级酒类销售统计。当前为 {score_text}，需用前30天真实回复、上架和补货数据修正。", PALE_GOLD)
    doc.add_heading("3. 竞争格局", level=1)
    rows = [[x.get("group"), x.get("examples"), x.get("threat"), x.get("response")] for x in root.get("competitor_groups", [])]
    add_table(doc, ["竞争/替代组", "例子", "威胁", "汾酒应对"], rows, [1.25, 1.8, 0.75, 2.7], 8.0)
    doc.add_heading("4. 第一批渠道与客户", level=1)
    add_table(doc, ["优先", "渠道", "为什么先做", "最小验证动作", "淘汰信号"], [
        [1, "持牌精品酒门店/线上酒类配送", "进口烈酒货架、公开联系方式、可看真实销量", "确认上架条件→3瓶/小批试销→7/14/30天库存", "不回传销量、乱价、只要独家"],
        [2, "中高端酒店/餐厅/宴会", "成人餐饮场景、品鉴和店员推荐可控", "F&B/采购会面→盲测→单一场景菜单试销", "无负责人、员工不培训、只摆不卖"],
        [3, "中餐/亚洲餐与中尼商务场景", "白酒教育距离较短、商务宴请清晰", "饭局/包间小杯品鉴→团购/礼赠询盘", "只依赖中国客群、无尼泊尔消费者反馈"],
        [4, "酒吧/调酒师", "可做低剂量试饮与口味教育", "封闭行业品鉴；公开内容另走合规", "要求违规公开推广或受众年轻化"],
        [5, "婚礼/企业礼赠", "订单价值高但强依赖价格、包装和合规", "先收采购条件和样品反馈", "价格/SKU未确认仍要求报价"],
    ], [0.55, 1.4, 1.55, 1.75, 1.25], 7.8)
    doc.add_heading("5. 数字与电商渠道", level=1)
    for platform in channels.get("ecommerce_and_delivery_status", []):
        doc.add_heading(str(platform.get("platform")), level=2)
        add_table(doc, ["字段", "当前结论", "状态/动作"], [
            ["运营与覆盖", humanize(platform.get("operating_status")), humanize(platform.get("coverage"))],
            ["酒类/进口烈酒", humanize(platform.get("sells_alcohol")), humanize(platform.get("imported_spirits"))],
            ["入驻主体/许可证", shown(platform.get("onboarding_entity")), shown(platform.get("licenses"))],
            ["费用/抽佣", shown(platform.get("fees_commission")), "NEEDS_VERIFY：书面询价"],
            ["配送/年龄验证", shown(platform.get("delivery")), shown(platform.get("age_verification"))],
            ["内容展示/价格促销/酒类广告", "BLOCKED：未取得当地及平台书面许可", "不可从商品在售反推广告获准"],
            ["联系与优先级", shown(platform.get("contact")), humanize(platform.get("priority"))],
            ["下一步", humanize(platform.get("next_step")), f"证据：{humanize(platform.get('source_ids'))}"],
        ], [1.55, 3.0, 1.95], 8.0)
    doc.add_heading("6. 线索数据库现状", level=1)
    verified = sum(1 for x in LEADS if str(x.get("current_status") or x.get("status") or x.get("verification_status") or "").lower() in ("已验证", "verified"))
    dual = sum(1 for x in LEADS if str(x.get("priority", "")).lower() == "high")
    add_table(doc, ["指标", "当前值", "解释"], [
        ["真实公开线索", len(LEADS), "未用虚假企业补足数量；缺失字段保留null/待验证"],
        ["状态已验证", verified, "按研究文件显式状态统计"],
        ["独立域名双源高优线索", dual, "两个独立域名且评分达标；仍需人工核验企业和负责人"],
        ["目标", "300+", "若当前不足，继续按区域/类别迭代，不虚构"],
    ], [1.5, 1.2, 3.8])
    doc.add_heading("7. 竞争策略：窄滩头而非全市场", level=1)
    add_bullets(doc, [
        "Eliminate：大历史课、政治关系、全国TVC、过早独家。",
        "Reduce：首单规模、一次性SKU数量、公开传播依赖。",
        "Raise：品鉴教育、店员脚本、7/14/30天回访、证据和数据回传。",
        "Create：‘小杯慢饮+本地食物验证+补货闸门’的90天试销包。",
    ])
    add_sources(doc, ALL_SOURCES)
    save(doc, "02_汾酒尼泊尔深度市场调研报告.docx")


def culture_items(key_candidates, fallback):
    value = first_present(culture, key_candidates, fallback)
    return as_list(value)


def doc_culture():
    doc = new_doc("尼泊尔生活习惯与消费文化报告", "人群、场景、节庆、传统酒与食物搭配的证据化执行指南", "report")
    add_callout(doc, "核心判断", "尼泊尔不是单一饮酒文化。民族/宗教、性别、城市、代际和家庭规范均可能改变酒的接受边界；传统酒的存在不等于消费者会接受中国白酒。")
    doc.add_heading("1. 人群分层", level=1)
    culture_population = culture.get("population_and_culture", {})
    segments = as_list(culture_population.get("operational_segments")) or culture_items(["consumer_segments", "population_segments", "segments"], [
        {"segment": "Kathmandu城市中高收入成年人", "scene": "商务宴请、餐厅、礼赠", "barrier": "与威士忌比较、对白酒陌生", "test": "小杯盲测+价格后测"},
        {"segment": "年轻城市白领", "scene": "朋友聚餐、酒吧、内容消费", "barrier": "刺激度、品牌距离、年龄合规", "test": "成年样本、低量试饮、无功效表达"},
        {"segment": "酒店/旅游消费者", "scene": "餐饮、度假、礼品", "barrier": "短停留、SKU解释成本", "test": "菜单单一场景试销"},
        {"segment": "家庭/婚礼采购者", "scene": "团聚、宴会、礼赠", "barrier": "宗教/家庭禁忌、价格/包装未知", "test": "采购访谈+家庭边界筛选"},
        {"segment": "中尼商务人士/中国游客", "scene": "商务宴请、中餐", "barrier": "容易形成狭窄华人专用定位", "test": "作为早期滩头，不替代尼泊尔消费者验证"},
    ])
    seg_rows = []
    for x in segments:
        if isinstance(x, dict):
            seg_rows.append([
                x.get("segment") or x.get("group") or x.get("name"),
                x.get("scene") or x.get("occasions") or x.get("role"),
                x.get("barrier") or x.get("risk") or x.get("allowed_now"),
                x.get("test") or x.get("validation") or "；".join(str(v) for v in as_list(x.get("proof_needed"))),
            ])
    required_segments = [
        ("Kathmandu城市中高收入成年人", "餐饮/礼赠", "价格和品类陌生", "成年小杯盲测"),
        ("年轻城市白领", "朋友聚餐/内容", "年龄与年轻化风险", "21+样本；演员25+"),
        ("夜生活/酒吧成年人", "持牌酒吧", "豪饮/酒驾", "调酒师封闭测试"),
        ("家庭型消费者", "家庭聚餐", "家庭/宗教拒酒", "先访谈边界"),
        ("婚礼/宴会采购者", "宴会/礼赠", "未成年人共场", "B2B询盘后审批"),
        ("商务宴请决策者", "公司宴请", "成功暗示/劝酒", "采购访谈"),
        ("中产家庭采购者", "居家/礼赠", "价格、包装未知", "概念与价格后测"),
        ("高收入群体", "高端餐饮/礼赠", "不可用收入推定偏好", "现有烈酒行为验证"),
        ("中国驻尼人员", "中餐/商务", "不能代表尼泊尔需求", "窄滩头测试"),
        ("中国游客", "酒店/旅游", "短停留与安全", "酒店采购数据"),
        ("中尼贸易相关人士", "商会/商务活动", "政治化风险", "一对一B2B"),
        ("礼赠采购者", "企业/节庆礼赠", "合规与SKU缺失", "采购条件表"),
        ("酒店/餐饮采购与服务团队", "菜单/培训/库存", "牌照与员工执行", "采购+店员SOP"),
    ]
    existing = " ".join(str(r[0]) for r in seg_rows)
    seg_rows.extend([list(r) for r in required_segments if r[0] not in existing])
    add_table(doc, ["人群", "高可能场景", "主要阻力/边界", "最小验证"], seg_rows, [1.55, 1.65, 1.7, 1.6], 8.0)
    doc.add_heading("2. 生活与消费场景", level=1)
    scene_rows = [
        ["家庭聚餐/朋友聚会", "分享与小杯尝试", "先确认家庭酒精边界，不把团聚自动等同饮酒", "愿意二次尝试、推荐、购买阻力"],
        ["商务宴请/中尼活动", "关系维护与解释型品鉴", "不使用政治口号，不暗示事业成功", "会后询盘、样品、试销"],
        ["酒店/高端餐饮", "菜单搭配、服务员推荐", "持牌场所和年龄控制", "杯售/瓶售、推荐率、补货"],
        ["酒吧夜生活", "调酒师教育、成年受众", "不能表现豪饮、醉酒或年轻化", "合规品鉴反馈，不以播放量为主"],
        ["婚礼/企业礼赠", "订单大、礼盒场景", "宗教/家庭差异、价格包装待确认", "采购意愿、可接受价、复购周期"],
        ["居家/即时配送", "方便复购", "年龄验证、配送和广告边界", "订单、重复购买、退款/投诉"],
        ["城市白领社交", "下班后小型聚会", "成年定向，不做压力饮酒", "首次尝试/再饮"],
        ["节庆团聚", "备货与礼赠研究", "不把节庆默认等同饮酒", "采购访谈"],
        ["企业礼赠", "批量订单", "对象接受酒类、税务与包装待确认", "询价/样品/转化"],
        ["中尼机构活动", "品类教育", "不使用政治背书", "B2B询盘"],
        ["中国游客服务", "熟悉品类", "不替代本地验证", "酒店点单/反馈"],
        ["旅游度假场景", "酒店餐饮", "户外/高海拔前后不促饮", "菜单销售"],
        ["厨师/调酒师培训", "建立服务方法", "封闭专业培训", "推荐正确率"],
        ["门店试饮", "降低首次认知门槛", "持牌、成年、定量、自愿", "试饮到购买"],
        ["宴会桌餐", "小杯配餐", "成年分桌与场所批准", "桌均销量/投诉"],
        ["线上查询到店", "门店定位", "仅书面允许后开放", "地图/到店/订单"],
    ]
    add_table(doc, ["场景", "价值假设", "合规/文化边界", "验证指标"], scene_rows, [1.25, 1.65, 2.0, 1.6], 8.0)
    doc.add_heading("3. 年度执行日历", level=1)
    festivals = first_present(culture, ["festival_and_season_calendar", "festival_calendar", "festivals", "annual_calendar"], [])
    rows = []
    if festivals:
        for x in as_list(festivals):
            if isinstance(x, dict):
                rows.append([x.get("window") or x.get("festival") or x.get("name"), x.get("period") or x.get("date") or "每年复核", x.get("participants") or x.get("audience") or x.get("consumer_context"), x.get("prep") or x.get("habit") or x.get("occasion") or x.get("consumer_context"), x.get("opportunity") or x.get("fit") or x.get("recommendation"), x.get("risk")])
    if not rows:
        rows = [
            ["Nepali New Year", "Apr；每年核对官方日历", "城市家庭/企业", "节庆与商务问候", "B2B礼赠候选", "价格/包装/广告待确认"],
            ["旅游春季", "Feb-Apr", "国际/国内旅游者", "酒店餐饮与旅行", "酒店/餐厅试销", "短停留、季节波动"],
            ["Dashain", "2026官方连续假期约Oct 17-23", "广泛家庭，具体边界不同", "团聚、礼物、祝福、仪式", "采购访谈；不默认适酒", "宗教/家庭禁忌"],
            ["Tihar", "2026约Nov 8-12", "家庭/兄弟姐妹/商户", "节灯与关系仪式", "礼赠候选", "不可商业化亵渎仪式"],
            ["旅游秋季", "Sep-Nov", "国际旅游/商务", "酒店、餐饮、度假", "酒店/餐厅", "与节庆资源冲突"],
            ["婚礼旺季", "按当年auspicious dates", "家庭/宴会采购", "宴会与礼赠", "B2B宴会渠道", "必须逐年当地确认"],
            ["中国春节/中秋/国庆", "按公历/农历", "中国人员与中尼机构", "商务/文化活动", "窄众B2B切入口", "不得等同尼泊尔大众需求"],
        ]
    supplemental_festivals = [
        ["旅游春季", "约3-5月；逐年核对", "酒店/游客/餐厅", "提前6-8周核库存与安全", "B2B菜单研究", "户外/高海拔禁促饮"],
        ["旅游秋季", "约10-11月；逐年核对", "酒店/游客/宴会", "与节庆物流联排", "酒店试销", "季节与节庆重叠"],
        ["企业年度礼赠", "公司财年/节前；待访谈", "企业采购/员工/客户", "先核对象与税务", "私域B2B", "不得公开节庆促销"],
        ["中尼商务文化活动", "按商协会日历", "企业/机构成年人", "逐场核许可与参与者", "专业品鉴/礼赠询盘", "禁止政治背书"],
        ["其他地方节庆", "按城市/社群日历", "本地社群", "由当地顾问逐项核对", "默认不营销", "宗教/民族挪用"],
    ]
    existing_festivals = " ".join(str(r[0]) for r in rows)
    rows.extend([r for r in supplemental_festivals if r[0] not in existing_festivals])
    add_table(doc, ["周期", "时间", "消费者/运营语境", "准备动作", "建议", "风险"], rows, [1.0, 1.0, 1.3, 1.35, 1.25, 1.1], 7.7)
    doc.add_heading("4. 传统酒语境", level=1)
    add_table(doc, ["酒饮", "已知文化/工艺事实", "与汾酒关系", "禁止说法"], [
        ["Aaila / Raksi", "蒸馏型传统酒；不同社群、地区和仪式角色不同", "可帮助解释‘蒸馏酒并非完全陌生’，口味不能类推", "‘就是尼泊尔版白酒’或‘汾酒更高级’"],
        ["Chhyang / Jaad", "发酵型谷物酒，酒精度和饮用方式不同", "只作为谷物发酵文化背景", "直接做风味/强度等同"],
        ["Tongba", "发酵小米饮用传统及器具/热水饮法", "可启发仪式与慢饮研究", "把器具挪用成品牌道具"],
    ], [1.05, 2.25, 1.8, 1.4], 8.0)
    doc.add_heading("5. 食物搭配候选（全部需真实品鉴）", level=1)
    pairings = first_present(culture, ["food_pairing_candidates", "food_pairings", "pairing_candidates"], [])
    rows = []
    for x in as_list(pairings):
        if isinstance(x, dict):
            rows.append([x.get("food") or x.get("dish"), x.get("hypothesis") or x.get("rationale"), x.get("test_serve") or x.get("serving") or x.get("format"), x.get("obstruction") or x.get("resistance") or x.get("barrier"), x.get("decision_rule") or x.get("explanation") or x.get("message"), x.get("tag") or x.get("status") or "NEEDS_VERIFY"])
    if not rows:
        for food, rationale in (("Momo", "脂香/蘸料可能与清香形成对比"), ("Sekuwa", "炭烤和肉香可能承接烈酒"), ("Newari cuisine", "本地宴饮语境但菜式/社群差异大"), ("Thakali cuisine", "复合餐盘需逐道测试"), ("中餐/烧烤", "白酒教育距离较短"), ("酒吧小食", "低量尝试和调饮可测试")):
            rows.append([food, rationale, "15-20ml小杯；纯饮/加水/加冰分组", "香气陌生、入口刺激", "先闻香、小口、配食，不劝饮", "NEEDS_VERIFY"])
    supplemental_pairings = [
        ("其他本地烤肉", "脂香/炭烤候选，按肉类与辣度分组"), ("酒吧小食", "低量纯饮/稀释对照"),
        ("传统酒同场描述性对照", "只比较可感知描述，不做优劣与替代"), ("企业礼赠开箱/无餐", "验证包装、解释成本与拒绝原因"),
    ]
    existing_pairings = " ".join(str(r[0]) for r in rows)
    for food, rationale in supplemental_pairings:
        if food not in existing_pairings:
            rows.append([food, rationale, "SKU待确认；5-10ml闻香/小口；配水", "刺激、文化误读或无餐承接", "由当地成年人和专业人员记录；未达阈值即淘汰", "NEEDS_VERIFY"])
    add_table(doc, ["食物", "理论依据", "服务方式候选", "可能阻力", "讲解", "状态"], rows, [1.0, 1.45, 1.4, 1.15, 1.1, 0.9], 7.5)
    add_sources(doc, [s for s in ALL_SOURCES if any(k in str(s.get("title", "")).lower() for k in ("alcohol", "dashain", "tihar", "tourism", "ferment", "dhs", "holiday", "climate"))][:24])
    save(doc, "03_尼泊尔生活习惯与消费文化报告.docx")


def doc_bridge():
    doc = new_doc("中国文化植入与中尼文化桥接矩阵", "让产品先进入尼泊尔真实生活，再解释其中国来源", "guide")
    add_callout(doc, "内容原则", "建议起始配比：55%尼泊尔真实生活、20%产品饮用教育、15%中尼共同文化、8%汾酒工艺与来源、2%纯形象；8周后按合规且与询盘/销量相关的数据调整。")
    bridges = first_present(culture, ["culture_bridge_matrix", "bridge_matrix", "cultural_bridges"], [])
    rows = []
    for x in as_list(bridges):
        if isinstance(x, dict):
            evidence = x.get("evidence_grade") or x.get("grade") or x.get("evidence_ids")
            if x.get("tag"):
                evidence = f"{x.get('tag')}｜{humanize(evidence)}" if evidence else x.get("tag")
            rows.append([
                x.get("nepal_fact") or x.get("local_fact"),
                x.get("china_link") or x.get("chinese_bridge") or x.get("chinese_culture"),
                x.get("fenjiu_expression") or x.get("expression"),
                x.get("scene") or x.get("best_scene") or x.get("use_scene"),
                x.get("risk"),
                humanize(evidence),
            ])
    if not rows:
        defaults = [
            ("待客与分享", "中国宴请中的共同品尝", "小杯、先闻香、小口，不劝饮", "持牌餐厅品鉴", "不能暗示社交成功", "B/NEEDS_VERIFY"),
            ("家庭团聚", "中国节庆团聚", "只讲‘共享时刻’，不强行对应节日", "家庭采购访谈", "部分家庭禁酒", "B/NEEDS_VERIFY"),
            ("尊重长辈", "敬意与礼物", "礼盒和递送礼仪待当地测试", "企业/家庭礼赠", "避免等级化和劝酒", "C/NEEDS_VERIFY"),
            ("手工与传统", "酿造工艺传承", "以可验证工艺事实解释香气", "店员培训", "禁止空泛‘几千年’", "B"),
            ("食物与酒", "餐酒搭配", "与Momo/Sekuwa等做盲测", "餐厅", "不得先宣布绝配", "INFERRED"),
            ("山地与自然", "山西产地/自然意象", "作为视觉背景，不做地缘类比", "贸易母片", "避免虚假风土功效", "C"),
            ("商务关系维护", "中国商务宴请", "一对一解释产品和试销", "中尼商会/企业", "不使用政治口号", "B"),
        ]
        rows = [list(x) for x in defaults]
    add_table(doc, ["尼泊尔本地事实", "可连接中国文化", "汾酒表达", "场景", "风险", "证据"], rows, [1.15, 1.2, 1.45, 1.0, 1.1, 0.6], 7.5)
    doc.add_heading("禁止成为主叙事的表达", level=1)
    add_bullets(doc, ["‘中国名酒’作为唯一购买理由", "‘几千年文化’但无可核验事实", "‘大国品牌’、‘一带一路’或政治关系换消费认同", "‘东方神秘’和对尼泊尔传统的挪用", "把Aaila/Raksi称作低端或原始", "把中国游客存在等同于尼泊尔消费者偏好"])
    doc.add_heading("8周内容学习规则", level=1)
    add_numbered(doc, ["每条内容先写清本地生活事实和证据。", "检查是否含酒类品牌露出；若有，进入合规闸门。", "仅在持牌、成年、获授权的场景做真实测试。", "用评论质量、询盘、门店查询和销量变化评估，而非播放量。", "若出现文化误读、宗教/民族争议或平台警告，立即下线并人工处理。"])
    add_sources(doc, [s for s in ALL_SOURCES if any(k in str(s.get("title", "")) for k in ("China", "Chinese", "Confucius", "Tourism")) or "China" in str(s.get("conclusion", ""))][:10])
    save(doc, "04_中国文化植入与中尼文化桥接矩阵.docx")


def execution_value(keys, fallback):
    return first_present(execution, keys, fallback)


def doc_90day():
    doc = new_doc("90天B2B与B2C市场执行手册", "从公开线索到品鉴、试销、上架、动销和补货", "guide")
    add_callout(doc, "执行重心", "建议前4周 B2B 70%、B2C反馈20%、内容/素材10%。只有当合规解除且B2B已有可购买点，才提高内容占比。")
    doc.add_heading("1. 标准推进路径", level=1)
    path = ["发现企业", "验证企业", "找到公开负责人/路由", "评分", "首次接触", "需求判断", "简版资料", "电话/线上沟通", "预约", "现场品鉴", "90天试销", "上架", "店员培训", "7天回访", "14天动销", "30天库存", "补货/淘汰"]
    add_table(doc, ["步骤", "AI完成", "人工完成", "退出条件"], [[i+1, p, "数据整理、个性化、提醒" if i < 8 else "记录/分析", "真实发送/电话/见面/品鉴/谈判/履约" if i >= 4 else "仅在需要时", "身份不明、拒绝、违规、无执行或无补货"] for i, p in enumerate(path)], [0.5, 1.45, 2.05, 2.5], 7.6)
    doc.add_heading("2. 每个平台怎么找客户", level=1)
    add_table(doc, ["平台", "搜索/发现", "验证", "首次动作", "规则"], [
        ["Google Maps/Search", "area + liquor store/hotel/restaurant/banquet/imported spirits", "地址、营业、评论、电话、官网/第二来源", "进入CRM，不直接群发", "不绕过验证码或高频抓取"],
        ["Facebook/Messenger", "企业Page、近期帖、About/contact", "近90天活跃+地址+第二来源", "一次个性化商务消息", "D1/D3/D7/D14后停止"],
        ["LinkedIn", "Procurement/F&B Manager/Owner/Corporate Purchase + Nepal", "公司任职、资历、共同联系人", "连接请求或InMail", "不收集非公开个人信息"],
        ["电话", "公开企业电话", "先确认企业和负责人路由", "30秒开场，约下一步", "不得AI冒充真人自动拨打"],
        ["Viber/WhatsApp", "仅用公开企业号码或已交换号码", "确认同意与关系", "预约/样品/订单跟进", "Viber business强调verified/opt-in，不作冷群发"],
        ["协会/名录", "酒店、餐饮、商会、旅游、双边商会", "会员/活动公开页面", "申请活动或引荐", "不是成交捷径，酒类活动仍需合规"],
    ], [1.05, 1.65, 1.55, 1.25, 1.0], 7.7)
    doc.add_heading("3. 12周执行周期", level=1)
    weeks = execution_value(["twelve_week_plan", "execution_weeks", "week_plan"], [])
    rows = []
    for x in as_list(weeks):
        if isinstance(x, dict):
            rows.append([x.get("week") or x.get("period"), x.get("focus") or x.get("objective") or x.get("goal"), humanize(x.get("actions") or x.get("deliverables")), "；".join(humanize(v) for v in (x.get("gate") or x.get("done_when"), x.get("kpi")) if v), x.get("owner") or "AI+人工"])
    if not rows:
        rows = [
            ["W1", "输入与合规", "内部数据、方案C、城市初分、文件/CRM", "知道可做/不可做", "AI+国代"],
            ["W2", "文化与150线索", "文化、节庆、桥接、区域图、150线索", "前50条可验证", "AI"],
            ["W3", "300线索与评分", "去重、双源、Top50、话术", "不虚构且有下一步", "AI"],
            ["W4", "商务/试销包", "资料、品鉴、培训、平台申请、内容脚本", "报价字段补齐或标阻断", "AI+人工"],
            ["W5-6", "第一轮接触", "接触30-50、有效回复10-20、沟通5-10、预约3-8", "拒绝原因入库", "人工发送/会面"],
            ["W7-8", "试销/反馈", "5-10终端、培训、销量、30-50反馈", "每店有7/14/30天记录", "人工+AI"],
            ["W9-10", "优化/补货", "渠道、口感、内容和话术调整", "出现首轮补货", "AI分析+人工"],
            ["W11-12", "复盘/决策", "回答渠道、场景、SKU、补货、扩城", "继续/调整/停止有证据", "决策人"],
        ]
    add_table(doc, ["周期", "目标", "动作/交付", "闸门", "负责人"], rows, [0.75, 1.05, 2.25, 1.55, 0.9], 7.5)
    doc.add_heading("4. KPI与决策阈值", level=1)
    add_table(doc, ["指标", "90天建议", "继续", "调整", "停止/暂停"], [
        ["原始/已验证/高优线索", "300+/150+/50+", "质量满足并持续转化", "不足则扩来源和关键词", "不得为数量造假"],
        ["接触/有效回复", "50-80 / 15-30", "回复率≥20%且有预约", "10-20%改客群/话术", "<10%且两轮修正无改善"],
        ["见面/品鉴", "5-12", "有真实产品反馈", "预约多到场少则改流程", "无法合法品鉴"],
        ["试销终端", "8-15", "完成培训和数据回传", "上架不推荐则强化店员", "只摆不卖且不配合"],
        ["补货终端率", "≥30%试销终端", "证明场景可复制", "15-29%再优化30天", "<15%且两轮无改善"],
        ["消费者反馈", "30-50份", "愿再饮/场景清晰", "刺激/价格阻力可修", "重大口感拒绝且无可行服务方式"],
        ["合规事件", "0", "无警告/投诉", "轻微问题立即整改", "监管/平台警告即停"],
    ], [1.35, 1.1, 1.4, 1.4, 1.25], 7.6)
    doc.add_heading("5. 未来7天动作", level=1)
    add_numbered(doc, ["国代一次性补齐SKU、价格、利润、库存、样品、配送和主体资料。", "当地律师/监管与平台书面确认酒类内容、产品列表、品鉴和B2B材料边界。", "从数据库筛Top50；仅双源且评分达标者进入人工。", "对前20家做D1个性化首触达，统一记录回复/拒绝原因。", "预约2-3场持牌场所、成年受众的小杯品鉴；准备盲测反馈表。", "对Cheers/Barmandoo/Drinks Nepal/Daraz分别询问入驻与年龄验证。", "不开公开酒类内容；先拍无公开发布的合规培训/品鉴素材用于内部复核。"])
    add_sources(doc, [s for s in ALL_SOURCES if any(k in str(s.get("title", "")) for k in ("Viber", "LinkedIn", "Digital 2025", "Video play"))])
    save(doc, "05_90天B2B与B2C市场执行手册.docx")


def doc_distributor():
    doc = new_doc("经销商发展与试销机制", "不先给独家：用90天实际进货、终端销售、数据回传和补货升级合作", "guide")
    add_callout(doc, "硬原则", "独家权不是见面奖励。只有真实进货、真实终端、消费者销售、补货、价格纪律和数据回传同时成立，才讨论更深合作。")
    doc.add_heading("1. 四阶段机制", level=1)
    add_table(doc, ["阶段", "验证内容", "输出", "升级门槛"], [
        ["线索", "真实经营、覆盖区域、客户、进口酒、负责人", "证据和评分", "身份可验证、分数≥50"],
        ["意向", "愿了解/品鉴/试销、覆盖终端、利润关注、支持需求", "会议纪要", "愿意给目标终端和时间表"],
        ["90天试销", "SKU、首批量、目标店、数据、价格、培训、补货", "试销单+周数据", "发生消费者销售并补货"],
        ["升级合作", "销售团队、覆盖、价格纪律、回款、数据", "扩区/扩量建议", "多周期稳定表现；独家另行审议"],
    ], [1.0, 2.25, 1.35, 1.9], 8.0)
    doc.add_heading("2. 90天试销条款清单", level=1)
    add_bullets(doc, ["区域、SKU、首批量、目标终端和上架时限明确", "不立即授予长期独家；任何排他都需书面、期限、销量与退出条款", "价格边界、促销审批、串货/乱价处理", "7/14/30/60/90天库存、销量、消费者反馈和补货记录", "店员培训、品鉴方式、年龄与责任饮酒规则", "样品、破损、退换、付款、配送和投诉责任", "停止条件：不回传、乱价、无真实终端、要求违规宣传、无补货"])
    doc.add_heading("3. 经销商评分（100分）", level=1)
    add_table(doc, ["维度", "权重", "证据", "0分条件"], [
        ["目标客户匹配", 15, "客户/终端结构", "无目标客群"], ["渠道覆盖", 15, "覆盖清单和近期开店/销售", "只口头宣称"],
        ["销售团队执行", 15, "人员、路线、周计划", "无销售人员"], ["试销意愿", 15, "接受小范围、数据和补货闸门", "只要独家"],
        ["负责人可触达", 10, "正式会面/授权", "身份不明"], ["信誉/回款", 10, "工商/贸易和参考", "重大风险"],
        ["终端动销", 10, "培训/陈列/回访能力", "只压货"], ["数据回传", 5, "同意模板和频率", "拒绝"], ["价格纪律", 5, "书面承诺/历史", "要求乱价"],
    ], [2.0, 0.8, 2.2, 1.5], 8.0)
    doc.add_heading("4. 品鉴与店员训练", level=1)
    add_numbered(doc, ["确认参与者均为成年人且场所/活动允许。", "先记录平时饮酒与对白酒认知。", "以15-20ml小杯为候选，纯饮/加水/加冰分组；实际SKU服务方式由品鉴决定。", "先闻香、小口、配食，不劝饮、不做健康/成功暗示。", "记录香气、刺激、回味、场景、可接受价格、再次饮用和购买阻力。", "店员只使用通过合规和母语审核的话术。"])
    doc.add_heading("5. 经销商资料包内容", level=1)
    add_bullets(doc, ["1页快速介绍（谁买/何时喝/为什么试）", "SKU/包装/价格/利润（待合作方补全）", "品牌和工艺可验证事实", "目标人群和场景", "90天试销、品鉴和店员培训", "消费者异议回答", "7/14/30天反馈与补货表"])
    doc.add_heading("6. 一页快速介绍模板", level=1)
    add_table(doc, ["模块", "可直接使用的内容"], [
        ["合作定位", "汾酒尼泊尔国代支持的90天、小范围、非长期独家试销"],
        ["目标场景", "持牌零售/配送、酒店餐饮宴会、中餐/亚洲餐与中尼商务"],
        ["首批产品", "NEEDS_VERIFY：SKU、ABV、包装、建议零售价、供价、MOQ、库存"],
        ["我们提供", "样品政策待确认；店员培训；品鉴SOP；7/14/30天复盘"],
        ["合作方提供", "合法资质、目标终端、员工执行、库存/销量/反馈数据、价格纪律"],
        ["成功标准", "发生消费者销售、数据可追溯、90天内出现补货；无合规事件"],
    ], [1.55, 4.95], 8.2)
    doc.add_heading("7. 店员训练卡与异议速答", level=1)
    add_table(doc, ["顾客问题/动作", "店员回答/动作", "不得做"], [
        ["这是什么？", "这是中国清香型白酒Fenjiu；先闻香，再自愿小口，具体SKU信息以确认版为准。", "只讲历史、强迫干杯"],
        ["会不会太烈？", "可选择5-10ml闻香/小口并配水与食物；也可以拒绝。", "健康、保暖、提神声称"],
        ["怎么喝？", "按本次获批SOP：定量小杯、慢饮、配水；稀释/加冰需以品鉴结果确认。", "shot挑战、豪饮"],
        ["和威士忌一样吗？", "属于不同品类；只描述实际香气与口感，不替顾客下结论。", "贬低其他品类"],
        ["多少钱？", "NEEDS_VERIFY：只能使用国代书面价格表。", "口头乱价/未批折扣"],
    ], [1.35, 3.6, 1.55], 8.0)
    doc.add_heading("8. 7/14/30天反馈与补货表", level=1)
    add_table(doc, ["节点", "必须回传", "判断与下一步"], [
        ["D7", "到货/上架、培训人数、开瓶/杯售、顾客问题", "未上架或未培训：纠正；无执行则暂停"],
        ["D14", "销售量、库存、试饮到购买、前三拒绝原因", "调整服务/话术；不得用曝光替代销量"],
        ["D30", "累计销量、剩余库存、毛利反馈、投诉、补货意向", "补货/再测试/淘汰；写入CRM"],
    ], [0.9, 3.6, 2.0], 8.1)
    save(doc, "06_经销商发展与试销机制.docx")


def get_topics():
    raw = execution_value(["short_video_topics", "video_topics", "content_topics", "topics_30"], [])
    return [x for x in as_list(raw) if isinstance(x, dict)]


def get_scripts():
    raw = execution_value(["shoot_ready_scripts", "video_scripts", "scripts", "shootable_scripts", "scripts_12"], [])
    return [x for x in as_list(raw) if isinstance(x, dict)]


def doc_social():
    doc = new_doc("TikTok及社交媒体推广策略", "当前按方案C：平台可用不等于酒类广告或公开推广可用", "report")
    add_callout(doc, "当前选择：方案C", "尼泊尔法律对酒类广告的禁止构成首要闸门。未取得当地书面法律意见、平台书面许可和国代主体/账户审批前，不发布带品牌、购买引导、价格或饮用画面的公共推广。", PALE_RED)
    doc.add_heading("1. 合规闸门", level=1)
    add_numbered(doc, ["确认平台在尼泊尔当前可访问并已注册/合规运营。", "取得尼泊尔律师/监管对自然内容、付费广告、创作者、直播、商品页和持牌场所内容的书面意见。", "向TikTok/Meta/YouTube取得Nepal酒类行业书面资格与账户要求。", "确认主体、许可证、演员年龄、ABV/责任提示、年龄门槛和落地页。", "国代合规负责人逐条签字；任何BLOCKED项未解除，不发布。"])
    doc.add_heading("2. A/B/C条件方案", level=1)
    add_table(doc, ["方案", "条件", "允许做", "仍禁止/注意"], [
        ["A", "当地法律+平台书面允许自然和付费", "成年定向、账号、内容、创作者、付费测试、合规落地页", "未成年人、<25岁演员（TikTok政策）、豪饮、功效、奖励饮酒、灰色账户"],
        ["B", "自然内容书面允许，付费不允许", "自然内容、餐饮/品鉴教育、成年创作者、主页门店信息", "任何付费、未经批准购买引导、违规品牌露出"],
        ["C（当前）", "公开酒类推广受禁止或规则未确认", "B2B资料、店员培训、私下一对一、持牌场所内部材料（均当地确认）", "不伪装公共文化内容为酒类广告，不规避审核"],
    ], [0.85, 1.65, 2.15, 1.85], 8.0)
    doc.add_heading("2.1 三套方案的执行配置", level=2)
    add_table(doc, ["模块", "方案A｜自然+付费获准", "方案B｜仅自然获准", "方案C｜当前"], [
        ["账号与定位", "国代实名账号；21+成年餐饮/产品教育", "同左；不做付费扩量", "不运营消费者酒类推广账号"],
        ["内容矩阵", "当地生活/配餐/教育/人物/购买信息", "删除付费素材和未批购买引导", "仅内部培训、B2B一对一材料"],
        ["创作者", "25+、授权、披露、成年受众；先2-3人小测", "仅自然合作且逐条审批", "不做公开达人酒类合作"],
        ["付费定向", "仅书面许可的地域/年龄/兴趣；小预算分组", "禁止", "禁止"],
        ["转化路径", "内容→合规落地页/门店定位→询盘/购买", "主页→获批门店信息/预约", "销售人员→持牌场所/采购会面"],
        ["追踪", "UTM、门店码、询盘、品鉴、销量、补货", "同左但无paid字段", "CRM来源、预约、试销、补货"],
        ["预算闸门", "预算待国代确认；先8周小测，补货前不扩量", "只计制作/人工；待确认", "公开传播预算=0"],
        ["停机", "警告/投诉/年龄异常/无购买点即停", "同左", "任何试图伪装公开推广即停"],
    ], [1.1, 2.0, 1.7, 1.7], 7.7)
    doc.add_heading("3. 内容支柱（仅用于获批后的脚本池）", level=1)
    add_table(doc, ["支柱", "内容", "验证", "风险"], [
        ["当地生活", "成年朋友/家庭/商务/婚礼/酒店", "当地人审场景真实性", "家庭/宗教边界"],
        ["食物搭配", "Momo/Sekuwa/Newari/Thakali/中餐", "真实盲测", "不得先说绝配"],
        ["产品教育", "Fenjiu/Baijiu、闻香、小杯慢饮", "理解度与刺激反馈", "不得劝饮/健康声称"],
        ["文化桥接", "待客、分享、礼物、手工", "本地文化审核", "中国化过重/政治化"],
        ["真实人物", "厨师、酒店、调酒师、成年消费者", "授权和年龄", "年轻化、虚假背书"],
        ["购买信息", "门店/平台/品鉴", "只在书面允许时", "价格/链接/商品页限制"],
    ], [1.0, 2.15, 1.6, 1.75], 8.0)
    topics = get_topics()
    if topics:
        doc.add_heading("4. 第一批30个选题", level=1)
        rows = []
        for i, x in enumerate(topics[:30], 1):
            rows.append([i, x.get("title") or x.get("topic") or x.get("name"), x.get("pillar") or x.get("category"), x.get("hook") or x.get("first_3_seconds"), x.get("sales_scene") or x.get("scene"), x.get("risk") or x.get("compliance")])
        add_table(doc, ["#", "选题", "支柱", "前三秒", "销售场景", "风险"], rows, [0.35, 1.55, 0.85, 1.45, 1.1, 1.2], 7.3)
    scripts = get_scripts()
    if scripts:
        doc.add_heading("5. 12条可拍脚本", level=1)
        for i, x in enumerate(scripts[:12], 1):
            doc.add_heading(f"脚本 {i}｜{x.get('title') or x.get('topic') or '未命名'}", level=2)
            add_table(doc, ["字段", "内容"], [
                ["前三秒钩子", x.get("hook_first_3s") or x.get("hook") or x.get("first_3_seconds")],
                ["镜头", humanize(x.get("shot_list") or x.get("shots") or x.get("storyboard"))],
                ["演员/场景/道具", "；".join(humanize(v) for v in (x.get("actors") or x.get("cast"), x.get("scene"), x.get("props")) if v)],
                ["台词", humanize(x.get("dialogue") or x.get("lines") or x.get("script"))],
                ["English subtitle", humanize(x.get("subtitle_en") or x.get("english_subtitle") or x.get("en_subtitle"))],
                ["Nepali草稿", humanize(x.get("subtitle_ne_draft") or x.get("nepali_draft") or x.get("ne_subtitle") or "NEEDS_VERIFY：母语者审核")],
                ["发布文案/关键词", "；".join(humanize(v) for v in (x.get("caption_en") or x.get("caption"), x.get("search_keywords") or x.get("keywords")) if v)],
                ["指标", humanize(x.get("target_metrics") or x.get("metrics") or x.get("target_metric"))],
                ["销售场景", x.get("sales_scene")],
                ["风险检查", x.get("risk_check") or x.get("risk") or x.get("compliance_check") or "BLOCKED until approval"],
            ], [1.3, 5.2], 8.3)
    doc.add_heading("6. 8周测试和指标", level=1)
    add_bullets(doc, ["W1-2：内部制作6条，不公开；先做合规与本地真实性评审。", "W3-4：若书面允许，测试不同内容支柱；否则保持内部培训素材。", "W5-6：只保留与评论质量、主页/地图、询盘和门店销售相关的两个方向。", "W7-8：测试创作者/门店/线下转化；无购买点不做传播扩量。", "记录3秒/6秒/完播/时长/收藏/分享/评论质量/主页/私信/地图/询盘/品鉴/上架/销量；平台指标口径分开。"])
    add_sources(doc, [s for s in ALL_SOURCES if any(k in str(s.get("title", "")).lower() for k in ("tiktok", "meta", "facebook", "google ads", "video play", "public health"))])
    save(doc, "07_TikTok及社交媒体推广策略.docx")


def doc_compliance_checklist():
    doc = new_doc("TikTok及社交媒体合规确认清单", "可直接发给尼泊尔合作方、当地律师和平台销售的书面确认表", "guide")
    items = [
        "TikTok/Meta/Instagram/YouTube在尼泊尔当前是否正常运营且注册要求满足？",
        "Public Health Service Act 2018 s45对品牌自然内容、产品教育、B2B材料、持牌场所内部展示和商品页分别如何适用？",
        "TikTok自然内容、付费广告、创作者、直播、商品链接在Nepal对alcohol是否允许？",
        "Meta/Facebook/Instagram自然内容与付费广告对alcohol是否允许？",
        "YouTube自然内容与Google/YouTube付费广告对alcohol是否允许？",
        "允许使用的当地公司主体、许可证和广告账户审批材料是什么？",
        "最低受众年龄、演员/创作者年龄、年龄门槛和地域限制是什么？",
        "是否必须显示ABV、责任饮酒、健康警示或本地语言声明？",
        "持杯、抿饮、调酒、食物搭配、礼赠、价格、折扣和配送分别允许到什么程度？",
        "是否允许门店定位、商品目录、下单链接、即时配送和品鉴预约？",
        "创作者是否需要披露合作？哪些行业资质/合同/授权要留档？",
        "出现警告、限流、删除或投诉时，联系谁、多久停投、如何保全证据？",
    ]
    researched_questions = culture.get("partner_confirmation_questions") or culture.get("partner_compliance_questions")
    if researched_questions:
        items = []
        for q in as_list(researched_questions):
            if isinstance(q, dict):
                items.append(q.get("question") or q.get("item") or q.get("text") or str(q))
            else:
                items.append(str(q))
    add_table(doc, ["#", "待书面确认问题", "法律意见", "平台意见", "国代签字", "状态"], [[i+1, x, "", "", "", "BLOCKED"] for i, x in enumerate(items)], [0.35, 3.45, 0.8, 0.8, 0.65, 0.45], 7.5)
    doc.add_heading("发布前逐条检查", level=1)
    add_bullets(doc, ["演员/创作者年龄文件和授权已存档", "不面向或吸引未成年人", "无豪饮、醉酒、酒驾、危险、怀孕或健康/成功暗示", "无奖品/奖励饮酒、灰色账户或审核规避", "ABV/责任提示/年龄限制符合书面意见", "每个品牌露出、价格和购买引导均在批准范围", "Nepali文案已母语审核", "国代合规负责人已签字", "投诉/下线联系人已明确"])
    add_sources(doc, [s for s in ALL_SOURCES if any(k in str(s.get("title", "")).lower() for k in ("tiktok", "meta", "facebook", "google ads", "public health"))])
    save(doc, "07A_TikTok及社交媒体合规确认清单.docx")


def doc_ai_sop():
    doc = new_doc("AI与人工最小介入SOP", "AI完成标准化工作；人工只承担法律、语言和现实世界动作", "guide")
    add_callout(doc, "边界", "AI可发现、整理、去重、评分、生成和复盘，但不得代替当地律师/监管/平台批准，不得编造企业、联系人或商业条件，也不得未经授权自动外联。")
    roles = [
        ("总控 Codex", "拆解、文件、合并、冲突、状态、报告", "跨Agent冲突/完成闸门"), ("市场研究", "城市、人群、竞争、节庆、场景", "证据与假设分层"),
        ("线索挖掘", "公开搜索、字段提取、去重、缺失标记", "不收集非公开个人信息"), ("线索评分", "城市/渠道/终端/创作者评分和每日队列", "只把双源高分交人工"),
        ("外联内容", "Email/Messenger/LinkedIn/Viber/电话脚本", "不自动冒充真人"), ("内容策略", "选题、脚本、日历、复盘", "合规闸门前不发布"),
        ("合规与证据", "官方规则、来源、年龄/功效/优惠风险", "不替代律师"), ("CRM复盘", "状态、提醒、回复/预约/试销/补货率、周报", "数据缺失要报警"),
        ("Claude复核", "长文、文化、反方、人感和落地性复核", "本机CLI不可用时生成复核任务包，不假装已调用"),
    ]
    add_table(doc, ["Agent", "职责", "硬边界"], roles, [1.25, 3.45, 1.8], 8.2)
    doc.add_heading("人工介入触发条件", level=1)
    triggers = ["置信度<80%", "两个A级来源冲突", "当地法律/平台资质", "价格/账期/折扣", "区域/渠道独家", "正式商务谈判", "客户要求电话/见面", "送样", "实际品鉴", "Nepali公开发布", "宗教/民族/政治", "平台警告/封禁", "客户投诉", "高价值客户负面反馈", "企业身份无法确认"]
    add_table(doc, ["触发", "人工只做什么", "完成后交回AI"], [[t, "做决定/现实动作并记录证据", "更新状态、结构化结果、生成下一步"] for t in triggers], [2.0, 2.25, 2.25], 8.0)
    doc.add_heading("每日自动运行", level=1)
    add_numbered(doc, ["08:30校验新线索、去重和证据字段。", "09:00按评分和上次动作生成不超过20条人工队列。", "外联后由人工/授权系统写回结果。", "17:00分类回复、更新状态和下一次动作。", "每周五计算回复、预约、品鉴、试销、补货和拒绝原因。", "任何合规警告立即BLOCKED并通知人工。"])
    doc.add_heading("数据质量闸门", level=1)
    add_bullets(doc, ["企业名称+城市+类别+至少一条公开URL才入库", "高优先级需要两条独立公开来源", "决策人缺失不降低企业真实性，但不得编造", "来源访问日期超过90天或主页长期不活跃时重新验证", "同企业用标准化名称、电话、域名和地址去重", "AI输出只能进入‘待验证’，人工/公开证据才能升级‘已验证’"])
    save(doc, "08_AI与人工最小介入SOP.docx")


def doc_outreach():
    doc = new_doc("商务外联话术与跟进模板", "英文可直接个性化；Nepali为母语审核前草稿", "guide")
    add_callout(doc, "发送规则", "先验证企业，再写一条与该企业真实业务相关的首句。D1/D3/D7/D14最多四次；明确拒绝、要求停止或身份不明时立即停止。")
    templates = [
        ("Messenger/Email 首次", "English", "Hello [Name/Team], I noticed [verified business fact]. We work with Fenjiu’s authorized national distributor in Nepal and are evaluating a small, compliant trial for selected licensed outlets. This is not a request for exclusivity. Would the person responsible for imported spirits or F&B purchasing be open to a 15-minute introduction?"),
        ("D3 跟进", "English", "A quick follow-up in case the earlier note reached the wrong person. We can share a one-page trial outline covering target adult occasions, staff guidance, data tracking and a no-long-term-exclusivity pilot. Who would be the right contact?"),
        ("D7 电话开场", "English", "Hello, I’m calling regarding a business introduction sent to [channel]. May I confirm who handles imported spirits/F&B purchasing? We are proposing a small 90-day trial with staff training and clear replenishment checks, subject to your licensing and compliance requirements."),
        ("预约品鉴", "English", "Thank you for your interest. We propose a closed tasting for verified adults at a licensed venue. The aim is to compare aroma, entry intensity, aftertaste, food fit and sales objections—not to push a large opening order. Please confirm date, attendees, venue permission and any house policy."),
        ("试销提案", "English", "Based on the tasting, we suggest one SKU, a defined outlet set, staff briefing, weekly stock/sales feedback and 7/14/30-day reviews. Price, margin, MOQ and sample support will be confirmed in writing by the national distributor. No long-term exclusivity is included."),
        ("最后一次", "English", "I’ll close the loop after this message so we do not over-contact you. If a compliant imported-spirit trial becomes relevant, reply with the right contact and preferred time. Otherwise we will not follow up further."),
        ("首次信息", "Nepali draft — NEEDS_VERIFY", "नमस्कार [नाम/टोली], हामीले [प्रमाणित व्यावसायिक तथ्य] देख्यौं। नेपालमा आधिकारिक राष्ट्रिय वितरकसँग मिलेर चयनित इजाजतप्राप्त आउटलेटका लागि सानो र अनुपालनयुक्त परीक्षण सम्भावना मूल्याङ्कन गर्दैछौं। आयातित स्पिरिट वा F&B खरिद हेर्ने व्यक्तिसँग १५ मिनेटको परिचय सम्भव छ?"),
        ("停止联系", "Nepali draft — NEEDS_VERIFY", "यो अन्तिम सन्देश हो ताकि हामीले तपाईंलाई अनावश्यक रूपमा सम्पर्क नगरौं। रुचि नभए हामी थप सम्पर्क गर्ने छैनौं।")
    ]
    add_table(doc, ["场景", "语言/状态", "模板"], templates, [1.2, 1.25, 4.05], 8.0)
    doc.add_heading("LinkedIn搜索", level=1)
    add_bullets(doc, ["Titles: Owner, Founder, Managing Director, Procurement Manager, Purchasing Manager, F&B Manager, Food and Beverage Director, General Manager, Corporate Administration, Event Manager", "Boolean: (procurement OR purchasing OR \"food and beverage\" OR F&B) AND (hotel OR restaurant OR distributor OR spirits) AND Nepal", "Account-first: 先锁定公司，再找角色；优先共同联系人和最近任职变动。"])
    doc.add_heading("LinkedIn与已读未回模板", level=1)
    add_table(doc, ["场景", "English模板", "Nepali状态"], [
        ["LinkedIn连接请求", "Hello [Name], I noticed your role at [verified company]. May I connect and send a one-page outline for a small Fenjiu licensed-outlet trial?", "NEEDS_VERIFY：由母语者根据English定稿"],
        ["连接后首条", "Thanks for connecting. Who handles imported spirits/F&B purchasing? We propose a limited 90-day trial, not exclusivity.", "NEEDS_VERIFY"],
        ["D3未回", "Should this go to a purchasing or F&B colleague? A short no is fine and I will close the loop.", "NEEDS_VERIFY"],
        ["Messenger已读未回", "Could you point me to imported spirits/F&B, or reply ‘not relevant’ and I will stop?", "NEEDS_VERIFY"],
        ["D14关闭", "I’ll close this outreach now and will not follow up again without a new reason.", "NEEDS_VERIFY"],
    ], [1.25, 4.25, 1.0], 7.6)
    doc.add_heading("常见异议", level=1)
    add_table(doc, ["异议", "回答框架", "禁止"], [
        ["消费者不认识白酒", "承认陌生；用小杯、食物和低风险试销教育", "夸大历史或民族认同"],
        ["比威士忌贵/难比较", "先确认SKU/价格；比较每次服务量、场景和毛利", "虚构利润或贬低威士忌"],
        ["要独家", "先90天、终端、数据和补货；达标再谈", "为换首单立刻给独家"],
        ["刺激", "记录；测试小杯、加水/加冰/食物，不劝饮", "健康功效或强迫适应"],
        ["能否公开推广", "当前方案C；需当地和平台书面批准", "暗示可绕审核"],
    ], [1.25, 3.45, 1.8], 8.1)
    save(doc, "09_商务外联话术与跟进模板.docx")


def doc_claude_pack():
    doc = new_doc("Claude复核任务包", "当前环境未检测到可用的本地 Claude CLI；本文件用于外部复核", "guide")
    add_callout(doc, "状态", "NEEDS_VERIFY：未执行Claude复核，不能把本任务包写成‘Claude已审核’。")
    doc.add_heading("复核范围", level=1)
    add_bullets(doc, ["02 市场调研结论是否过度推断", "03/04 尼泊尔文化和中国文化桥接是否冒犯、僵硬或过度中国化", "05/06 是否真正能从线索推进到补货", "07/07A 是否把平台规则和当地法律混淆", "08 AI/人工边界是否现实", "09 外联语气是否像真实商务沟通", "全部文件是否存在伪数据、空泛建议和内部矛盾"])
    doc.add_heading("需一并提供的文件", level=1)
    add_bullets(doc, ["01-09全部DOCX（含07A）", "10-16全部XLSX", "17_机器可读取的渠道数据.json", "19_最终执行报告.docx", "说明：不得仅上传18任务包，否则无法完成事实与执行性复核"])
    doc.add_heading("给Claude的任务", level=1)
    add_para(doc, "请以反方审稿人身份审查本交付包。不要重写所有文本。逐条列出：P0法律/合规风险，P1事实或逻辑错误，P2执行性缺口，P3语言/文化问题；每条给文件名、原文摘要、问题、证据要求、最小修复。特别检查：是否把传统酒文化等同于白酒需求，是否把中尼关系等同品牌偏好，是否用播放量替代补货，是否在方案C下仍暗示公开酒类推广。")
    doc.add_heading("输出格式", level=1)
    add_table(doc, ["优先级", "文件", "问题", "证据/理由", "最小修复", "是否阻断"], [["P0/P1/P2/P3", "", "", "", "", "是/否"] for _ in range(8)], [0.65, 1.05, 1.5, 1.4, 1.4, 0.5], 7.5)
    save(doc, "18_Claude复核任务包.docx")


def doc_final():
    doc = new_doc("最终执行报告", "汾酒尼泊尔B2B、B2C深度调研与90天执行系统｜当前事实、缺口与未来7天", "report")
    city_source = channels.get("city_scoring") or root.get("city_scores", [])
    first_city = next((x for x in city_source if x.get("city") == "Kathmandu Valley"), city_source[0] if city_source else {})
    add_callout(doc, "一句话结论", "用Kathmandu Valley作为唯一首城，以持牌零售/电商、酒店餐饮宴会和中餐/中尼商务场景建立8-15家试销终端；公开社交推广当前按方案C阻断；90天以≥30%试销终端补货和零合规事件决定是否继续。")
    doc.add_heading("1. 本轮结论", level=1)
    add_table(doc, ["问题", "结论", "标签"], [
        ["第一城市", "Kathmandu Valley", f"COMPUTED {first_city.get('total', 'NEEDS_VERIFY')}/100"],
        ["前三渠道", "持牌精品酒门店/配送；酒店餐饮宴会；中餐/亚洲餐与中尼商务", "INFERRED"],
        ["最快找客户", "Google Maps/Search发现；官网/Facebook验证；电话/Messenger/Viber路由；LinkedIn找决策角色", "INFERRED"],
        ["最适合消费者场景", "持牌场所的小杯食物品鉴、商务宴请、价格/SKU确认后的礼赠", "INFERRED/NEEDS_VERIFY"],
        ["TikTok", "当前方案C；平台可用但酒类公开推广阻断", "KNOWN/BLOCKED"],
        ["中国文化", "先尼泊尔生活事实，再饮用教育和共同文化，最后才是品牌工艺", "INFERRED"],
        ["三大困难", "合规；商业数据缺失；首次口感与无补货", "KNOWN/INFERRED"],
        ["90天要验证", "一个可复制的合法场景、8-15家试销、≥30%补货、零合规事件", "COMPUTED/NEEDS_VERIFY"],
    ], [1.3, 4.4, 0.8], 8.2)
    doc.add_heading("2. 已生成文件", level=1)
    deliverables = [
        ["00_项目执行说明与阅读顺序.docx", "交付地图、标签与边界", "快速进入项目", "所有关键人"],
        ["00_待合作方补充的内部数据清单.docx", "SKU、价格、库存、利润、主体等21项", "补齐正式报价前置条件", "国代、品牌方"],
        ["01_执行前影响面与阻断检查.docx", "法律、平台、文化、来源检查", "确定可做与必须停", "负责人、法务"],
        ["02_汾酒尼泊尔深度市场调研报告.docx", "城市、竞争、渠道、进入策略", "市场决策", "决策人、销售"],
        ["03_尼泊尔生活习惯与消费文化报告.docx", "人群、场景、节庆、传统酒、配餐", "本地化与品鉴设计", "内容、销售"],
        ["04_中国文化植入与中尼文化桥接矩阵.docx", "本地事实到中国文化和汾酒表达的桥接", "避免文化误读", "内容、培训"],
        ["05_90天B2B与B2C市场执行手册.docx", "平台获客、12周动作、KPI", "现场执行", "执行负责人"],
        ["06_经销商发展与试销机制.docx", "四阶段合作、试销条款、评分", "从线索推进到补货", "渠道负责人"],
        ["07_TikTok及社交媒体推广策略.docx", "A/B/C方案、30选题、12脚本", "获批后的内容执行", "内容、合规"],
        ["07A_TikTok及社交媒体合规确认清单.docx", "律师、平台、国代书面确认问题", "解除传播阻断", "法务、平台负责人"],
        ["08_AI与人工最小介入SOP.docx", "Agent分工、人工触发、数据闸门", "日常协同", "运营、AI协作人员"],
        ["09_商务外联话术与跟进模板.docx", "D1/D3/D7/D14与异议处理", "真实B2B接触", "本地销售"],
        ["10_渠道商及终端线索数据库.xlsx", "383条公开线索与双源队列", "筛选、验证、外联", "渠道执行人员"],
        ["11_城市_渠道商_终端_创作者评分表.xlsx", "四类100分模型", "优先级排序", "项目负责人"],
        ["12_90天CRM执行看板.xlsx", "Top50队列与转化KPI", "跟进与复盘", "销售、负责人"],
        ["13_短视频内容日历与脚本库.xlsx", "30选题、12脚本、12周日历", "获批后拍摄与测试", "内容团队"],
        ["14_困难_风险_解决方案矩阵.xlsx", "风险、信号、处理与备用方案", "风险管理", "负责人、法务"],
        ["15_人工介入队列.xlsx", "只需人工处理的现实动作", "最小人工协作", "国代、本地执行"],
        ["16_来源与证据索引.xlsx", "去重来源、别名、结论与访问状态", "证据追溯", "审阅者、研究人员"],
        ["17_机器可读取的渠道数据.json", "城市、平台、竞品、线索、来源结构化数据", "系统接入与再计算", "数据/AI人员"],
        ["18_Claude复核任务包.docx", "P0-P3独立复核任务", "Claude可用时补做反方复核", "外部复核人员"],
        ["19_最终执行报告.docx", "结论、数据、状态、人工事项、未来7天", "决策与启动", "所有关键人"],
    ]
    add_table(doc, ["文件", "内容", "用途", "推荐阅读者"], deliverables, [2.45, 1.55, 1.45, 1.05], 7.1)
    doc.add_heading("3. 关键数据", level=1)
    verified = sum(1 for x in LEADS if str(x.get("current_status") or x.get("status") or x.get("verification_status") or "").lower() in ("已验证", "verified"))
    dual = sum(1 for x in LEADS if str(x.get("priority", "")).lower() == "high")
    categories = {}
    for x in LEADS:
        cat = x.get("customer_type") or x.get("category") or x.get("type") or x.get("category_group") or "未分类"
        categories[str(cat)] = categories.get(str(cat), 0) + 1
    add_table(doc, ["指标", "值", "状态解释"], [
        ["线索总数", len(LEADS), "真实公开来源；不足300不补假数据"],
        ["显式已验证", verified, "仅按数据状态统计"],
        ["独立域名双源高优", dual, "两个独立域名且评分达标；不等于已联系或业务验证"],
        ["来源总数", len(ALL_SOURCES), "去重后的证据索引"],
        ["城市评分", " / ".join(f"{x.get('city')} {x.get('total')}" for x in (channels.get("city_scoring") or root.get("city_scores", []))), "COMPUTED、需真实销售修正"],
        ["平台状态", "TikTok可访问；酒类公开推广C", "动态状态，执行前复查"],
    ], [1.35, 2.1, 3.05], 8.2)
    if categories:
        doc.add_heading("线索类别分布", level=2)
        add_table(doc, ["类别", "数量"], sorted(categories.items(), key=lambda x: (-x[1], x[0])), [5.2, 1.3], 8.4)
    add_callout(doc, "数据库覆盖缺口", "当前公开线索以酒类零售、餐饮、酒店/度假为主；经销商/批发商、企业礼赠采购、中国企业等类别尚未形成可核验样本。不得把类别缺失写成市场不存在，需通过国代名录、商会和人工引荐补充。", PALE_GOLD)
    doc.add_heading("4. 判断标签", level=1)
    add_bullets(doc, [
        "KNOWN：国代角色边界；尼泊尔法律第45条；TikTok可访问；在线酒类零售和国内/进口烈酒竞争存在。",
        "INFERRED：Kathmandu首城、前三渠道、小杯食物品鉴、B2B 70%起步。",
        "COMPUTED：城市评分、线索/来源统计、90天KPI和扩城阈值。",
        "NEEDS_VERIFY：SKU、价格、利润、库存、配送、样品、平台入驻、Nepali、食物搭配、消费者接受、决策人。",
        "BLOCKED：公开酒类推广/付费广告、未经审核Nepali、正式报价和ROI、独家承诺。",
    ])
    doc.add_heading("5. 人工只需处理的事项", level=1)
    add_table(doc, ["优先", "为什么必须人工", "AI已完成", "人工下一步", "完成定义"], [
        [1, "法律/平台必须有责任主体", "法规与问题清单", "律师/平台/国代书面签字", "每个内容/商品/品鉴场景有明确允许/禁止"],
        [2, "商业数据不公开", "内部数据表", "国代填SKU/价/利/库存/样品/配送", "可出正式一页试销报价"],
        [3, "本地语言与文化", "English+Nepali草稿/桥接矩阵", "母语者逐条审", "无歧义、无冒犯、可公开/外联"],
        [4, "电话、见面、品鉴和谈判", "Top50、话术、表单", "对Top20发送并约2-3场", "CRM有真实回复/拒绝/预约"],
        [5, "口感和搭配只能真实测试", "候选搭配和问卷", "成年样本盲测", "30-50份结构化反馈"],
    ], [0.45, 1.3, 1.25, 1.65, 1.85], 7.5)
    doc.add_heading("6. 未来7天", level=1)
    add_numbered(doc, ["完成合作方内部数据表。", "取得s45和各平台书面合规意见。", "数据库筛Top50，双源/评分/下一步齐全。", "首批20家D1联系；D3跟进计划进CRM。", "安排2-3场持牌、成年、小杯品鉴。", "询问4个平台入驻/抽佣/年龄/内容规则。", "周末输出首轮回复、预约、拒绝原因和风险周报。"])
    doc.add_heading("7. 完成状态", level=1)
    add_bullets(doc, [
        "研究与文件生成：已生成；结构/渲染验证以交付时QA记录为准。",
        f"公开线索数据库：已生成 {len(LEADS)} 条；显式业务验证=0；独立域名双源高优={dual}。",
        "商业条件：NEEDS_VERIFY；本地法律/平台书面批准：BLOCKED。",
        "人工外联、品鉴、试销、补货和业务通过：尚未发生。",
        "Claude复核：未执行；已提供Word任务包。",
    ])
    save(doc, "19_最终执行报告.docx")


def main():
    doc_reading_order()
    doc_internal_inputs()
    doc_impact()
    doc_market()
    doc_culture()
    doc_bridge()
    doc_90day()
    doc_distributor()
    doc_social()
    doc_compliance_checklist()
    doc_ai_sop()
    doc_outreach()
    doc_claude_pack()
    doc_final()
    print(f"Generated 14 DOCX files in {OUT}")


if __name__ == "__main__":
    main()
